"""One reader model for verified Climate-FCV web and DOCX outputs."""

from __future__ import annotations

import html
import re
import unicodedata
from urllib.parse import urlsplit
from pathlib import Path
from typing import Any

from docx import Document

from climate_question_bank import CLIMATE_LITERATURE_REFERENCES
from sector_lenses.climate_judgments import ALLOWED
from sector_lenses.climate_verified_schemas import (
    validate_summary_fragment,
    validate_summary_overview,
)


DIMENSIONS = (
    (
        "relevance",
        "Climate-FCV relevance",
        "How strongly do climate and FCV pressures intersect in this "
        "project's context and objectives?",
    ),
    (
        "sensitivity",
        "Climate & FCV sensitivity",
        "Is the project designed to recognise and avoid worsening climate- "
        "and FCV-related risks (do no harm)?",
    ),
    (
        "responsiveness",
        "Climate & FCV responsiveness",
        "Does the project actively build climate resilience and address the "
        "drivers of fragility, conflict, and violence?",
    ),
    (
        "operationalization",
        "From intent to delivery (operationalization)",
        "Are these intentions turned into concrete requirements, roles, "
        "indicators, and triggers the team can act on?",
    ),
)
NO_RECOMMENDATION_MESSAGE = (
    "No operational priorities were identified in this assessment. Review the "
    "core questions and points to check below."
)
INCOMPLETE_RECOMMENDATION_MESSAGE = (
    "The recommendation stage could not be completed because every generated "
    "recommendation failed the automated checks. Do not treat this "
    "Recommendations Note as complete; rerun the analysis or contact support."
)
# Headline rating shown above the core-question cards: how sensitive the project
# is to climate and FCV considerations, on a Limited -> Moderate -> Strong scale
# derived from the (retained) internal `sensitivity` judgment. Higher is better
# (more strongly designed to recognise and avoid worsening climate/FCV risk).
_SENSITIVITY_SCALE_LABELS = (
    "Very Limited", "Limited", "Moderate", "Strong", "Very Strong",
)
SENSITIVITY_RATING_QUESTION = (
    "How sensitive is this project to climate and FCV considerations?"
)
SENSITIVITY_RATING_CAVEAT = (
    "This is a subjective judgement on the part of this AI tool and does not "
    "constitute an official WBG rating."
)
_SENSITIVITY_RATING = {
    "very_strong": {
        "label": "Very Strong", "level": 5, "tone": "good",
        "description": "The project is very strongly designed to recognise climate "
        "and conflict risks and to avoid making them worse.",
    },
    "strong": {
        "label": "Strong", "level": 4, "tone": "good",
        "description": "The project is strongly designed to recognise climate and "
        "conflict risks and to avoid making them worse.",
    },
    "moderate": {
        "label": "Moderate", "level": 3, "tone": "mid",
        "description": "The project recognises several climate and conflict risks, "
        "but gaps remain in how it avoids making them worse.",
    },
    "limited": {
        "label": "Limited", "level": 2, "tone": "low",
        "description": "The project shows limited attention to recognising climate "
        "and conflict risks and avoiding harm - an area to strengthen.",
    },
    "very_limited": {
        "label": "Very Limited", "level": 1, "tone": "low",
        "description": "The project shows very limited attention to recognising "
        "climate and conflict risks and avoiding harm - a priority to strengthen.",
    },
    "unclear": {
        "label": "Not yet clear", "level": 0, "tone": "unclear",
        "description": "There is not yet enough in the document to judge how well "
        "the project recognises climate and conflict risks and avoids harm.",
    },
}
CORE_QUESTIONS_INTRO = (
    "This section works through the core questions that the World Bank's "
    "climate-and-fragility guidance asks of a project in a conflict-affected "
    "setting. Each answer draws only on your document and the evidence gathered "
    "for this run, and is kept separate from the executive summary above so it "
    "adds new detail rather than repeating it."
)


HEADINGS = (
    "Executive readout",
    "Core climate-FCV questions",
    "Ranked operational priorities",
    "Points to check before the decision meeting",
    "What to keep an eye on",
)
# Lay-comprehensible intro shared by the points-to-check section across all three
# render surfaces (server HTML, DOCX, and the frontend renderer).
POINTS_TO_CHECK_INTRO = (
    "These are smaller things for the team to confirm or consider before the concept "
    "or decision meeting. They are not the main recommendations above, and none is a "
    "reason to stop the project - think of them as a checklist."
)
# Reader-facing priority rows only. Model-internal routing/authority fields and
# coded reference lists (routing_status, authority_basis, recommendation_basis,
# *_ids) remain available to validation but are excluded from visible surfaces.
PRIORITY_FIELDS = (
    ("Decision", "decision"),
    ("Minimum action", "minimum_action"),
    ("Enhanced action", "enhanced_action"),
    ("Activation condition", "enhanced_activation"),
    ("Who", "responsible_function"),
    ("Completion evidence", "completion_evidence"),
    ("Completion evidence status", "completion_evidence_status"),
    ("Confidence", "confidence"),
    ("Limitation", "limitation"),
    ("Caution", "caution"),
)
SMOKE_RUNTIME_WARNING = (
    "Smoke test: validates workflow completion only; "
    "not a quality benchmark."
)
CANDIDATE_PREVIEW_WARNING = (
    "Candidate country evidence: preview; not approved."
)
ADVISORY_NOTICE = (
    "This automated screening supports task-team judgment and does not "
    "constitute an institutional adequacy or compliance decision."
)
_PLACEHOLDER = re.compile(
    r"(?:\[\s*(?:tbd|todo|insert|placeholder)[^\]]*\]|"
    r"\b(?:tbd|todo|lorem ipsum)\b)",
    re.IGNORECASE,
)
_BARE_PLACEHOLDER = re.compile(
    r"^\s*placeholder(?:\s+text)?[.!]?\s*$",
    re.IGNORECASE,
)


def _scrub_placeholder_text(text: str) -> str:
    """Strip model-emitted placeholder cues from a single reader string.

    A single stray '[insert ...]'/'[tbd]'/bare 'tbd'/'todo' from the model must
    not hard-fail an entire (paid) run at the reader-integrity gate. Remove the
    cue and keep the surrounding prose, mirroring the deterministic
    vocabulary-scrub pattern used elsewhere in the app. Genuinely structural
    problems (empty required fields, truncation) still surface downstream.
    """
    if not text:
        return text
    # Normalise em/en dashes (and any mis-decoded replacement char) to ASCII
    # hyphens. Model prose sometimes uses em dashes which (a) mojibake to a stray
    # replacement character in the exported HTML/DOCX and (b) violate house style;
    # doing this in the shared scrub fixes every reader surface at once.
    text = text.replace("—", " - ").replace("–", "-").replace("�", "-")
    cleaned = _PLACEHOLDER.sub("", text)
    if _BARE_PLACEHOLDER.fullmatch(cleaned.strip()):
        cleaned = ""
    cleaned = re.sub(r"\(\s*\)", "", cleaned)            # empty parens left behind
    cleaned = re.sub(r"[ \t]+([.,;:!?])", r"\1", cleaned)  # space before punctuation
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)          # collapse space runs (keep newlines)
    return cleaned.strip()


def _scrub_placeholders(value: object) -> object:
    """Recursively scrub placeholder cues from every string leaf in a reader."""
    if isinstance(value, str):
        return _scrub_placeholder_text(value)
    if isinstance(value, dict):
        return {key: _scrub_placeholders(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_scrub_placeholders(item) for item in value]
    if isinstance(value, tuple):
        return tuple(_scrub_placeholders(item) for item in value)
    return value


def _mapping(value: object) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _records(value: object) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(
        value, list
    ) else []


def _text(value: object) -> str:
    return str(value or "").strip()


def _normalized_source_title(value: object) -> str:
    """Return a stable, punctuation-insensitive key for source matching."""

    normalized = unicodedata.normalize("NFKD", _text(value)).casefold()
    normalized = normalized.replace("&", " and ")
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", " ", ascii_text).strip()


def _is_public_world_bank_url(value: object) -> bool:
    """Accept only a fail-closed HTTPS authority on World Bank's DNS domain."""

    if not isinstance(value, str):
        return False
    try:
        parsed = urlsplit(value)
    except ValueError:
        return False
    authority = parsed.netloc
    if (
        parsed.scheme.casefold() != "https"
        or not authority
        or not authority.isascii()
        or any(character in authority for character in ("%", "@", ":", "[", "]"))
        or parsed.username is not None
        or parsed.password is not None
    ):
        return False
    hostname = authority.casefold()
    labels = hostname.split(".")
    if (
        len(hostname) > 253
        or any(
            not label
            or len(label) > 63
            or label.startswith("-")
            or label.endswith("-")
            or label.startswith("xn--")
            or not re.fullmatch(r"[a-z0-9-]+", label)
            for label in labels
        )
    ):
        return False
    return hostname == "worldbank.org" or hostname.endswith(".worldbank.org")


_SENTENCE_CLOSING_MARKS = "\"')]}\u201d\u2019"


def _complete_sentence(value: object) -> str:
    """Normalize a verified prose fragment into one complete sentence."""

    text = re.sub(r"\s+", " ", _text(value)).strip()
    if not text:
        return ""
    terminal_text = text.rstrip(_SENTENCE_CLOSING_MARKS)
    return text if terminal_text and terminal_text[-1] in ".!?" else f"{text}."


def _distinct_texts(values: list[object], limit: int) -> list[str]:
    """Keep non-empty strings in first-seen order, ignoring case/space repeats."""

    selected: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = _text(value)
        key = re.sub(r"\s+", " ", text).casefold()
        if not text or key in seen:
            continue
        seen.add(key)
        selected.append(text)
        if len(selected) >= limit:
            break
    return selected


def _deduplicated_questions(matches: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Deduplicate by stable ID, otherwise by normalized verified source content."""

    selected: list[dict[str, Any]] = []
    seen: set[tuple[str, ...]] = set()
    for question in matches:
        question_id = _text(question.get("question_id")).casefold()
        identity = (
            ("question_id", question_id)
            if question_id
            else (
                "content",
                _normalized_source_title(question.get("source")),
                re.sub(r"\s+", " ", _text(question.get("summary"))).casefold(),
                re.sub(r"\s+", " ", _text(question.get("watch"))).casefold(),
            )
        )
        if identity in seen:
            continue
        seen.add(identity)
        selected.append(question)
    return selected


def _guidance_project_use(matches: list[dict[str, Any]]) -> str:
    """Return one short project-specific follow-up from verified reader fields."""

    watches = _distinct_texts(
        [question.get("watch") for question in matches],
        limit=1,
    )
    if watches:
        return (
            "For this project, use the source to address this follow-up: "
            + _complete_sentence(watches[0])
        )
    questions = _distinct_texts(
        [question.get("question") for question in matches],
        limit=1,
    )
    if questions:
        return (
            "For this project, use the source to examine this question: "
            + _complete_sentence(questions[0])
        )
    return ""


def build_climate_guidance_items(
    core_questions: object,
    sources: object,
) -> list[dict[str, str]]:
    """Build ranked, matched-only project guidance from verified reader fields."""

    matched_by_source: dict[str, list[dict[str, Any]]] = {}
    for question in _records(core_questions):
        source_key = _normalized_source_title(question.get("source"))
        if source_key:
            matched_by_source.setdefault(source_key, []).append(question)

    ranked: list[tuple[int, int, dict[str, str]]] = []
    emitted_source_keys: set[str] = set()
    for catalog_order, source in enumerate(_records(sources)):
        title = _text(source.get("title"))
        source_key = _normalized_source_title(title)
        matches = _deduplicated_questions(matched_by_source.get(source_key, []))
        url = source.get("url")
        if (
            not title
            or source_key in emitted_source_keys
            or not matches
            or not _is_public_world_bank_url(url)
        ):
            continue
        project_use = _guidance_project_use(matches)
        if not project_use:
            continue
        practical_value = _text(source.get("practical_value")) or _text(source.get("description"))
        ranked.append((
            -len(matches),
            catalog_order,
            {
                "title": title,
                "url": str(url),
                "practical_value": practical_value,
                "project_use": project_use,
            },
        ))
        emitted_source_keys.add(source_key)
    return [item for _, _, item in sorted(ranked)[:4]]


def _field_text(value: object) -> str:
    if isinstance(value, (list, tuple)):
        return ", ".join(_text(item) for item in value if _text(item))
    return _text(value)


def _rank(value: object) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return 999


def _no_priority_message(model: dict[str, object]) -> str:
    """Return a bounded reader explanation without internal run diagnostics."""

    return _text(model.get("recommendation_message")) or NO_RECOMMENDATION_MESSAGE



_PRIORITY_TITLE_PREFIX = re.compile(
    r"^\s*Priority\s+\d+\s*(?:[-:.\u00b7\u2013\u2014\u2022]\s*)?",
    re.IGNORECASE,
)


def _normalize_priority_title(value: object) -> str:
    original = _text(value)
    normalized = _PRIORITY_TITLE_PREFIX.sub("", original).strip()
    return normalized or original


def _first_reader_text(*values: object) -> str:
    for value in values:
        text = _scrub_placeholder_text(_text(value))
        if text:
            return text
    return ""


def _project_cycle_for_operation(
    operation_context: dict[str, object], priority: dict[str, Any]
) -> dict[str, str]:
    document_type = _text(operation_context.get("document_type")).casefold()
    if document_type in {"pcn", "pid", "concept note", "project concept note"}:
        primary_label, secondary_label = "At concept stage", "During preparation"
    elif document_type in {"additional financing", "af"}:
        primary_label = "In the Additional Financing package"
        secondary_label = "Before approval"
    elif document_type == "restructuring":
        primary_label = "In the restructuring package"
        secondary_label = "During implementation"
    elif document_type in {
        "pad",
        "project appraisal document",
        "project paper",
        "program paper",
        "program document",
    }:
        primary_label = "In the current project document"
        secondary_label = "Before approval"
    else:
        primary_label = "At the current review stage"
        secondary_label = "Before the next decision point"
    return {
        "primary_label": primary_label,
        "primary_text": _first_reader_text(
            priority.get("minimum_action"), priority.get("decision")
        ),
        "secondary_label": secondary_label,
        "secondary_text": _first_reader_text(
            priority.get("completion_evidence"),
            priority.get("enhanced_action"),
            priority.get("limitation"),
        ),
    }

def _priority_summary(priorities: list[dict[str, Any]]) -> dict[str, object]:
    titles = [_text(item.get("title")) for item in priorities]
    count = len(titles)
    if count == 0:
        statement = (
            "No final operational priority was admitted after validation and "
            "semantic review."
        )
    else:
        noun = "priority" if count == 1 else "priorities"
        statement = (
            "Drawing on the overview and core climate-FCV questions, the analysis "
            f"identifies {count} main operational {noun} for strengthening climate "
            "resilience, conflict sensitivity and implementation readiness in this "
            "project. These are followed by secondary points to check before the "
            "decision meeting and issues to keep under review as preparation advances."
        )
    return {"count": count, "titles": titles, "statement": statement}


def build_reader_model(assessment: dict[str, object]) -> dict[str, object]:
    """Project a verified assessment into the only reader-facing structure."""

    analysis = _mapping(assessment.get("analysis"))
    existing_responses = [
        {
            "response_id": _text(response.get("response_id")),
            "project_fact_ids": [
                _text(item)
                for item in response.get("project_fact_ids", [])
                if _text(item)
            ] if isinstance(response.get("project_fact_ids"), (list, tuple)) else [],
            "pathway_ids": [
                _text(item)
                for item in response.get("pathway_ids", [])
                if _text(item)
            ] if isinstance(response.get("pathway_ids"), (list, tuple)) else [],
            "description": _text(response.get("description")),
            "limitation": _text(response.get("limitation")),
        }
        for response in _records(analysis.get("existing_responses"))[:12]
    ]
    raw_judgments = _mapping(assessment.get("judgments"))
    judgments = []
    for key, title, description in DIMENSIONS:
        judgment = _mapping(raw_judgments.get(key))
        judgments.append({
            "dimension": key,
            "title": title,
            "description": description,
            "value": _text(judgment.get("value")),
            "rationale": _text(judgment.get("rationale")),
            "evidence_ids": [
                _text(item)
                for item in judgment.get("evidence_ids", [])
                if _text(item)
            ] if isinstance(judgment.get("evidence_ids"), (list, tuple)) else [],
        })

    core_questions = [
        {
            "question_id": _text(q.get("question_id")),
            "theme": _text(q.get("theme")),
            "question": _text(q.get("question")),
            "source": _text(q.get("source")),
            "summary": _text(q.get("summary")),
            "evidence_ids": [
                _text(e) for e in q.get("evidence_ids", []) if _text(e)
            ] if isinstance(q.get("evidence_ids"), (list, tuple)) else [],
            "watch": _text(q.get("watch")),
        }
        for q in _records(assessment.get("core_questions"))
    ]
    by_dimension = {item["dimension"]: item for item in judgments}
    sensitivity_judgment = by_dimension.get("sensitivity", {})
    sensitivity_value = _text(sensitivity_judgment.get("value")).lower() or "unclear"
    rating = _SENSITIVITY_RATING.get(
        sensitivity_value, _SENSITIVITY_RATING["unclear"]
    )
    overview_summary = _text(assessment.get("overview_summary"))
    canonical_summary_text = {
        key: assessment.get(key)
        for key in (
            "facts", "derived_assertions", "analysis",
            "overview_summary", "core_questions", "context_evidence", "sources",
        )
    }
    # Use the normalized four-dimension reader records as canonical judgment
    # support, rather than trusting an unprojected saved-result mapping.
    canonical_summary_text["judgments"] = judgments
    fallback_canonical_text = {
        key: assessment.get(key)
        for key in (
            "facts", "derived_assertions", "analysis", "overview_summary",
            "core_questions", "context_evidence", "sources",
        )
    }
    summary_overview = _summary_overview_paragraphs(
        _mapping(assessment.get("summary_overview")),
        executive_readout=assessment.get("executive_readout"),
        canonical_text=canonical_summary_text,
    )
    summary_overview_generated = bool(summary_overview)
    if not summary_overview:
        summary_overview = _legacy_summary_overview(
            overview_summary,
            judgments,
            _mapping(assessment.get("validation")).get("reason_codes", ()),
            canonical_text=fallback_canonical_text,
            known_evidence_ids=_summary_known_evidence_ids(assessment),
        )
    summary_overview_status = _text(assessment.get("summary_overview_status"))
    if not summary_overview_generated:
        summary_overview_status = "fallback"
    elif summary_overview_status not in {"generated", "fallback"}:
        summary_overview_status = "generated" if summary_overview else "fallback"
    climate_sensitivity_rating = {
        "value": sensitivity_value,
        "label": rating["label"],
        "level": rating["level"],
        "tone": rating["tone"],
        "scale": list(_SENSITIVITY_SCALE_LABELS),
        "question": SENSITIVITY_RATING_QUESTION,
        # 3-4 sentence plain-language overall summary embedded in the top overview
        # block; empty for older/blank runs, in which case the card shows only the
        # level gloss (graceful degradation).
        "overview_summary": overview_summary,
        "description": rating["description"],
        "caveat": SENSITIVITY_RATING_CAVEAT,
        "evidence_ids": [
            _text(e) for e in sensitivity_judgment.get("evidence_ids", []) if _text(e)
        ] if isinstance(sensitivity_judgment.get("evidence_ids"), (list, tuple)) else [],
    }

    priorities = sorted(
        _records(assessment.get("priorities")),
        key=lambda item: (_rank(item.get("rank")), _text(item.get("title"))),
    )[:5]
    flags = _records(assessment.get("review_readiness_flags"))[:4]
    minor_climate_points = [
        {
            "point": _text(item.get("point")),
            "why": _text(item.get("why")),
            "how_to_check": _text(item.get("how_to_check")),
            "residual_gap_ids": [
                _text(g) for g in item.get("residual_gap_ids", []) if _text(g)
            ] if isinstance(item.get("residual_gap_ids"), (list, tuple)) else [],
        }
        for item in _records(assessment.get("minor_climate_points"))
    ][:3]
    validation = _mapping(assessment.get("validation"))
    diagnostics = _mapping(assessment.get("recommendation_diagnostics"))
    diagnostic_reasons = diagnostics.get("reason_codes", [])
    all_suppressed = (
        diagnostics.get("parsed_candidate_count", 0) > 0
        and diagnostics.get("final_priority_count", 0) == 0
        and isinstance(diagnostic_reasons, (list, tuple))
        and "RECOMMENDATIONS_ALL_SUPPRESSED" in diagnostic_reasons
    )
    recommendation_status = "incomplete" if all_suppressed else "complete"
    recommendation_message = (
        INCOMPLETE_RECOMMENDATION_MESSAGE
        if all_suppressed
        else NO_RECOMMENDATION_MESSAGE if not priorities else ""
    )
    executive = _text(assessment.get("executive_readout")) or _text(
        assessment.get("judgment_summary")
    )
    raw_operation_context = _mapping(assessment.get("operation_context"))
    operation_context = {
        key: raw_operation_context.get(key)
        for key in (
            "document_type",
            "instrument_type",
            "country_scope",
            "is_mpa",
            "has_ipf_component",
            "preparation_regime",
            "processing_model",
            "es_regime",
            "warning_codes",
        )
    } if raw_operation_context else {}
    repair_actions = _mapping(assessment.get("manifest")).get(
        "repair_actions", []
    )
    drafting_withheld = (
        _text(operation_context.get("instrument_type")).casefold()
        in {"", "unknown"}
        and isinstance(repair_actions, (list, tuple))
        and "DRAFTING_CURRENT_UNRESOLVED_ROUTE_DROPPED" in repair_actions
    )
    reader_priorities = []
    for priority in priorities:
        reader_priority = dict(priority)
        reader_priority["title"] = _normalize_priority_title(priority.get("title"))
        reader_priority["project_cycle"] = _project_cycle_for_operation(
            operation_context, reader_priority
        )
        reader_priorities.append(reader_priority)
    return _scrub_placeholders({
        "executive_readout": executive,
        "operation_context": operation_context,
        "drafting_route_status": (
            "withheld_unresolved_instrument" if drafting_withheld else "available"
        ),
        "overview_summary": overview_summary,
        "judgments": judgments,
        "summary_overview": summary_overview,
        "summary_overview_status": summary_overview_status,
        "climate_sensitivity_rating": climate_sensitivity_rating,
        "core_questions": core_questions,
        "existing_responses": existing_responses,
        "priorities": reader_priorities,
        "recommendation_status": recommendation_status,
        "recommendation_message": recommendation_message,
        "review_readiness_flags": [dict(item) for item in flags],
        "minor_climate_points": minor_climate_points,
        "priority_summary": _priority_summary(reader_priorities),
        "evidence_status": _text(assessment.get("evidence_status")) or "approved",
        "technical_annex": {
            "run_id": _text(assessment.get("run_id")),
            "schema_version": _text(assessment.get("schema_version")),
            "bank_release_id": _text(assessment.get("bank_release_id")),
            "validation_status": _text(validation.get("status")),
            "summary_overview_status": summary_overview_status,
            "recommendation_candidate_count": diagnostics.get(
                "raw_candidate_count", 0
            ),
            "recommendation_admitted_count": diagnostics.get(
                "admitted_count", 0
            ),
            "recommendation_final_count": diagnostics.get(
                "final_priority_count", 0
            ),
            "semantic_reviewer_invoked": diagnostics.get(
                "reviewer_invoked", False
            ),
            "live_research_count": _mapping(
                assessment.get("manifest")
            ).get("live_research_count", 0),
            "semantic_reviewer_verdict": _text(
                diagnostics.get("reviewer_verdict")
            ) or "not_invoked",
            "recommendation_reason_codes": diagnostics.get("reason_codes", []),
            "unsupported_numeric_tokens": diagnostics.get(
                "unsupported_numeric_tokens", []
            ),
            "semantic_review_object_ids": [
                _text(item)
                for item in diagnostics.get("semantic_review_object_ids", [])
                if _text(item)
            ][:12] if isinstance(
                diagnostics.get("semantic_review_object_ids"), list
            ) else [],
            "candidate_suppressions": _records(
                diagnostics.get("candidate_suppressions")
            )[:3],
        },
        "advisory_notice": ADVISORY_NOTICE,
    })


_METHODOLOGY_NOTE = (
    "Here is how this analysis was put together, in plain terms. The tool first "
    "reads your document and pulls out the concrete facts it can find - what the "
    "project will do, where, for whom, and what safeguards it already includes. "
    "It then looks at how climate pressures and conflict or fragility affect each "
    "other in this setting, and checks every finding against those facts. It only "
    "offers a recommendation when the evidence clearly supports one, and a second "
    "automated check removes anything that over-reaches. Nothing here is invented: "
    "each point is tied to specific project evidence retained by the validation "
    "pipeline, even though internal reference codes are not shown in the reader."
)

_ID_TYPE_LABELS = {
    "PF": "Project fact", "RG": "Residual gap", "ER": "Existing response",
    "PW": "Climate-FCV pathway", "CE-LIVE": "Live research", "CE": "Context evidence",
}


def _id_type(identifier: str) -> str:
    if identifier.upper().startswith("CE-LIVE"):
        return "CE-LIVE"
    prefix = identifier.split("-", 1)[0].upper()
    return prefix


def _chain_prose(chain: list) -> str:
    parts = [_text(c) for c in chain if _text(c)]
    if len(parts) >= 3:
        middle = ", ".join(parts[1:-1])
        return f"{parts[0]}, leading through {middle}, to {parts[-1]}."
    return "; ".join(parts) + ("." if parts else "")


def build_evidence_trail(assessment: dict[str, object]) -> dict[str, object]:
    """Project a plain-language evidence trail from the raw verified assessment.

    Deterministic; resolves only IDs actually cited in judgments/priorities.
    Resolution priority for a cited ID: facts > residual gaps > existing
    responses > pathways, then CE-LIVE / CE stubs, then "reference not resolved".
    """
    analysis = _mapping(assessment.get("analysis"))
    facts = {_text(f.get("claim_id")): f for f in _records(assessment.get("facts"))}
    responses = {_text(r.get("response_id")): r
                 for r in _records(analysis.get("existing_responses"))}
    gaps = {_text(g.get("gap_id")): g
            for g in _records(analysis.get("residual_gaps"))}
    pathways_raw = _records(analysis.get("pathways"))
    pathways_by_id = {_text(p.get("pathway_id")): p for p in pathways_raw}

    pathways = []
    for p in pathways_raw:
        chain_prose = _chain_prose(
            p.get("chain") if isinstance(p.get("chain"), list) else []
        )
        if not chain_prose:
            # A pathway with no chain prose (e.g. a thin model run) would render
            # as a naked "Climate -> FCV:" label; drop it rather than show a stub.
            continue
        direction = _text(p.get("direction"))
        label = ("Climate -> FCV" if direction == "climate_to_fcv"
                 else "FCV -> Climate" if direction == "fcv_to_climate"
                 else direction or "Pathway")
        pathways.append({
            "direction_label": label,
            "chain_prose": chain_prose,
            "anchor_ids": [_text(a) for a in p.get("project_anchor_ids", []) if _text(a)]
            if isinstance(p.get("project_anchor_ids"), (list, tuple)) else [],
        })

    cited: list[str] = []
    judgments = _mapping(assessment.get("judgments"))
    for value in judgments.values():
        jm = _mapping(value)
        eids = jm.get("evidence_ids")
        if isinstance(eids, (list, tuple)):
            for eid in eids:
                if _text(eid):
                    cited.append(_text(eid))
    for pr in _records(assessment.get("priorities")):
        for field in ("project_anchor_ids", "pathway_ids", "existing_response_ids",
                      "residual_gap_ids", "instrument_claim_ids"):
            vals = pr.get(field)
            if isinstance(vals, (list, tuple)):
                cited.extend(_text(v) for v in vals if _text(v))

    seen: set[str] = set()
    evidence_key = []
    for cid in cited:
        if cid in seen:
            continue
        seen.add(cid)
        t = _id_type(cid)
        label = _ID_TYPE_LABELS.get(t, "Reference")
        if cid in facts:
            f = facts[cid]
            text = " ".join(
                x for x in (_text(f.get("subject")), _text(f.get("predicate")),
                             _text(f.get("object"))) if x
            ) or _text(f.get("supporting_excerpt"))
        elif cid in gaps:
            text = _text(gaps[cid].get("statement"))
        elif cid in responses:
            text = _text(responses[cid].get("description"))
        elif cid in pathways_by_id:
            text = _chain_prose(
                pathways_by_id[cid].get("chain")
                if isinstance(pathways_by_id[cid].get("chain"), list) else []
            )
        elif t == "CE-LIVE":
            text = "Accepted live-research evidence for this run."
        elif t == "CE":
            text = "Context evidence cited for this run (summary not stored)."
        else:
            text = "Reference not resolved."
        if not _text(text):
            # Skip entries that resolve to nothing (e.g. a cited pathway whose
            # chain came back empty) rather than showing a bare code with no text.
            continue
        evidence_key.append({"id": cid, "type_label": label, "text": text})

    diag = _mapping(assessment.get("recommendation_diagnostics"))
    manifest = _mapping(assessment.get("manifest"))
    diagnostics = {
        "candidate_count": diag.get("raw_candidate_count", 0),
        "admitted_count": diag.get("admitted_count", 0),
        "final_count": diag.get("final_priority_count", 0),
        "reviewer_verdict": _text(diag.get("reviewer_verdict")) or "not_invoked",
        "live_research_count": manifest.get("live_research_count", 0),
        "bank_release": _text(assessment.get("bank_release_id")),
        "evidence_status": _text(assessment.get("evidence_status")),
    }
    limitations_raw = analysis.get("evidence_limitations")
    if isinstance(limitations_raw, (list, tuple)):
        limitations = " ".join(
            _text(item) for item in limitations_raw if _text(item)
        )
    else:
        limitations = _text(limitations_raw)

    return {
        "methodology_note": _METHODOLOGY_NOTE,
        "pathways": pathways,
        "limitations": limitations,
        "evidence_key": evidence_key,
        "diagnostics": diagnostics,
    }


def attach_provenance(reader: dict[str, object], assessment: dict[str, object]) -> dict[str, object]:
    """Attach the evidence trail and static literature references to a reader.

    Called after build_reader_model/validate_reader_model so it is additive and
    never affects reader-integrity validation.
    """
    # Scrub dashes/placeholders here too: attach_provenance runs AFTER
    # build_reader_model's _scrub_placeholders pass, so model-generated
    # evidence-trail text and source descriptions would otherwise keep em/en
    # dashes (house style is ASCII hyphens) and any stray placeholder cue.
    reader["evidence_trail"] = _scrub_placeholders(build_evidence_trail(assessment))
    reader["sources"] = _scrub_placeholders(
        [dict(entry) for entry in CLIMATE_LITERATURE_REFERENCES]
    )
    reader["guidance_items"] = build_climate_guidance_items(reader.get("core_questions"), reader.get("sources"))
    return reader


def _all_strings(value: object):
    if isinstance(value, str):
        yield value
    elif isinstance(value, dict):
        for item in value.values():
            yield from _all_strings(item)
    elif isinstance(value, list):
        for item in value:
            yield from _all_strings(item)


def validate_reader_model(model: dict[str, object]) -> tuple[str, ...]:
    """Return deterministic reader-integrity reason codes."""

    issues: list[str] = []
    executive = _text(model.get("executive_readout"))
    if executive:
        words = len(executive.split())
        if not 300 <= words <= 900:
            issues.append("EXECUTIVE_LENGTH_INVALID")
        if executive[-1:] not in ".?!":
            issues.append("EXECUTIVE_SENTENCE_INCOMPLETE")

    judgments = _records(model.get("judgments"))
    if [item.get("dimension") for item in judgments] != [
        key for key, _, _ in DIMENSIONS
    ]:
        issues.append("JUDGMENT_DIMENSIONS_INCOMPLETE")
    for judgment in judgments:
        dimension = _text(judgment.get("dimension"))
        if _text(judgment.get("value")) not in ALLOWED.get(dimension, set()):
            issues.append("JUDGMENT_VALUE_INVALID")
        rationale = _text(judgment.get("rationale"))
        if not rationale or rationale[-1:] not in ".?!":
            issues.append("JUDGMENT_RATIONALE_INCOMPLETE")

    priorities = _records(model.get("priorities"))
    ranks = [_rank(item.get("rank")) for item in priorities]
    if ranks != list(range(1, len(priorities) + 1)):
        issues.append("PRIORITY_RANK_ORDER_INVALID")
    titles = [_text(item.get("title")).casefold() for item in priorities]
    if len(titles) != len(set(titles)):
        issues.append("DUPLICATE_PRIORITY_TITLE")
    required = (
        "recommendation_id",
        "title",
        "decision",
        "minimum_action",
        "confidence",
    )
    operation_context = _mapping(model.get("operation_context"))
    drafting_withheld = (
        _text(model.get("drafting_route_status"))
        == "withheld_unresolved_instrument"
        and _text(operation_context.get("instrument_type")).casefold()
        in {"", "unknown"}
    )
    drafting_required = (
        "target_document",
        "target_section",
        "drafting_status",
        "text",
        "project_basis_ids",
        "gap_basis_ids",
        "guidance_ids",
    )
    for priority in priorities:
        current = _mapping(priority.get("current_document_drafting"))
        if (
            not current or any(key not in current for key in drafting_required)
        ) and not drafting_withheld:
            issues.append("CURRENT_DRAFTING_INCOMPLETE")
        optional_value = priority.get("operational_instrument_drafting")
        if optional_value is not None and not isinstance(optional_value, dict):
            issues.append("OPERATIONAL_DRAFTING_MALFORMED")
        for block in (current, _mapping(optional_value)):
            if not block:
                continue
            draft_text = _text(block.get("text"))
            if not draft_text or draft_text[-1:] not in ".?!":
                issues.append("DRAFTING_TEXT_INCOMPLETE")

    if any(not all(_text(item.get(key)) for key in required) for item in priorities):
        issues.append("PRIORITY_FIELD_INCOMPLETE")
    summary = _mapping(model.get("priority_summary"))
    summary_titles = summary.get("titles")
    if (
        summary.get("count") != len(priorities)
        or summary_titles != [_text(item.get("title")) for item in priorities]
        or not _text(summary.get("statement"))
    ):
        issues.append("PRIORITY_SUMMARY_MISMATCH")


    if any(
        _PLACEHOLDER.search(value) or _BARE_PLACEHOLDER.fullmatch(value)
        for value in _all_strings(model)
    ):
        issues.append("UNRESOLVED_PLACEHOLDER")
    if any(
        value.rstrip().endswith(("[", "{", "…"))
        for value in _all_strings(model)
        if value.strip()
    ):
        issues.append("TRUNCATED_READER_FIELD")
    return tuple(dict.fromkeys(issues))


def _heading(level: int, text: str) -> str:
    return f"<h{level}>{html.escape(text)}</h{level}>"


def _drafting_html(label: str, value: object) -> str:
    block = _mapping(value)
    if not block:
        return ""
    parts = ['<section class="climate-drafting-block">']
    parts.append(_heading(4, label))
    for field_label, key in (
        ("Target document", "target_document"),
        ("Target section", "target_section"),
        ("Drafting status", "drafting_status"),
        ("Guidance basis", "guidance_ids"),
    ):
        field_value = _field_text(block.get(key))
        if field_value:
            parts.append(
                f"<p><strong>{html.escape(field_label)}:</strong> "
                f"{html.escape(field_value)}</p>"
            )
    parts.append(
        '<div class="climate-drafting-text">'
        + html.escape(_text(block.get("text")))
        + "</div>"
    )
    parts.append("</section>")
    return "".join(parts)



def _project_cycle_html(value: object) -> str:
    cycle = _mapping(value)
    primary_label = _text(cycle.get("primary_label"))
    primary_text = _text(cycle.get("primary_text"))
    if not primary_label or not primary_text:
        return ""
    parts = ['<section class="climate-project-cycle">']
    parts.append(_heading(4, "Where this fits in the project cycle"))
    parts.append(
        f"<p><strong>{html.escape(primary_label)}:</strong> "
        f"{html.escape(primary_text)}</p>"
    )
    secondary_label = _text(cycle.get("secondary_label"))
    secondary_text = _text(cycle.get("secondary_text"))
    if secondary_label and secondary_text:
        parts.append(
            f"<p><strong>{html.escape(secondary_label)}:</strong> "
            f"{html.escape(secondary_text)}</p>"
        )
    parts.append("</section>")
    return "".join(parts)


_RATING_TONE_COLORS = {
    "good": "#1A9850", "mid": "#E8A33D", "low": "#D73027", "unclear": "#9aa4b2",
}


def _sensitivity_rating_html(rating: dict[str, object]) -> str:
    """Render the headline climate & FCV sensitivity rating as a scale."""
    level = _rank(rating.get("level"))
    tone = _text(rating.get("tone")) or "unclear"
    active = _RATING_TONE_COLORS.get(tone, "#9aa4b2")
    scale = rating.get("scale") if isinstance(rating.get("scale"), list) else []
    segments = []
    for index, label in enumerate(scale, start=1):
        is_active = index == level
        bg = active if is_active else "#EEF0F3"
        color = "#fff" if is_active else "#6b7280"
        weight = "700" if is_active else "400"
        segments.append(
            f'<span style="flex:1;text-align:center;padding:6px 4px;background:{bg};'
            f'color:{color};font-weight:{weight};font-size:12px">'
            f"{html.escape(_text(label))}</span>"
        )
    scale_html = (
        '<div style="display:flex;gap:3px;margin:7px 0 9px;border-radius:6px;'
        'overflow:hidden;max-width:360px">' + "".join(segments) + "</div>"
    )
    summary = _text(rating.get("overview_summary"))
    summary_html = (
        f'<p style="margin:0 0 10px;font-size:14px;line-height:1.5">'
        f"{html.escape(summary)}</p>"
        if summary
        else ""
    )
    return (
        '<section class="climate-overview-panel climate-sens-rating" '
        'style="background:#fff;border:1px solid #D7E1E7;border-left:5px solid '
        f'{active};border-radius:10px;padding:18px 20px;margin:0 0 24px">'
        # Graphic first: the "How sensitive" question, rating label, and scale.
        + f'<p style="margin:0 0 2px"><strong>'
        f'{html.escape(_text(rating.get("question")))}</strong></p>'
        f'<p style="margin:0;font-size:18px;font-weight:700;color:{active}">'
        f'{html.escape(_text(rating.get("label")))}</p>'
        + scale_html
        + f'<p style="margin:0 0 4px">{html.escape(_text(rating.get("description")))}</p>'
        # Then the overall summary text below the graphic.
        + summary_html
        + '<p style="margin:6px 0 0;font-size:12px;color:#6b7280">'
        + html.escape(_text(rating.get("caveat")))
        + "</p></section>"
    )


def render_reader_html(model: dict[str, object]) -> str:
    """Render escaped HTML from the canonical reader dictionary."""

    parts = ['<article class="climate-verified-assessment">']
    if _text(model.get("runtime_mode")) == "smoke":
        parts.append(
            '<p class="climate-smoke-warning">'
            + html.escape(SMOKE_RUNTIME_WARNING)
            + "</p>"
        )
    if _text(model.get("evidence_status")) == "preview; not approved":
        parts.append(
            '<p class="climate-preview-warning">'
            + html.escape(CANDIDATE_PREVIEW_WARNING)
            + "</p>"
        )
    operation_context = _mapping(model.get("operation_context"))
    if operation_context:
        instrument = _text(operation_context.get("instrument_type")) or "Unknown"
        document_type = _text(operation_context.get("document_type")) or "Unknown"
        preparation = (
            _text(operation_context.get("preparation_regime"))
            or "unresolved_policy_source"
        ).replace("_", " ")
        es_regime = (
            _text(operation_context.get("es_regime")) or "UNRESOLVED"
        ).replace("_", " ")
        mpa_label = "MPA program" if operation_context.get("is_mpa") else "Not identified as MPA"
        parts.append(
            '<section class="climate-operation-context"><h2>'
            "How this operation was routed</h2><dl>"
            f"<div><dt>Instrument</dt><dd>{html.escape(instrument)}</dd></div>"
            f"<div><dt>Document</dt><dd>{html.escape(document_type)}</dd></div>"
            f"<div><dt>Preparation</dt><dd>{html.escape(preparation)}</dd></div>"
            f"<div><dt>E&amp;S route</dt><dd>{html.escape(es_regime)}</dd></div>"
            f"<div><dt>Program layer</dt><dd>{html.escape(mpa_label)}</dd></div>"
            "</dl>"
        )
        if instrument.casefold() == "unknown" or document_type.casefold() == "unknown":
            parts.append(
                "<p>Operational context could not be resolved safely, so "
                "document-targeted guidance was withheld.</p>"
            )
        parts.append("</section>")
    # Overview at the very top: the headline sensitivity rating card carries the
    # 3-4 sentence plain-language overall summary, so the reader gets the whole
    # takeaway up front. The fuller Executive readout follows as detail below.
    rating = _mapping(model.get("climate_sensitivity_rating"))
    if rating:
        parts.append(_sensitivity_rating_html(rating))

    parts.append(_heading(2, HEADINGS[0]))
    for _exec_para in re.split(r"\n\s*\n+", _text(model.get("executive_readout")).strip()):
        _exec_para = _exec_para.strip()
        if not _exec_para:
            continue
        _m = re.match(r"^(.*?[.!?])(\s+)([\s\S]*)$", _exec_para)
        if _m:
            parts.append(
                f"<p><strong>{html.escape(_m.group(1))}</strong>"
                f"{html.escape(_m.group(2) + _m.group(3))}</p>"
            )
        else:
            parts.append(f"<p><strong>{html.escape(_exec_para)}</strong></p>")

    parts.append(_heading(2, HEADINGS[1]))
    parts.append(f"<p>{html.escape(CORE_QUESTIONS_INTRO)}</p>")
    for question in _records(model.get("core_questions")):
        parts.append('<section class="climate-core-question">')
        parts.append(_heading(3, _text(question.get("question"))))
        source = _text(question.get("source"))
        if source:
            parts.append(
                '<p class="climate-core-source">For further insights on why this '
                "matters, see: <em>" + html.escape(source) + "</em></p>"
            )
        for para in re.split(r"\n\s*\n+", _text(question.get("summary")).strip()):
            para = para.strip()
            if para:
                parts.append("<p>" + html.escape(para) + "</p>")
        parts.append("</section>")

    parts.append(_heading(2, HEADINGS[2]))
    priorities = _records(model.get("priorities"))
    priority_summary = _mapping(model.get("priority_summary"))
    if priorities and _text(priority_summary.get("statement")):
        parts.append(
            '<p class="climate-priority-summary">'
            + html.escape(_text(priority_summary.get("statement")))
            + "</p>"
        )
    if not priorities:
        parts.append(f"<p>{html.escape(_no_priority_message(model))}</p>")
    for priority_index, priority in enumerate(priorities):
        rank = _rank(priority.get("rank"))
        identifier = _text(priority.get("recommendation_id"))
        open_attribute = " open" if priority_index == 0 else ""
        parts.append(
            f'<details class="climate-priority-disclosure"{open_attribute}>'
        )
        parts.append(
            '<summary><h3 class="climate-priority-title">'
            + html.escape(
                f"{rank}. {_text(priority.get('title'))} ({identifier})"
            )
            + "</h3></summary>"
        )
        parts.append('<div class="climate-priority-body">')
        narrative = _text(priority.get("narrative"))
        if narrative:
            for para in re.split(r"\n\s*\n+", narrative.strip()):
                para = para.strip()
                if para:
                    parts.append(f"<p>{html.escape(para)}</p>")
        for label, key in PRIORITY_FIELDS:
            value = _field_text(priority.get(key))
            if value:
                parts.append(
                    f"<p><strong>{html.escape(label)}:</strong> "
                    f"{html.escape(value)}</p>"
                )
            if key == "minimum_action":
                parts.append(_drafting_html(
                    "Current document drafting",
                    priority.get("current_document_drafting"),
                ))
                parts.append(_drafting_html(
                    "Operational instrument drafting",
                    priority.get("operational_instrument_drafting"),
                ))
        parts.append(_project_cycle_html(priority.get("project_cycle")))
        parts.append("</div></details>")

    # Secondary points remain visible, prose-led and locally numbered.
    minor_points = _records(model.get("minor_climate_points"))
    doc_flags = _records(model.get("review_readiness_flags"))
    if minor_points or doc_flags:
        parts.append(_heading(2, HEADINGS[3]))
        parts.append(f"<p>{html.escape(POINTS_TO_CHECK_INTRO)}</p>")
        if minor_points:
            parts.append("<h3>Smaller climate &amp; fragility points to consider</h3>")
            parts.append(
                "<p>These are smaller, climate- and fragility-specific points that were "
                "not large enough to become a recommendation above, but are still worth "
                "a look.</p>"
            )
            for index, point in enumerate(minor_points, start=1):
                parts.append('<section class="climate-numbered-item">')
                parts.append(
                    f'<span class="climate-item-number">{index:02d}</span>'
                    '<div class="climate-numbered-content">'
                )
                parts.append(_heading(4, _text(point.get("point"))))
                parts.append("<p>" + html.escape(_text(point.get("why"))) + "</p>")
                parts.append(
                    "<p><strong>How to address:</strong> "
                    + html.escape(_text(point.get("how_to_check")))
                    + "</p></div></section>"
                )

        if doc_flags:
            parts.append("<h3>Document points to confirm</h3>")
        for index, flag in enumerate(doc_flags, start=1):
            parts.append('<section class="climate-numbered-item">')
            parts.append(
                f'<span class="climate-item-number">{index:02d}</span>'
                '<div class="climate-numbered-content">'
            )
            parts.append(_heading(4, _text(flag.get("flag"))))
            parts.append(
                "<p><strong>Why it matters:</strong> "
                + html.escape(_text(flag.get("why_it_matters")))
                + "</p><p><strong>Suggested verification:</strong> "
                + html.escape(_text(flag.get("suggested_verification")))
                + "</p></div></section>"
            )

    watch_items = [
        (_text(q.get("question")), _text(q.get("watch")))
        for q in _records(model.get("core_questions"))
        if _text(q.get("watch"))
    ]
    if watch_items:
        parts.append(_heading(2, HEADINGS[4]))
        parts.append(
            "<p>These are things to monitor as the project develops. They are not "
            "actions to take now - just points to keep in view.</p>"
        )
        for index, (question_text, watch_text) in enumerate(watch_items, start=1):
            lead = (
                f"<strong>{html.escape(question_text)}</strong> "
                if question_text else ""
            )
            parts.append(
                '<section class="climate-numbered-item climate-watch-item">'
                f'<span class="climate-item-number">{index:02d}</span>'
                '<div class="climate-numbered-content">'
                f"<p>{lead}{html.escape(watch_text)}</p>"
                "</div></section>"
            )

    guidance_items = [
        item for item in _records(model.get("guidance_items"))
        if _is_public_world_bank_url(item.get("url"))
    ]
    if guidance_items:
        parts.append(_heading(2, "Relevant WBG guidance for this project"))
        parts.append(
            '<details class="climate-guidance-disclosure"><summary>'
            "Where the team can go for more detailed follow-up</summary>"
            '<div class="climate-guidance-body">'
            "<p>These sources are selected because they speak directly to the "
            "current findings. This is not a standard reading list.</p>"
        )
        for item in guidance_items:
            title = html.escape(_text(item.get("title")))
            url = html.escape(_text(item.get("url")))
            parts.append('<article class="climate-guidance-item">')
            parts.append(
                f'<h3><a href="{url}" target="_blank" rel="noopener">'
                f"{title}</a></h3>"
            )
            practical_value = _text(item.get("practical_value"))
            if practical_value:
                parts.append(
                    '<p class="climate-guidance-value">'
                    + html.escape(practical_value)
                    + "</p>"
                )
            project_use = _text(item.get("project_use"))
            if project_use:
                parts.append(
                    '<p class="climate-guidance-use">'
                    + html.escape(project_use)
                    + "</p>"
                )
            parts.append("</article>")
        parts.append("</div></details>")

    trail = _mapping(model.get("evidence_trail"))
    sources = model.get("sources") or []
    if trail or sources:
        parts.append(
            '<details class="climate-fold"><summary>'
            "Method, limitations, and sources</summary>"
        )
        methodology_note = _text(trail.get("methodology_note"))
        if methodology_note:
            parts.append(f"<p>{html.escape(methodology_note)}</p>")
        pathways = _records(trail.get("pathways"))
        if pathways:
            parts.append("<h4>Pathways</h4>")
            parts.append(
                "<p>Pathways are the specific ways climate pressures and fragility feed "
                "into each other in this project. Each one is a short chain from a cause "
                "to an effect.</p><ul>"
            )
            for pathway in pathways:
                parts.append(
                    f'<li><strong>{html.escape(_text(pathway.get("direction_label")))}:'
                    f'</strong> {html.escape(_text(pathway.get("chain_prose")))}</li>'
                )
            parts.append("</ul>")
        limitations = _field_text(trail.get("limitations"))
        if limitations:
            parts.append("<h4>Limitations</h4>")
            parts.append(f"<p>{html.escape(limitations)}</p>")
        if sources:
            parts.append("<h4>Sources &amp; further reading</h4>")
            parts.append(
                "<p>This analysis draws on the World Bank's core guidance on climate "
                "action in fragile and conflict-affected settings. Linked titles open the "
                "World Bank publication page if you want to read more.</p><ul>"
            )
            for source in sources:
                source_map = _mapping(source)
                title = html.escape(_text(source_map.get("title")))
                description = html.escape(_text(source_map.get("description")))
                url = source_map.get("url")
                if isinstance(url, str) and url.startswith("https://"):
                    label = (
                        f'<a href="{html.escape(url)}" target="_blank" '
                        f'rel="noopener">{title}</a>'
                    )
                    tail = f" - {description}" if description else ""
                else:
                    note = (
                        " (reference only; no public link is shown until one is confirmed)"
                    )
                    label = title
                    tail = f" - {description}{note}" if description else note
                parts.append(f"<li>{label}{tail}</li>")
            parts.append("</ul>")
        parts.append("</details>")

    parts.append(
        '<p class="climate-advisory">'
        + html.escape(_text(model.get("advisory_notice")))
        + "</p></article>"
    )
    return "".join(parts)


def _docx_field(document: Document, label: str, value: object) -> None:
    text = _field_text(value)
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(text)


def _docx_drafting(document: Document, label: str, value: object) -> None:
    block = _mapping(value)
    if not block:
        return
    document.add_heading(label, level=3)
    _docx_field(document, "Target document", block.get("target_document"))
    _docx_field(document, "Target section", block.get("target_section"))
    _docx_field(document, "Drafting status", block.get("drafting_status"))
    _docx_field(document, "Guidance basis", block.get("guidance_ids"))
    document.add_paragraph(_text(block.get("text")))


def _docx_project_cycle(document: Document, value: object) -> None:
    cycle = _mapping(value)
    primary_label = _text(cycle.get("primary_label"))
    primary_text = _text(cycle.get("primary_text"))
    if not primary_label or not primary_text:
        return
    document.add_heading("Where this fits in the project cycle", level=3)
    _docx_field(document, primary_label, primary_text)
    secondary_label = _text(cycle.get("secondary_label"))
    secondary_text = _text(cycle.get("secondary_text"))
    if secondary_label and secondary_text:
        _docx_field(document, secondary_label, secondary_text)


def write_reader_docx(model: dict[str, object], path: str | Path) -> Path:
    """Write the same reader dictionary to a compact Word document."""

    output = path if hasattr(path, "write") else Path(path)
    document = Document()
    if _text(model.get("runtime_mode")) == "smoke":
        paragraph = document.add_paragraph(SMOKE_RUNTIME_WARNING)
        if paragraph.runs:
            paragraph.runs[0].bold = True
    if _text(model.get("evidence_status")) == "preview; not approved":
        paragraph = document.add_paragraph(CANDIDATE_PREVIEW_WARNING)
        if paragraph.runs:
            paragraph.runs[0].bold = True
    operation_context = _mapping(model.get("operation_context"))
    if operation_context:
        document.add_heading("How this operation was routed", level=1)
        _docx_field(document, "Instrument", operation_context.get("instrument_type"))
        _docx_field(document, "Document", operation_context.get("document_type"))
        _docx_field(
            document,
            "Preparation",
            _text(operation_context.get("preparation_regime")).replace("_", " "),
        )
        _docx_field(
            document,
            "E&S route",
            _text(operation_context.get("es_regime")).replace("_", " "),
        )
        _docx_field(
            document,
            "Program layer",
            "MPA program" if operation_context.get("is_mpa") else "Not identified as MPA",
        )
        if (
            _text(operation_context.get("instrument_type")).casefold() == "unknown"
            or _text(operation_context.get("document_type")).casefold() == "unknown"
        ):
            document.add_paragraph(
                "Operational context could not be resolved safely, so "
                "document-targeted guidance was withheld."
            )
    # Overview at the very top: the summary + rating come first, then the fuller
    # Executive readout as detail below (parity with the HTML surface).
    rating = _mapping(model.get("climate_sensitivity_rating"))
    if rating:
        # Graphic first: the rating question, label and scale.
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{_text(rating.get('question'))} ").bold = True
        scale = rating.get("scale") if isinstance(rating.get("scale"), list) else []
        paragraph.add_run(
            f"Rating: {_text(rating.get('label'))}"
            + (f" (scale: {' - '.join(_text(s) for s in scale)})" if scale else "")
        )
        document.add_paragraph(_text(rating.get("description")))
        # Then the overall summary text below the graphic.
        summary = _text(rating.get("overview_summary"))
        if summary:
            document.add_paragraph(summary)
        caveat = document.add_paragraph(_text(rating.get("caveat")))
        if caveat.runs:
            caveat.runs[0].italic = True
    document.add_heading(HEADINGS[0], level=1)
    for _exec_para in re.split(r"\n\s*\n+", _text(model.get("executive_readout")).strip()):
        _exec_para = _exec_para.strip()
        if not _exec_para:
            continue
        _m = re.match(r"^(.*?[.!?])(\s+)([\s\S]*)$", _exec_para)
        _paragraph = document.add_paragraph()
        if _m:
            _paragraph.add_run(_m.group(1)).bold = True
            _paragraph.add_run(_m.group(2) + _m.group(3))
        else:
            _paragraph.add_run(_exec_para).bold = True

    document.add_heading(HEADINGS[1], level=1)
    document.add_paragraph(CORE_QUESTIONS_INTRO)
    for question in _records(model.get("core_questions")):
        document.add_heading(_text(question.get("question")), level=2)
        source = _text(question.get("source"))
        if source:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(
                f"For further insights on why this matters, see: {source}"
            )
            run.italic = True
        for para in re.split(r"\n\s*\n+", _text(question.get("summary")).strip()):
            para = para.strip()
            if para:
                document.add_paragraph(para)

    document.add_heading(HEADINGS[2], level=1)
    priorities = _records(model.get("priorities"))
    priority_summary = _mapping(model.get("priority_summary"))
    if priorities and _text(priority_summary.get("statement")):
        document.add_paragraph(_text(priority_summary.get("statement")))
    if not priorities:
        document.add_paragraph(_no_priority_message(model))
    for priority in priorities:
        rank = _rank(priority.get("rank"))
        identifier = _text(priority.get("recommendation_id"))
        document.add_heading(
            f"{rank}. {_text(priority.get('title'))} ({identifier})",
            level=2,
        )
        narrative = _text(priority.get("narrative"))
        if narrative:
            for para in re.split(r"\n\s*\n+", narrative.strip()):
                para = para.strip()
                if para:
                    document.add_paragraph(para)
        for label, key in PRIORITY_FIELDS:
            _docx_field(document, label, priority.get(key))
            if key == "minimum_action":
                _docx_drafting(
                    document, "Current document drafting",
                    priority.get("current_document_drafting"),
                )
                _docx_drafting(
                    document, "Operational instrument drafting",
                    priority.get("operational_instrument_drafting"),
                )
        _docx_project_cycle(document, priority.get("project_cycle"))

    minor_points = _records(model.get("minor_climate_points"))
    doc_flags = _records(model.get("review_readiness_flags"))
    if minor_points or doc_flags:
        document.add_heading(HEADINGS[3], level=1)
        document.add_paragraph(POINTS_TO_CHECK_INTRO)
        if minor_points:
            document.add_heading(
                "Smaller climate & fragility points to consider", level=2
            )
            document.add_paragraph(
                "These are smaller, climate- and fragility-specific points that were "
                "not large enough to become a recommendation above, but are still "
                "worth a look."
            )
            for index, point in enumerate(minor_points, start=1):
                document.add_heading(
                    f"{index:02d} {_text(point.get('point'))}", level=3
                )
                document.add_paragraph(_text(point.get("why")))
                _docx_field(document, "How to address", point.get("how_to_check"))

        if doc_flags:
            document.add_heading("Document points to confirm", level=2)
        for index, flag in enumerate(doc_flags, start=1):
            document.add_heading(
                f"{index:02d} {_text(flag.get('flag'))}", level=3
            )
            _docx_field(document, "Why it matters", flag.get("why_it_matters"))
            _docx_field(
                document,
                "Suggested verification",
                flag.get("suggested_verification"),
            )

    watch_items = [
        (_text(q.get("question")), _text(q.get("watch")))
        for q in _records(model.get("core_questions"))
        if _text(q.get("watch"))
    ]
    if watch_items:
        document.add_heading(HEADINGS[4], level=1)
        document.add_paragraph(
            "These are things to monitor as the project develops. They are not "
            "actions to take now - just points to keep in view."
        )
        for index, (question_text, watch_text) in enumerate(watch_items, start=1):
            prefix = f"{question_text}: " if question_text else ""
            document.add_paragraph(f"{index:02d} {prefix}{watch_text}")

    guidance_items = [
        item for item in _records(model.get("guidance_items"))
        if _is_public_world_bank_url(item.get("url"))
    ]
    if guidance_items:
        document.add_heading("Relevant WBG guidance for this project", level=1)
        document.add_heading("Where the team can go for more detailed follow-up", level=2)
        for item in guidance_items:
            document.add_heading(_text(item.get("title")), level=3)
            practical_value = _text(item.get("practical_value"))
            if practical_value:
                document.add_paragraph(practical_value)
            project_use = _text(item.get("project_use"))
            if project_use:
                document.add_paragraph(project_use)
            document.add_paragraph(_text(item.get("url")))
    trail = _mapping(model.get("evidence_trail"))
    sources = model.get("sources") or []
    if trail or sources:
        document.add_heading("Method, limitations, and sources", level=1)
        methodology_note = _text(trail.get("methodology_note"))
        if methodology_note:
            document.add_paragraph(methodology_note)
        pathways = _records(trail.get("pathways"))
        if pathways:
            document.add_heading("Pathways", level=2)
            document.add_paragraph(
                "Pathways are the specific ways climate pressures and fragility feed into "
                "each other in this project. Each one is a short chain from a cause to an "
                "effect."
            )
            for pathway in pathways:
                document.add_paragraph(
                    f'{_text(pathway.get("direction_label"))}: '
                    f'{_text(pathway.get("chain_prose"))}'
                )
        limitations = _field_text(trail.get("limitations"))
        if limitations:
            document.add_heading("Limitations", level=2)
            document.add_paragraph(limitations)
        if sources:
            document.add_heading("Sources & further reading", level=2)
            document.add_paragraph(
                "This analysis draws on the World Bank's core guidance on climate action "
                "in fragile and conflict-affected settings."
            )
            for source in sources:
                source_map = _mapping(source)
                title = _text(source_map.get("title"))
                description = _text(source_map.get("description"))
                url = source_map.get("url")
                if isinstance(url, str) and url.startswith("https://"):
                    tail = f" - {description}" if description else ""
                    line = f"{title}{tail} ({url})"
                else:
                    note = (
                        " (reference only; no public link is shown until one is confirmed)"
                    )
                    line = (
                        f"{title} - {description}{note}"
                        if description else f"{title}{note}"
                    )
                document.add_paragraph(line)


    document.add_paragraph(_text(model.get("advisory_notice")))
    document.save(output)
    return output

_summary_overview_paragraphs = validate_summary_overview


def _summary_known_evidence_ids(assessment: dict[str, object]) -> set[str]:
    """Collect only authoritative assessment-register IDs for fallback."""
    known: set[str] = set()

    def add_records(records: object, identifier_key: str) -> None:
        for record in _records(records):
            identifier = _text(record.get(identifier_key))
            if identifier:
                known.add(identifier)

    add_records(assessment.get("facts"), "claim_id")
    add_records(assessment.get("derived_assertions"), "assertion_id")
    analysis = _mapping(assessment.get("analysis"))
    register_names = (
        ("existing_responses", "response_id"),
        ("pathways", "pathway_id"),
        ("residual_gaps", "gap_id"),
    )
    for register_name, identifier_key in register_names:
        add_records(analysis.get(register_name), identifier_key)
        add_records(assessment.get(register_name), identifier_key)
    add_records(assessment.get("responses"), "response_id")
    add_records(assessment.get("gaps"), "gap_id")
    add_records(assessment.get("context_evidence"), "evidence_id")
    return known


def _summary_sentences(value: object) -> list[str]:
    text = _scrub_placeholder_text(_text(value))
    if not text:
        return []
    return [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+", re.sub(r"\s+", " ", text))
        if sentence.strip() and sentence.strip()[-1:] in ".?!"
    ]


def _legacy_summary_overview(
    overview_summary: object,
    judgments: list[dict[str, Any]],
    validation_reason_codes: object = (),
    canonical_text: object = (),
    known_evidence_ids: set[str] | None = None,
) -> list[str]:
    """Build a bounded legacy display from overview plus valid rating rationales."""
    sentences = _summary_sentences(overview_summary)
    valid_dimensions = {key for key, _, _ in DIMENSIONS}
    known_evidence_ids = set(known_evidence_ids or ())
    reason_codes = {
        _text(item) for item in validation_reason_codes
        if _text(item)
    } if isinstance(validation_reason_codes, (list, tuple, set)) else set()
    for judgment in judgments:
        dimension = _text(judgment.get("dimension"))
        value = _text(judgment.get("value"))
        rationale = _text(judgment.get("rationale"))
        evidence_ids = judgment.get("evidence_ids")
        dimension_issue = any(
            code.startswith(f"{dimension.upper()}_")
            or code.startswith("JUDGMENT_")
            for code in reason_codes
        )
        if (
            dimension in valid_dimensions
            and value in ALLOWED.get(dimension, set())
            and rationale
            and rationale[-1:] in ".?!"
            and validate_summary_fragment(rationale, canonical_text=canonical_text)
            and isinstance(evidence_ids, (list, tuple))
            and any(_text(item) in known_evidence_ids for item in evidence_ids)
            and not dimension_issue
        ):
            sentences.extend(_summary_sentences(rationale))
    bounded: list[str] = []
    word_count = 0
    for sentence in sentences:
        sentence_words = len(sentence.split())
        if not sentence_words or word_count + sentence_words > 220:
            break
        bounded.append(sentence)
        word_count += sentence_words
    if len(bounded) < 2:
        return bounded
    split_at = max(1, min(len(bounded) - 1, len(bounded) // 2))
    return [" ".join(bounded[:split_at]), " ".join(bounded[split_at:])]
