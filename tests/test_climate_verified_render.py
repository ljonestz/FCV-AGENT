from __future__ import annotations

from io import BytesIO

from docx import Document

from sector_lenses.climate_verified_render import (
    HEADINGS,
    build_reader_model,
    render_reader_html,
    validate_reader_model,
    write_reader_docx,
)


def _assessment() -> dict[str, object]:
    sentence = (
        "The project evidence supports a material Climate-FCV pathway, while "
        "the documented response remains at an early operational stage. "
    )
    return {
        "schema_version": "climate-verified-v2.1",
        "run_id": "run-1",
        "bank_release_id": "ssd-2026.08",
        "evidence_status": "preview; not approved",
        "executive_readout": sentence * 25,
        "judgments": {
            "relevance": {"value": "high", "rationale": "Material pathway."},
            "sensitivity": {
                "value": "moderate",
                "rationale": "Some relevant risks are recognized.",
            },
            "responsiveness": {
                "value": "emerging",
                "rationale": "Potential benefits are developing.",
            },
            "operationalization": {
                "value": "partial",
                "rationale": "Delivery arrangements remain incomplete.",
            },
        },
        "priorities": [
            {
                "recommendation_id": f"REC-00{index}",
                "rank": index,
                "title": f"Priority {index}",
                "decision": "Make a documented design decision.",
                "minimum_action": "Complete the proportionate minimum action.",
                "enhanced_action": None,
                "enhanced_activation": None,
                "responsible_function": "Task team",
                "routing_status": "standard_document_advisory",
                "authority_basis": "none_verified",
                "recommendation_basis": "project_evidence",
                "project_anchor_ids": ["PF-001"],
                "pathway_ids": ["PW-001"],
                "existing_response_ids": ["ER-001"],
                "residual_gap_ids": ["RG-001"],
                "instrument_claim_ids": [],
                "completion_evidence": "Updated project section",
                "completion_evidence_status": "updated_section",
                "confidence": "medium",
                "limitation": "Detailed parameters remain to be confirmed.",
                "caution": "Avoid unintended exclusion.",
                "current_document_drafting": {
                    "target_document": "PCN",
                    "target_section": "Project Description",
                    "drafting_status": "advisory_proposal",
                    "text": (
                        "Suggested targeted text for the current project document. "
                        * 12
                    ).strip(),
                    "project_basis_ids": ["PF-001"],
                    "gap_basis_ids": ["RG-001"],
                    "guidance_ids": ["GUIDE-PCN-DESIGN"],
                },
                "operational_instrument_drafting": ({
                    "target_document": "Security Risk Management Plan",
                    "target_section": "Continuity arrangements",
                    "drafting_status": "existing_commitment",
                    "text": (
                        "Distinct operational instrument text for continuity. " * 12
                    ).strip(),
                    "project_basis_ids": ["PF-001"],
                    "gap_basis_ids": ["RG-001"],
                    "guidance_ids": ["GUIDE-FCV-CONTINUITY"],
                } if index == 1 else None),
            }
            for index in range(1, 5)
        ],
        "review_readiness_flags": [
            {
                "flag_id": "RF-001",
                "category": "document_inconsistency",
                "flag": "Two sections state different financing totals.",
                "why_it_matters": "The controlling scope cannot be verified.",
                "suggested_verification": "Confirm the controlling total.",
            }
        ],
        "validation": {"status": "passed"},
        "recommendation_diagnostics": {
            "raw_candidate_count": 3,
            "parsed_candidate_count": 3,
            "valid_candidate_count": 3,
            "admitted_count": 3,
            "final_priority_count": 3,
            "reviewer_invoked": False,
            "reviewer_verdict": "not_invoked",
            "reason_codes": [],
            "unsupported_numeric_tokens": [],
            "semantic_review_object_ids": [],
            "candidate_suppressions": [],
        },
    }


def test_reader_has_four_dimensions_priority_cap_and_safe_annex():
    model = build_reader_model(_assessment())

    assert len(model["judgments"]) == 4
    assert len(model["priorities"]) == 3
    assert model["priority_summary"] == {
        "count": 3,
        "titles": ["Priority 1", "Priority 2", "Priority 3"],
        "statement": "Three final operational priorities are presented: Priority 1; Priority 2; Priority 3.",
    }
    assert "overall_rating" not in model
    assert model["evidence_status"] == "preview; not approved"
    assert model["technical_annex"] == {
        "run_id": "run-1",
        "schema_version": "climate-verified-v2.1",
        "bank_release_id": "ssd-2026.08",
        "validation_status": "passed",
        "recommendation_candidate_count": 3,
        "recommendation_admitted_count": 3,
        "recommendation_final_count": 3,
        "semantic_reviewer_invoked": False,
        "semantic_reviewer_verdict": "not_invoked",
        "recommendation_reason_codes": [],
        "unsupported_numeric_tokens": [],
        "semantic_review_object_ids": [],
        "candidate_suppressions": [],
        "live_research_count": 0,
    }


def test_judgment_evidence_ids_render_from_tuple_and_list():
    # The pipeline stores judgments via dataclasses.asdict(), which preserves
    # Judgment.evidence_ids as a tuple. The reader must surface those IDs, not
    # drop them because they are a tuple rather than a list.
    assessment = _assessment()
    assessment["judgments"]["relevance"]["evidence_ids"] = ("PF-001", "CE-001")
    assessment["judgments"]["sensitivity"]["evidence_ids"] = ["PF-002"]

    model = build_reader_model(assessment)
    by_dimension = {item["dimension"]: item for item in model["judgments"]}

    assert by_dimension["relevance"]["evidence_ids"] == ["PF-001", "CE-001"]
    assert by_dimension["sensitivity"]["evidence_ids"] == ["PF-002"]


def test_reader_exposes_bounded_existing_responses_with_safe_display_fields():
    assessment = _assessment()
    assessment["analysis"] = {
        "existing_responses": [
            {
                "response_id": f"ER-{index:03d}",
                "project_fact_ids": ["PF-001"],
                "pathway_ids": ["PW-001"],
                "description": f"Documented response {index}.",
                "limitation": f"Remaining limitation {index}.",
                "internal_review_note": "Must not reach the reader.",
            }
            for index in range(1, 14)
        ]
    }

    model = build_reader_model(assessment)

    assert model["existing_responses"] == [
        {
            "response_id": f"ER-{index:03d}",
            "project_fact_ids": ["PF-001"],
            "pathway_ids": ["PW-001"],
            "description": f"Documented response {index}.",
            "limitation": f"Remaining limitation {index}.",
        }
        for index in range(1, 13)
    ]


def test_priority_narrative_renders_in_reader_html_and_docx():
    assessment = _assessment()
    assessment["priorities"][0]["narrative"] = (
        "First paragraph tells the story of the gap and what to do.\n\n"
        "Second paragraph covers who leads it and what done looks like."
    )
    model = build_reader_model(assessment)
    assert model["priorities"][0]["narrative"].startswith("First paragraph")

    html = render_reader_html(model)
    assert "First paragraph tells the story of the gap" in html
    assert "Second paragraph covers who leads it" in html

    buffer = BytesIO()
    write_reader_docx(model, buffer)
    buffer.seek(0)
    text = "\n".join(p.text for p in Document(buffer).paragraphs)
    assert "First paragraph tells the story of the gap" in text
    assert "Second paragraph covers who leads it" in text


def test_priority_summary_count_must_match_final_priorities():
    model = build_reader_model(_assessment())
    model["priority_summary"]["count"] = 2

    assert "PRIORITY_SUMMARY_MISMATCH" in validate_reader_model(model)


def test_reader_annex_preserves_bounded_candidate_suppression_path():
    assessment = _assessment()
    detail = {
        "recommendation_id": "REC-001",
        "stage": "validation",
        "reason_codes": ["RECOMMENDATION_NUMBER_UNSUPPORTED"],
        "unsupported_numeric_fields": [
            {"field": "minimum_action", "tokens": ["30"]}
        ],
    }
    assessment["recommendation_diagnostics"]["candidate_suppressions"] = [
        detail
    ]
    assessment["recommendation_diagnostics"]["semantic_review_object_ids"] = [
        "REC-001"
    ]

    model = build_reader_model(assessment)

    assert model["technical_annex"]["candidate_suppressions"] == [detail]
    assert model["technical_annex"]["semantic_review_object_ids"] == [
        "REC-001"
    ]


def test_reader_validation_rejects_placeholder_and_duplicate_titles():
    model = build_reader_model(_assessment())
    model["priorities"][1]["title"] = model["priorities"][0]["title"]
    model["priorities"][0]["minimum_action"] = "[TBD]"

    issues = validate_reader_model(model)

    assert "DUPLICATE_PRIORITY_TITLE" in issues
    assert "UNRESOLVED_PLACEHOLDER" in issues


def test_reader_allows_readiness_flag_to_describe_project_placeholder():
    model = build_reader_model(_assessment())
    flag = model["review_readiness_flags"][0]
    flag["category"] = "material_placeholder"
    flag["flag"] = "The climate screening field remains a placeholder."

    assert validate_reader_model(model) == ()


def test_reader_uses_tolerant_integrity_bounds_for_executive_length():
    model = build_reader_model(_assessment())
    model["executive_readout"] = ("word " * 699) + "word."

    assert "EXECUTIVE_LENGTH_INVALID" not in validate_reader_model(model)

    model["executive_readout"] = ("word " * 249) + "word."
    assert "EXECUTIVE_LENGTH_INVALID" in validate_reader_model(model)

    model["executive_readout"] = ("word " * 949) + "word."
    assert "EXECUTIVE_LENGTH_INVALID" in validate_reader_model(model)


def test_html_and_docx_share_headings_and_priority_order():
    model = build_reader_model(_assessment())
    assert validate_reader_model(model) == ()

    html = render_reader_html(model)
    output = BytesIO()
    write_reader_docx(model, output)
    output.seek(0)
    document = Document(output)
    document_text = "\n".join(
        paragraph.text for paragraph in document.paragraphs
    )

    assert [html.index(heading) for heading in HEADINGS] == sorted(
        html.index(heading) for heading in HEADINGS
    )
    assert [document_text.index(heading) for heading in HEADINGS] == sorted(
        document_text.index(heading) for heading in HEADINGS
    )
    for index in range(1, 4):
        identifier = f"REC-00{index}"
        assert identifier in html
        assert identifier in document_text
        assert f"Priority {index}" in html
        assert f"Priority {index}" in document_text
    for expected in (
        "standard_document_advisory",
        "none_verified",
        "project_evidence",
        "PF-001",
        "PW-001",
        "RG-001",
        "Current document drafting",
        "Operational instrument drafting",
        "Suggested targeted text for the current project document.",
        "Distinct operational instrument text for continuity.",
        "GUIDE-PCN-DESIGN",
        "GUIDE-FCV-CONTINUITY",
    ):
        assert expected in html
        assert expected in document_text
    assert "REC-004" not in html
    assert "REC-004" not in document_text
    assert not any(
        paragraph.text.rstrip().endswith(("[", "{", "..."))
        for paragraph in document.paragraphs
    )
    assert html.count("Operational instrument drafting") == 1
    assert document_text.count("Operational instrument drafting") == 1
    assert html.index("Current document drafting") < html.index(
        "Operational instrument drafting"
    )

def test_zero_priority_message_is_shared_by_html_and_docx():
    assessment = _assessment()
    assessment["priorities"] = []
    model = build_reader_model(assessment)

    rendered = render_reader_html(model)
    stream = BytesIO()
    write_reader_docx(model, stream)
    stream.seek(0)
    document_text = "\n".join(
        paragraph.text for paragraph in Document(stream).paragraphs
    )

    message = "No recommendation passed the admission threshold for this run."
    assert message in rendered
    assert message in document_text


def test_semantic_review_suppression_is_explained_in_html_and_docx():
    assessment = _assessment()
    assessment["priorities"] = []
    assessment["recommendation_diagnostics"] = {
        "raw_candidate_count": 3,
        "admitted_count": 3,
        "final_priority_count": 0,
        "reviewer_invoked": True,
        "reviewer_verdict": "revise",
        "reason_codes": ["PROJECT_FACT_PROVENANCE_UNSUPPORTED"],
    }
    model = build_reader_model(assessment)

    rendered = render_reader_html(model)
    stream = BytesIO()
    write_reader_docx(model, stream)
    stream.seek(0)
    document_text = "\n".join(
        paragraph.text for paragraph in Document(stream).paragraphs
    )

    message = (
        "3 recommendation candidates passed deterministic admission but were "
        "withheld after semantic review. Review outcome: revise. See the "
        "technical annex."
    )
    assert message in rendered
    assert message in document_text
    assert "No recommendation passed the admission threshold" not in rendered
    assert "No recommendation passed the admission threshold" not in document_text


def test_html_escapes_model_authored_content():
    assessment = _assessment()
    assessment["priorities"][0]["title"] = "<script>alert('x')</script>"
    model = build_reader_model(assessment)

    rendered = render_reader_html(model)

    assert "<script>" not in rendered
    assert "&lt;script&gt;" in rendered


def test_docx_writer_accepts_an_in_memory_stream():
    model = build_reader_model(_assessment())
    stream = BytesIO()

    returned = write_reader_docx(model, stream)

    assert returned is stream
    stream.seek(0)
    assert Document(stream).paragraphs[0].text == HEADINGS[0]


def test_smoke_runtime_is_watermarked_in_html_and_docx():
    model = build_reader_model(_assessment())
    model["runtime_mode"] = "smoke"
    model["technical_annex"]["runtime_mode"] = "smoke"

    rendered = render_reader_html(model)
    stream = BytesIO()
    write_reader_docx(model, stream)
    stream.seek(0)
    document_text = "\n".join(
        paragraph.text for paragraph in Document(stream).paragraphs
    )

    warning = (
        "Smoke test: validates workflow completion only; "
        "not a quality benchmark."
    )
    assert warning in rendered
    assert warning in document_text


def test_quality_runtime_does_not_show_smoke_watermark():
    model = build_reader_model(_assessment())
    model["runtime_mode"] = "quality"

    assert "Smoke test:" not in render_reader_html(model)
