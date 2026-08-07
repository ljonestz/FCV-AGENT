from __future__ import annotations

from sector_lenses.climate_verified_render import build_evidence_trail


def _assessment():
    return {
        "bank_release_id": "ssd-2026.08",
        "evidence_status": "preview; not approved",
        "manifest": {"live_research_count": 2},
        "facts": [
            {"claim_id": "PF-1", "subject": "The project", "predicate": "targets",
             "object": "flood-prone fisheries", "supporting_excerpt": "excerpt"},
            {"claim_id": "PF-9", "subject": "unused", "predicate": "x", "object": "y",
             "supporting_excerpt": "not cited"},
        ],
        "analysis": {
            "existing_responses": [
                {"response_id": "ER-1", "description": "Flood-resilient standards mandated."}],
            "pathways": [
                {"pathway_id": "PW-1", "direction": "climate_to_fcv",
                 "chain": ["flood displacement", "tenure disruption", "resource conflict"],
                 "project_anchor_ids": ["PF-1"]}],
            "residual_gaps": [
                {"gap_id": "RG-1", "statement": "Adaptive triggers not specified."}],
        },
        "judgments": {
            "relevance": {"value": "high", "evidence_ids": ["PF-1", "PW-1"], "rationale": "r"},
            "sensitivity": {"value": "moderate", "evidence_ids": ["RG-1"], "rationale": "r"},
        },
        "priorities": [
            {"recommendation_id": "REC-1", "project_anchor_ids": ["PF-1"],
             "pathway_ids": ["PW-1"], "existing_response_ids": ["ER-1"],
             "residual_gap_ids": ["RG-1"], "instrument_claim_ids": []}],
        "recommendation_diagnostics": {
            "raw_candidate_count": 3, "admitted_count": 3, "final_priority_count": 3,
            "reviewer_verdict": "not_invoked"},
    }


def test_evidence_trail_projects_methodology_pathways_key_diagnostics():
    t = build_evidence_trail(_assessment())
    assert t["methodology_note"] and isinstance(t["methodology_note"], str)
    assert len(t["pathways"]) == 1
    p = t["pathways"][0]
    assert p["direction_label"] == "Climate -> FCV"
    assert "flood displacement" in p["chain_prose"] and "resource conflict" in p["chain_prose"]
    key = {e["id"]: e for e in t["evidence_key"]}
    assert set(key) == {"PF-1", "PW-1", "RG-1", "ER-1"}
    assert "flood-prone fisheries" in key["PF-1"]["text"]
    assert key["RG-1"]["text"] == "Adaptive triggers not specified."
    assert key["ER-1"]["text"] == "Flood-resilient standards mandated."
    assert "PF-9" not in key
    d = t["diagnostics"]
    assert d["candidate_count"] == 3 and d["final_count"] == 3
    assert d["reviewer_verdict"] == "not_invoked"
    assert d["live_research_count"] == 2
    assert d["bank_release"] == "ssd-2026.08"


def test_unresolvable_cited_id_gets_neutral_label():
    a = _assessment()
    a["judgments"]["relevance"]["evidence_ids"] = ["PF-1", "CE-LIVE-7"]
    t = build_evidence_trail(a)
    key = {e["id"]: e for e in t["evidence_key"]}
    assert "CE-LIVE-7" in key
    assert key["CE-LIVE-7"]["type_label"].lower().startswith("live") or \
           "not resolved" in key["CE-LIVE-7"]["text"].lower()


def test_chain_prose_keeps_all_elements_when_more_than_three():
    from sector_lenses.climate_verified_render import _chain_prose
    out = _chain_prose(["pressure", "mid1", "mid2", "consequence"])
    assert "pressure" in out and "mid1" in out and "mid2" in out
    assert out.rstrip(".").endswith("consequence")


def test_chain_prose_handles_short_and_nonstring_chains():
    from sector_lenses.climate_verified_render import _chain_prose
    short = _chain_prose(["a", "b"])
    assert "a" in short and "b" in short and short.endswith(".")
    # None / non-string elements are tolerated (dropped), no crash
    mixed = _chain_prose(["a", None, "c", "d"])
    assert mixed.endswith(".") and "a" in mixed


from io import BytesIO

from docx import Document

from sector_lenses.climate_verified_render import (
    attach_provenance,
    build_reader_model,
    render_reader_html,
    validate_reader_model,
    write_reader_docx,
)
from tests.test_climate_verified_render import _assessment as _reader_assessment


def test_attach_provenance_is_additive_and_non_gating():
    model = build_reader_model(_reader_assessment())
    before = validate_reader_model(model)
    raw = _assessment()  # from this test module: has analysis/facts/diagnostics
    attach_provenance(model, raw)
    assert "evidence_trail" in model and "sources" in model
    assert model["evidence_trail"]["pathways"]  # projected
    assert len(model["sources"]) >= 5
    assert all("title" in s for s in model["sources"])
    assert validate_reader_model(model) == before


def _reader_with_trail():
    model = build_reader_model(_reader_assessment())
    attach_provenance(model, _assessment())
    return model


def test_annex_renders_in_html_and_docx():
    model = _reader_with_trail()
    html = render_reader_html(model)
    assert "How this analysis was produced" in html
    assert "flood-prone fisheries" in html            # evidence key resolved text
    assert ("Sources &amp; further reading" in html) or ("Sources & further reading" in html)
    assert "FCV-Sensitive Climate Action Framework" in html  # a static framework
    stream = BytesIO()
    write_reader_docx(model, stream)
    stream.seek(0)
    text = "\n".join(p.text for p in Document(stream).paragraphs)
    assert "How this analysis was produced" in text
    assert "FCV-Sensitive Climate Action Framework" in text
    assert "flood-prone fisheries" in text   # resolved evidence-key text in DOCX


def test_provenance_normalises_em_en_dashes_in_evidence_trail():
    """attach_provenance runs after build_reader_model's scrub, so evidence-trail
    text (model-generated) must still have em/en dashes normalised to ASCII."""
    model = build_reader_model(_reader_assessment())
    raw = _assessment()
    # Inject em/en dashes into model-generated evidence text.
    raw["analysis"]["residual_gaps"][0]["statement"] = (
        "Adaptive triggers — e.g. flood thresholds – not specified."
    )
    raw["facts"][0]["object"] = "flood–prone fisheries"
    attach_provenance(model, raw)
    key = {e["id"]: e for e in model["evidence_trail"]["evidence_key"]}
    for entry in key.values():
        assert "—" not in entry["text"] and "–" not in entry["text"]
    # The RG statement is preserved as prose, just with ASCII hyphens.
    assert "Adaptive triggers" in key["RG-1"]["text"]
    assert "not specified" in key["RG-1"]["text"]


def test_lay_comprehensibility_annex_and_sources(tmp_path):
    """WS1: plain-language intros, source descriptions, name-only marking, unified label."""
    model = _reader_with_trail()
    html = render_reader_html(model)
    stream = BytesIO()
    write_reader_docx(model, stream)
    stream.seek(0)
    docx_text = "\n".join(p.text for p in Document(stream).paragraphs)

    for text in (html, docx_text):
        # Provenance intros explaining what pathways / the evidence key are.
        assert "short chain from a cause to an effect" in text
        assert "traced to its source" in text
        # Source descriptions render for a lay reader.
        assert "do no harm" in text                       # FCV-Sensitive description
        assert "pathway to peace" in text                 # Defueling Conflict description
        # Name-only sources (url=None) are explicitly marked, never fabricated.
        assert "no public link is shown until one is confirmed" in text
        # Parity: the TTL-friendly points-to-check label is used on every surface.
        assert "Points to check before the decision meeting" in text
        assert "Review readiness flags for task-team verification" not in text
