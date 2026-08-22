"""Phase 4 — regime-aware action_timing resolution in extract_priorities."""

import app as app_module

import io

from docx import Document

def _block(timing):
    return (
        '%%%JSON_START%%%{"fcv_rating":"Moderately addressed",'
        '"fcv_responsiveness_rating":"Emerging","sensitivity_summary":"x",'
        '"responsiveness_summary":"y","risk_exposure":{"risks_to":"a","risks_from":"b"},'
        '"priorities":[{"title":"T","fcv_dimension":"Contextual","tag":"[S]","the_gap":"g",'
        '"why_it_matters":"w","actions":[],"who_acts":"TTL","when":"soon","resources":"r",'
        '"pad_sections":"IV","action_timing":"' + timing + '"}]}%%%JSON_END%%%'
    )


def test_new_model_remaps_before_appraisal_to_td_review():
    r = app_module.extract_priorities(
        _block("required-before-appraisal"), uploaded_doc_names=[],
        preparation_regime="new_model", instrument="IPF",
    )
    t = r["priorities"][0]["action_timing"]
    assert t == "before-TD-review"
    assert "appraisal" not in t


def test_new_model_keeps_valid_new_model_timing():
    r = app_module.extract_priorities(
        _block("before-IR"), uploaded_doc_names=[],
        preparation_regime="new_model", instrument="IPF",
    )
    assert r["priorities"][0]["action_timing"] == "before-IR"


def test_legacy_keeps_before_appraisal():
    r = app_module.extract_priorities(
        _block("required-before-appraisal"), uploaded_doc_names=[],
        preparation_regime="legacy_transitional", instrument="IPF",
    )
    assert r["priorities"][0]["action_timing"] == "required-before-appraisal"


def test_default_call_unchanged_legacy_behaviour():
    # No regime kwargs (existing callers) -> legacy validation, unchanged.
    r = app_module.extract_priorities(_block("required-before-appraisal"), uploaded_doc_names=[])
    assert r["priorities"][0]["action_timing"] == "required-before-appraisal"
    # An invalid legacy value still nulls out, as before.
    r2 = app_module.extract_priorities(_block("whenever-you-like"), uploaded_doc_names=[])
    assert r2["priorities"][0]["action_timing"] is None


def test_new_model_timing_labels_wired_into_ui_and_docx():
    """Task 4.2: the 11 new-model timings render as pills (3 index.html maps) + DOCX."""
    from pathlib import Path
    root = Path(__file__).resolve().parent.parent
    html = (root / "index.html").read_text(encoding="utf-8")
    appsrc = (root / "app.py").read_text(encoding="utf-8")
    # Human label present in all three index.html timing maps (summary, pill, downloadHTML)
    assert html.count("Before Technical Design review") >= 3
    assert html.count("During implementation support") >= 3
    for key in ("shortly-after-OIS", "before-TD-review", "before-IR", "before-One-Review", "before-Board"):
        assert key in html, key
    # DOCX timing_map (app.py) carries the labels too
    assert "Before Technical Design review" in appsrc
    assert "'before-One-Review': 'Before One Review'" in appsrc


def _download_docx_for_priority(priority):
    response = app_module.app.test_client().post(
        "/api/download-report",
        json={
            "summary": "# Test project\nSummary.",
            "priorities": [priority],
            "metadata": {"date_str": "22 August 2026"},
        },
    )
    assert response.status_code == 200
    return Document(io.BytesIO(response.data))


def test_docx_renders_canonical_project_cycle_after_actions():
    document = _download_docx_for_priority({
        "title": "Strengthen delivery sequencing",
        "the_gap": "The current design does not sequence key decisions.",
        "why_it_matters": "Unclear sequencing could delay delivery.",
        "actions": [{
            "document_element": "Implementation arrangements",
            "guidance": "Set out the decision sequence.",
        }],
        "project_cycle": {
            "primary_label": "Before appraisal",
            "primary_text": "Confirm the decision sequence in the current document.",
            "secondary_label": "During implementation",
            "secondary_text": "Track whether the agreed sequence is being followed.",
        },
        "concise": {
            "project_cycle": {
                "primary_label": "Legacy concise label",
                "primary_text": "This must not be used by Detailed exports.",
            }
        },
    })
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert "Where this fits in the project cycle" in text
    assert "Before appraisal" in text
    assert "Confirm the decision sequence in the current document." in text
    assert "During implementation" in text
    assert "Track whether the agreed sequence is being followed." in text
    assert text.index("Where this fits in the project cycle") > text.index("Set out the decision sequence.")

    cycle_index = paragraphs.index("Where this fits in the project cycle")
    assert document.paragraphs[cycle_index + 1].runs[0].bold is True
    assert document.paragraphs[cycle_index + 2].runs[0].bold is not True
    assert document.paragraphs[cycle_index + 3].runs[0].bold is True
    assert document.paragraphs[cycle_index + 4].runs[0].bold is not True


def test_docx_omits_project_cycle_when_only_legacy_concise_data_exists():
    document = _download_docx_for_priority({
        "title": "Retain canonical timing only",
        "the_gap": "The detailed record has no lifecycle block.",
        "concise": {
            "project_cycle": {
                "primary_label": "Legacy concise label",
                "primary_text": "Legacy concise timing must not leak into exports.",
            }
        },
    })
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Where this fits in the project cycle" not in text
    assert "Legacy concise label" not in text
    assert "Legacy concise timing must not leak into exports." not in text
