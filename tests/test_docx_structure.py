from __future__ import annotations

import base64
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app import extract_docx_content, extract_docx_text
from docx_structure import DocxUnit, extract_docx_reader_parts, extract_docx_units


def _paragraph_xml(text: str, *, style: str | None = None):
    paragraph = OxmlElement("w:p")
    if style:
        properties = OxmlElement("w:pPr")
        style_element = OxmlElement("w:pStyle")
        style_element.set(qn("w:val"), style)
        properties.append(style_element)
        paragraph.append(properties)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    paragraph.append(run)
    return paragraph


def _append_sdt(parent, children, *, checked: bool | None = None):
    sdt = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    if checked is not None:
        checkbox = OxmlElement("w14:checkbox")
        state = OxmlElement("w14:checked")
        state.set(qn("w14:val"), "1" if checked else "0")
        checkbox.append(state)
        properties.append(checkbox)
    content = OxmlElement("w:sdtContent")
    for child in children:
        content.append(child)
    sdt.append(properties)
    sdt.append(content)
    parent.append(sdt)
    return sdt


def _fixture_document(*, conflicting: bool = False) -> Document:
    document = Document()
    document.add_paragraph("Before metadata")

    outer = document.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    cell.text = "Basic Information"
    nested = Document().add_table(rows=0, cols=2)
    for label, value in (
        ("Operation ID", "P511185"),
        ("Financing Instrument", "Investment Project Financing (IPF)"),
        (
            "Environmental and Social Risk Classification",
            "Substantial",
        ),
    ):
        row = nested.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    if conflicting:
        row = nested.add_row()
        row.cells[0].text = "Financing Instrument"
        row.cells[1].text = "Program-for-Results Financing (PforR)"

    # Move the nested table's OOXML below an SDT so the regular python-docx
    # table wrappers cannot accidentally flatten it into the outer cell.
    nested_table = nested._tbl
    nested_table.getparent().remove(nested_table)
    _append_sdt(cell._tc, [nested_table])

    option = OxmlElement("w:p")
    option_run = OxmlElement("w:r")
    option_text = OxmlElement("w:t")
    option_text.text = "[ ] Multiphase Programmatic Approach (MPA)"
    option_run.append(option_text)
    option.append(option_run)
    cell._tc.append(option)

    checked = OxmlElement("w:p")
    checked_run = OxmlElement("w:r")
    checked_text = OxmlElement("w:t")
    checked_text.text = "[x] Climate resilience component"
    checked_run.append(checked_text)
    checked.append(checked_run)
    cell._tc.append(checked)

    hidden = OxmlElement("w:p")
    hidden_run = OxmlElement("w:r")
    hidden_props = OxmlElement("w:rPr")
    hidden_props.append(OxmlElement("w:vanish"))
    hidden_run.append(hidden_props)
    hidden_text = OxmlElement("w:t")
    hidden_text.text = "Hidden instruction"
    hidden_run.append(hidden_text)
    hidden.append(hidden_run)
    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.text = " PAGE"
    instruction_run.append(instruction)
    hidden.append(instruction_run)
    deleted = OxmlElement("w:del")
    deleted_text = OxmlElement("w:delText")
    deleted_text.text = "Deleted text"
    deleted.append(deleted_text)
    hidden.append(deleted)
    cell._tc.append(hidden)
    return document


def _docx_bytes(*, conflicting: bool = False) -> bytes:
    stream = BytesIO()
    _fixture_document(conflicting=conflicting).save(stream)
    return stream.getvalue()


def _units(*, conflicting: bool = False) -> list[DocxUnit]:
    return extract_docx_units(_fixture_document(conflicting=conflicting))


def test_shared_walker_reads_sdt_nested_table_once_in_order():
    units = _units()
    texts = [unit.text for unit in units]

    assert texts[0] == "Before metadata"
    assert "Operation ID: P511185" in texts
    assert "Financing Instrument: Investment Project Financing (IPF)" in texts
    assert "Environmental and Social Risk Classification: Substantial" in texts
    assert texts.count("Operation ID: P511185") == 1
    assert texts.count(
        "Financing Instrument: Investment Project Financing (IPF)"
    ) == 1
    assert "Multiphase Programmatic Approach" not in "\n".join(texts)
    assert "[x] Climate resilience component" in texts
    assert "Hidden instruction" not in "\n".join(texts)
    assert "Deleted text" not in "\n".join(texts)
    assert "PAGE" not in "\n".join(texts)

    financing = next(
        unit for unit in units if unit.field_name == "Financing Instrument"
    )
    assert financing.field_value == "Investment Project Financing (IPF)"
    assert financing.table_coordinates is not None


def test_shared_walker_has_stable_locations_and_no_nested_duplicate_blocks():
    first = _units()
    second = _units()

    assert [unit for unit in first] == [unit for unit in second]
    assert [unit.table_coordinates for unit in first if unit.table_coordinates]
    assert all(
        unit.table_coordinates is None
        or len(unit.table_coordinates) >= 2
        for unit in first
    )


def test_app_docx_extraction_uses_shared_units_and_public_return_contract():
    encoded = base64.b64encode(_docx_bytes()).decode("ascii")
    extracted, part_count = extract_docx_text(encoded, "fixture.docx")

    assert "P511185" in extracted
    assert "Financing Instrument: Investment Project Financing (IPF)" in extracted
    assert "Environmental and Social Risk Classification: Substantial" in extracted
    assert "Multiphase Programmatic Approach" not in extracted
    assert part_count >= 5


def test_app_docx_content_returns_json_sidecar_and_public_api_stays_two_values():
    encoded = base64.b64encode(_docx_bytes()).decode("ascii")

    text, part_count, structured_fields = extract_docx_content(
        encoded, "fixture.docx"
    )
    public_text, public_count = extract_docx_text(encoded, "fixture.docx")

    assert (public_text, public_count) == (text, part_count)
    assert structured_fields
    financing = next(
        field for field in structured_fields
        if field["field_name"] == "Financing Instrument"
    )
    assert financing["field_value"] == "Investment Project Financing (IPF)"
    assert isinstance(financing["location"], str)
    assert isinstance(financing["paragraph_index"], int)
    assert isinstance(financing["table_coordinates"], list)

    import json

    json.dumps(structured_fields)


def test_checked_word_checkbox_content_control_is_retained():
    document = Document()
    paragraph = document.add_paragraph()
    _append_sdt(
        paragraph._p,
        [_paragraph_xml("Selected instrument")],
        checked=True,
    )

    units = extract_docx_units(document)

    assert any("Selected instrument" in unit.text for unit in units)



def test_unchecked_word_checkbox_content_control_is_suppressed():
    document = Document()
    paragraph = document.add_paragraph()
    _append_sdt(
        paragraph._p,
        [_paragraph_xml("Multiphase Programmatic Approach (MPA)")],
        checked=False,
    )

    units = extract_docx_units(document)

    assert "Multiphase Programmatic Approach" not in " ".join(
        unit.text for unit in units
    )


def test_checked_and_unchecked_bracket_options_keep_only_selected_text():
    document = Document()
    document.add_paragraph("[ ] Multiphase Programmatic Approach (MPA)")
    document.add_paragraph("[x] Investment Project Financing (IPF)")

    text = " ".join(unit.text for unit in extract_docx_units(document))

    assert "Multiphase Programmatic Approach" not in text
    assert "Investment Project Financing (IPF)" in text


def test_header_row_followed_by_value_row_preserves_each_structured_pair():
    document = Document()
    table = document.add_table(rows=2, cols=3)
    headers = (
        "Operation ID",
        "Financing Instrument",
        "Environmental and Social Risk Classification",
    )
    values = (
        "P511185",
        "Investment Project Financing (IPF)",
        "Substantial",
    )
    for index, text in enumerate(headers):
        table.cell(0, index).text = text
    for index, text in enumerate(values):
        table.cell(1, index).text = text

    units = extract_docx_units(document)
    fields = {
        unit.field_name: unit.field_value
        for unit in units
        if unit.field_name
    }

    assert fields == dict(zip(headers, values))


def test_all_empty_canonical_header_and_value_rows_emit_empty_structured_fields():
    document = Document()
    table = document.add_table(rows=2, cols=3)
    for index, text in enumerate(
        (
            "Operation ID",
            "Financing Instrument",
            "Environmental and Social Risk Classification",
        )
    ):
        table.cell(0, index).text = text

    fields = [
        (unit.field_name, unit.field_value)
        for unit in extract_docx_units(document)
        if unit.field_name
    ]

    assert fields == [
        ("Operation ID", ""),
        ("Financing Instrument", ""),
        ("Environmental and Social Risk Classification", ""),
    ]


def _clear_cell_content(cell) -> None:
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)


def _nested_table(text: str):
    nested_document = Document()
    nested = nested_document.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = text
    table_element = nested._tbl
    table_element.getparent().remove(table_element)
    return table_element


def test_ordinary_cell_units_follow_xml_order_around_nested_tables():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    first_cell = table.cell(0, 0)
    _clear_cell_content(first_cell)
    first_cell._tc.append(_paragraph_xml("cell 0 before"))
    first_cell._tc.append(_nested_table("nested cell"))
    first_cell._tc.append(_paragraph_xml("cell 0 after"))

    second_cell = table.cell(0, 1)
    _clear_cell_content(second_cell)
    second_cell._tc.append(_paragraph_xml("cell 1"))

    texts = [unit.text for unit in extract_docx_units(document)]

    assert texts == [
        "cell 0 before",
        "nested cell",
        "cell 0 after",
        "cell 1",
    ]


def test_nested_table_coordinates_include_parent_path_without_colliding():
    document = Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _clear_cell_content(cell)
    cell._tc.append(_paragraph_xml("outer cell"))
    cell._tc.append(_nested_table("nested cell"))

    units = extract_docx_units(document)
    outer = next(unit for unit in units if unit.text == "outer cell")
    nested = next(unit for unit in units if unit.text == "nested cell")

    assert outer.table_coordinates is not None
    assert nested.table_coordinates is not None
    assert nested.table_coordinates[:len(outer.table_coordinates)] == (
        outer.table_coordinates
    )
    assert len(nested.table_coordinates) > len(outer.table_coordinates)
    assert nested.table_coordinates != outer.table_coordinates
    assert nested.table_coordinates not in {
        unit.table_coordinates for unit in units
        if unit is not nested and unit.table_coordinates is not None
    }


def test_same_row_structured_pairs_preserve_extra_cells_and_nested_tables():
    document = Document()
    table = document.add_table(rows=1, cols=5)
    for index, text in enumerate(
        (
            "Operation ID",
            "P511185",
            "Financing Instrument",
            "Investment Project Financing (IPF)",
        )
    ):
        table.cell(0, index).text = text

    extra = table.cell(0, 4)
    _clear_cell_content(extra)
    extra._tc.append(_paragraph_xml("Extra visible metadata"))
    extra._tc.append(_nested_table("Nested visible metadata"))

    units = extract_docx_units(document)
    fields = [
        (unit.field_name, unit.field_value)
        for unit in units
        if unit.field_name
    ]
    texts = [unit.text for unit in units]

    assert fields == [
        ("Operation ID", "P511185"),
        ("Financing Instrument", "Investment Project Financing (IPF)"),
    ]
    assert texts.index("Operation ID: P511185") < texts.index(
        "Financing Instrument: Investment Project Financing (IPF)"
    )
    assert texts.count("Extra visible metadata") == 1
    assert texts.count("Nested visible metadata") == 1


def test_structured_row_paragraph_index_counts_sdt_content_before_following_text():
    document = Document()
    table = document.add_table(rows=2, cols=2)
    for cell, text in zip(
        table.rows[0].cells,
        ("Operation ID", "Financing Instrument"),
    ):
        _clear_cell_content(cell)
        _append_sdt(cell._tc, [_paragraph_xml(text)])
    for cell, text in zip(
        table.rows[1].cells,
        ("P511185", "Investment Project Financing (IPF)"),
    ):
        _clear_cell_content(cell)
        _append_sdt(cell._tc, [_paragraph_xml(text)])
    following = document.add_paragraph("after structured metadata")

    units = extract_docx_units(document)
    fields = [unit for unit in units if unit.field_name]
    after = next(unit for unit in units if unit.text == following.text)

    assert len(fields) == 2
    assert all(unit.paragraph_index is not None for unit in fields)
    assert after.paragraph_index is not None
    assert after.paragraph_index > max(
        unit.paragraph_index for unit in fields
        if unit.paragraph_index is not None
    )
    assert after.paragraph_index >= 4


def test_checked_sdt_wrappers_around_rows_and_cells_preserve_coordinates():
    document = Document()
    table = document.add_table(rows=0, cols=2)
    row = table.add_row()
    row_element = row._tr
    row_element.getparent().remove(row_element)

    first_cell = row.cells[0]
    second_cell = row.cells[1]
    first_cell._tc.getparent().remove(first_cell._tc)
    second_cell._tc.getparent().remove(second_cell._tc)
    wrapped_row = _append_sdt(table._tbl, [], checked=True)
    row_content = wrapped_row.find(qn("w:sdtContent"))
    row_content.append(row_element)
    wrapped_cell = _append_sdt(row_element, [], checked=True)
    cell_content = wrapped_cell.find(qn("w:sdtContent"))
    cell_content.append(first_cell._tc)
    row_element.append(second_cell._tc)
    first_cell._tc.append(_paragraph_xml("wrapped row"))
    second_cell._tc.append(_paragraph_xml("wrapped cell"))

    units = extract_docx_units(document)

    assert [(unit.text, unit.table_coordinates) for unit in units] == [
        ("wrapped row", (0, 0)),
        ("wrapped cell", (0, 1)),
    ]


def test_unchecked_sdt_row_wrapper_suppresses_all_row_content():
    document = Document()
    table = document.add_table(rows=0, cols=1)
    row = table.add_row()
    row_element = row._tr
    row_element.getparent().remove(row_element)
    wrapper = _append_sdt(table._tbl, [row_element], checked=False)

    assert extract_docx_units(document) == []
    assert wrapper is not None


def test_structured_pairing_rejects_canonical_labels_as_values():
    document = Document()
    table = document.add_table(rows=1, cols=3)
    for index, text in enumerate(("Operation ID", "Financing Instrument", "P511185")):
        table.cell(0, index).text = text

    units = extract_docx_units(document)

    assert not [unit for unit in units if unit.field_name]
    assert [unit.text for unit in units] == [
        "Operation ID",
        "Financing Instrument",
        "P511185",
    ]


def test_header_row_followed_by_label_row_is_not_consumed_as_metadata():
    document = Document()
    table = document.add_table(rows=2, cols=3)
    for index, text in enumerate(("Operation ID", "Financing Instrument", "P511185")):
        table.cell(0, index).text = text
    for index, text in enumerate(("Operation ID", "Financing Instrument", "Substantial")):
        table.cell(1, index).text = text

    units = extract_docx_units(document)

    assert not [unit for unit in units if unit.field_name]
    assert [unit.text for unit in units] == [
        "Operation ID",
        "Financing Instrument",
        "P511185",
        "Operation ID",
        "Financing Instrument",
        "Substantial",
    ]


def test_empty_labelled_financing_instrument_is_structured_without_dangling_colon():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Financing Instrument"
    table.cell(0, 1).text = ""

    units = extract_docx_units(document)
    financing = next(unit for unit in units if unit.field_name == "Financing Instrument")

    assert financing.field_value == ""
    assert financing.text == "Financing Instrument"


def test_nested_tables_inside_structured_cells_follow_explicit_pair_order():
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Operation ID"
    table.cell(0, 1).text = "Financing Instrument"
    table.cell(1, 0).text = "P511185"
    table.cell(1, 1).text = "Investment Project Financing (IPF)"
    for cell, text in (
        (table.cell(0, 0), "header nested"),
        (table.cell(1, 0), "value nested"),
        (table.cell(0, 1), "second header nested"),
        (table.cell(1, 1), "second value nested"),
    ):
        cell._tc.append(_nested_table(text))

    assert [unit.text for unit in extract_docx_units(document)] == [
        "Operation ID: P511185",
        "header nested",
        "value nested",
        "Financing Instrument: Investment Project Financing (IPF)",
        "second header nested",
        "second value nested",
    ]


def test_reader_parts_keep_generic_rows_joined_and_interleaved_with_paragraphs():
    from docx_structure import extract_docx_reader_parts

    document = Document()
    document.add_paragraph("before")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "cell1"
    table.cell(0, 1).text = "cell2"
    document.add_paragraph("after")
    stream = BytesIO()
    document.save(stream)

    encoded = base64.b64encode(stream.getvalue()).decode("ascii")
    extracted, part_count = extract_docx_text(encoded, "generic.docx")

    assert extracted == "before\n\ncell1 | cell2\n\nafter"
    assert part_count == 3
    assert extract_docx_reader_parts(document) == ["before", "cell1 | cell2", "after"]


def test_reader_parts_keep_multiline_cell_text_within_cell_and_join_cells():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    first = table.cell(0, 0)
    first.text = "line one"
    first.add_paragraph("line two")
    table.cell(0, 1).text = "cell two"

    assert extract_docx_reader_parts(document) == [
        "line one\nline two | cell two"
    ]


def test_bracket_options_are_removed_anywhere_without_dropping_unrelated_prose():
    document = Document()
    document.add_paragraph("Selection: [ ] MPA; [x] IPF")

    text = " ".join(unit.text for unit in extract_docx_units(document))

    assert "MPA" not in text
    assert "Selection:" in text
    assert "[x] IPF" in text


def test_unchecked_sdt_cell_is_suppressed_without_shifting_sibling_coordinate():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    first_cell = table.cell(0, 0)
    second_cell = table.cell(0, 1)
    first_element = first_cell._tc
    first_element.getparent().remove(first_element)
    row = table.rows[0]._tr
    wrapper = _append_sdt(row, [first_element], checked=False)
    row.remove(wrapper)
    row.insert(0, wrapper)
    second_cell.text = "visible sibling"

    units = extract_docx_units(document)

    assert [(unit.text, unit.table_coordinates) for unit in units] == [
        ("visible sibling", (0, 1)),
    ]
    assert wrapper is not None
