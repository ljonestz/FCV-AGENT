"""Application contract tests for selector, payload, prompts, and Stage-3 provenance."""

import json
import io
from pathlib import Path

import pytest

import app as app_module
import climate_question_bank
from sector_lenses import (
    CLIMATE_NATIVE_SCHEMA_VERSION,
    build_climate_stage2_prompt,
    load_registry,
)


FIXTURE_ROOT = Path(__file__).parent / "fixtures" / "sector_lenses"
SOUTH_SUDAN_FIXTURE = (
    Path(__file__).parent
    / "fixtures"
    / "climate"
    / "south_sudan_dual_use.json"
)
_CANONICAL_BASELINE = {
    "sensitivity_rating": "Adequate",
    "responsiveness_rating": "Emerging",
    "sensitivity_reasoning": "Conflict-sensitive delivery is explicit.",
    "responsiveness_reasoning": "A root-cause pathway is present.",
    "evidence_trail": [{
        "claim": "Seasonal access affects named project activities.",
        "source_ids": ["peace-social-dividends"],
        "project_anchor": "Project activity in the project area",
    }],
}


def _add_specific_climate_paths(payload):
    """Add the canonical envelope and compact paths to positive fixtures."""

    payload.setdefault("schema_version", CLIMATE_NATIVE_SCHEMA_VERSION)
    payload.setdefault("fcv_baseline", _CANONICAL_BASELINE)
    for lens in payload.get("lenses", []):
        if lens.get("lens_id") != "climate":
            continue
        lens.setdefault(
            "executive_summary",
            lens.get("materiality_summary")
            or "Climate and FCV pressures affect the project.",
        )
        lens.setdefault("operating_context", {
            "fcv_setting": "Access and institutions shape delivery.",
            "climate_setting": "Material hazards shape project conditions.",
            "intersection": (
                "Climate and FCV pressures interact around project access."
            ),
        })
        for interaction in lens.get("interaction_readout", []):
            direction = interaction.get("direction_id", "")
            if interaction.get("pathways") or direction not in {
                "climate-fcv-on-project", "project-on-climate-fcv",
            }:
                continue
            interaction["pathways"] = [{
                "pathway_id": f"{direction}-1",
                "pressure": "Material climate or project pressure",
                "mechanism": "Access and allocation conditions change.",
                "project_implication": (
                    "A named project activity faces a distributional risk."
                ),
                "design_response": "Apply a specific access safeguard.",
                "project_elements": ["Project activity"],
                "geographies": ["Project area"],
                "affected_groups": ["Affected users"],
                "systems_or_assets": [],
                "time_horizons": ["project-lifetime"],
                "research_claim_ids": [],
                "confidence": "medium",
                "evidence_gap": "Site-level evidence remains incomplete.",
            }]
    return payload


def test_production_catalogue_has_manual_climate_without_suggestions():
    client = app_module.app.test_client()

    response = client.get("/api/sector-lenses")

    assert response.status_code == 200
    payload = response.get_json()
    assert payload["warnings"] == []
    assert len(payload["lenses"]) == 1
    climate = payload["lenses"][0]
    assert climate["id"] == "climate"
    assert climate["activation"] == "manual"
    assert [section["id"] for section in climate["readout_sections"]] == [
        "invest-in", "deliver-through",
    ]
    assert app_module.detect_lens_suggestions(
        "Climate resilience adaptation drought mitigation transition",
        app_module.SECTOR_LENS_REGISTRY,
    ) == []


def test_metadata_response_includes_ranked_lens_suggestions(monkeypatch):
    monkeypatch.setattr(app_module, "SECTOR_LENS_REGISTRY", load_registry(FIXTURE_ROOT))
    monkeypatch.setattr(app_module, "detect_document_type_from_text", lambda text, client: "PAD")
    monkeypatch.setattr(app_module, "get_client", lambda: object())
    client = app_module.app.test_client()
    text = ("Irrigation and food security are material project activities. " * 4).strip()

    response = client.post("/api/detect-document-type", json={"doc_text": text})

    assert response.status_code == 200
    payload = response.get_json()
    assert payload["document_type"] == "PAD"
    assert payload["lens_suggestions"][0]["lens_id"] == "test-agriculture"
    assert payload["lens_suggestions"][0]["selected_by_default"] is True


def test_lens_detection_failure_does_not_erase_document_metadata(monkeypatch):
    monkeypatch.setattr(app_module, "detect_document_type_from_text", lambda text, client: "PAD")
    monkeypatch.setattr(app_module, "get_client", lambda: object())
    monkeypatch.setattr(
        app_module, "detect_lens_suggestions",
        lambda text, registry: (_ for _ in ()).throw(RuntimeError("detector unavailable")),
    )

    response = app_module.app.test_client().post(
        "/api/detect-document-type", json={"doc_text": "Project text " * 20}
    )

    assert response.status_code == 200
    assert response.get_json()["document_type"] == "PAD"
    assert response.get_json()["lens_suggestions"] == []


def test_payload_and_prompt_context_resolve_authoritative_versions():
    registry = load_registry(FIXTURE_ROOT)
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["test-agriculture", "unknown"],
        "lens_versions": {"test-agriculture": "0.1.0"},
    })

    context = app_module.build_lens_stage_context(state, stage=2, registry=registry)

    assert state.active_lenses == ["test-agriculture", "unknown"]
    assert context["active_lenses"] == [
        {"id": "test-agriculture", "version": "1.2.0", "position": "primary"}
    ]
    assert {item["code"] for item in context["warnings"]} == {
        "unknown_lens", "version_mismatch"
    }
    assert "%%%LENS_DIAGNOSTIC_START%%%" in context["prompt"]
    assert "Agriculture synthesis guidance" in context["prompt"]
    assert context["restart_required"] is True


def test_priority_parser_preserves_valid_lens_badges_and_drops_unknown_values():
    block = {
        "fcv_rating": "Moderate",
        "fcv_responsiveness_rating": "Emerging",
        "sensitivity_summary": "Summary",
        "responsiveness_summary": "Summary",
        "risk_exposure": {"risks_to": "A", "risks_from": "B"},
        "priorities": [{
            "title": "Target access safeguards",
            "the_gap": "Gap",
            "why_it_matters": "Why",
            "evidence": "Evidence",
            "actions": [],
            "lens_ids": ["climate", "energy", 7, "climate"],
            "lens_relevance": "Material overlap with targeting.",
        }],
    }

    parsed = app_module.extract_priorities(
        "%%%JSON_START%%%" + json.dumps(block) + "%%%JSON_END%%%"
    )

    assert parsed["error"] is False
    assert parsed["priorities"][0]["lens_ids"] == ["climate", "energy"]
    assert parsed["priorities"][0]["lens_relevance"] == "Material overlap with targeting."
    restricted = app_module.extract_priorities(
        "%%%JSON_START%%%" + json.dumps(block) + "%%%JSON_END%%%",
        active_lens_ids=["climate"],
    )
    assert restricted["priorities"][0]["lens_ids"] == ["climate"]


def test_no_active_lens_has_no_prompt_or_core_behaviour_change():
    context = app_module.build_lens_stage_context(
        app_module.AnalysisState.from_payload({}), stage=2
    )

    assert context["active_lenses"] == []
    assert context["prompt"] == ""
    assert context["warnings"] == []


def test_downloaded_report_has_sector_source_and_evidence_appendix(monkeypatch):
    from docx import Document

    monkeypatch.setattr(app_module, "SECTOR_LENS_REGISTRY", load_registry(FIXTURE_ROOT))
    response = app_module.app.test_client().post("/api/download-report", json={
        "summary": "# Test project\nSummary.",
        "priorities": [{
            "title": "Target access safeguards",
            "the_gap": "Gap",
            "lens_ids": ["test-agriculture", "invented-lens"],
            "lens_relevance": "Material targeting link.",
        }],
        "active_lenses": [{
            "id": "test-agriculture", "version": "1.2.0", "position": "primary"
        }],
        "lens_diagnostic": {"findings": [{
            "lens_ids": ["test-agriculture"], "status": "partially_addressed",
            "core_mappings": ["dnh:3"], "source_ids": ["agri-guidance"],
            "evidence": ["Water-user groups are used for targeting."],
            "mechanism": "water access grievance",
            "geography": "project area",
            "action_target": "beneficiary targeting",
        }]},
        "metadata": {"date_str": "21 July 2026"},
    })

    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Appendix: Sector-Lens Sources and Evidence" in text
    assert "Agriculture Test Guidance" in text
    assert "Water-user groups are used for targeting." in text
    assert "Sector lenses: test-agriculture" in text
    assert "invented-lens" not in text



@pytest.mark.parametrize(
    ("state", "expected_notice"),
    [
        ("bank+research", ""),
        ("bank-only", "Live web research was unavailable for this run."),
        ("research-only", "No reviewed country-bank release was available."),
        ("thematic-only", "No reviewed country-bank release or accepted live research was available."),
    ],
)
def test_docx_surfaces_climate_grounding_state_and_reviewed_bank_sources(
    monkeypatch, state, expected_notice,
):
    from docx import Document

    has_bank = state in {"bank+research", "bank-only"}
    grounding = {
        "state": state,
        "content_version": "ssd-pilot-2026-07",
        "country_iso3": "SSD",
        "research_status": "accepted" if "research" in state else "empty",
        "bank_manifest": {"bank_status": "ok" if has_bank else "unavailable"},
        "sources": [{
            "source_id": "SSD-SRC-001",
            "title": "South Sudan reviewed climate-FCV source",
            "organization": "Trusted institute",
            "publication_date": "2025",
            "url": "https://example.org/ssd-source",
            "provenance": ["bank"],
        }] if has_bank else [],
    }
    incoming_manifest = {
        "bank_status": "ok" if has_bank else "unavailable",
        "content_version": "ssd-pilot-2026-07" if has_bank else None,
        "country_iso3": "SSD" if has_bank else None,
    }
    calls = []

    def rematerialize(manifest, research, **kwargs):
        calls.append((manifest, research, kwargs))
        return grounding, research

    monkeypatch.setattr(
        app_module, "resolve_climate_grounding", rematerialize
    )
    response = app_module.app.test_client().post(
        "/api/download-report",
        json={
            "summary": "# Grounding state test\nSummary.",
            "active_lenses": [{
                "id": "climate", "version": "1.1.0", "position": "primary",
            }],
            "lens_diagnostic": {"error": True},
            "climate_grounding": {
                "state": state,
                "bank_manifest": incoming_manifest,
                "sources": [{
                    "source_id": "SSD-SRC-999",
                    "title": "FORGED CLIENT SOURCE",
                    "url": "https://malicious.example/source",
                    "provenance": ["bank"],
                }],
                "prompt_context": "SECRET EVIDENCE PACKET",
            },
            "climate_research": {"status": "failed"},
            "metadata": {"date_str": "31 July 2026"},
        },
    )

    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert calls
    assert calls[0][0] == incoming_manifest
    assert "FORGED CLIENT SOURCE" not in text
    assert "SECRET EVIDENCE PACKET" not in text
    if expected_notice:
        assert expected_notice in text
    else:
        assert "No reviewed country-bank release" not in text
    if has_bank:
        assert "Reviewed country evidence bank" in text
        assert "Content version: ssd-pilot-2026-07" in text
        assert "South Sudan reviewed climate-FCV source" in text
        assert "https://example.org/ssd-source" in text
    else:
        assert "Reviewed country evidence bank" not in text
def test_downloaded_report_has_climate_readout_and_context_sources():
    from docx import Document

    response = app_module.app.test_client().post("/api/download-report", json={
        "summary": "# Test project\nSummary.",
        "priorities": [{
            "title": "Inclusive seasonal access",
            "the_gap": "Seasonal users are not represented in access decisions.",
            "country_category_relevance": "Legacy differentiated note.",
            "climate_links": {
                "status": "linked",
                "interaction_pathway_ids": ["climate-fcv-on-project-1"],
                "dividend_pathway_ids": [
                    "institutional-capacity-legitimacy"
                ],
                "finding_ids": ["climate-finding-1"],
                "contribution": "Protects legitimate seasonal access.",
                "strengthening_effect": (
                    "Adds representation and transparent monitoring."
                ),
            },
        }],
        "fcv_rating": "Adequate",
        "fcv_responsiveness_rating": "Emerging",
        "sensitivity_summary": "The project recognizes key FCV risks and remaining gaps.",
        "responsiveness_summary": "The project could make a stronger contribution to resilience.",
        "risk_exposure": {
            "risks_to": "Core fallback risk to the project.",
            "risks_from": "Core fallback risk from the project.",
        },
        "active_lenses": [{
            "id": "climate", "version": "1.1.0", "position": "primary"
        }],
        "lens_diagnostic": {
            "schema_version": CLIMATE_NATIVE_SCHEMA_VERSION,
            "fcv_baseline": _CANONICAL_BASELINE,
            "lenses": [{
            "lens_id": "climate",
            "applicability": "material",
            "materiality_level": "high",
            "materiality_summary": "Drought and fragility affect delivery.",
            "executive_summary": (
                "Drought, access, and allocation shape project delivery."
            ),
            "operating_context": {
                "fcv_setting": "Insecurity constrains access in project areas.",
                "climate_setting": "Drought and floods affect site access.",
                "intersection": "Hazards and insecurity constrain delivery.",
            },
            "analysis_emphasis": ["adaptation"],
            "source_ids": ["peace-social-dividends", "context-ccdr"],
            "interaction_readout": [{
                "direction_id": "climate-fcv-on-project",
                "summary": "Drought, insecurity, and weak access could disrupt delivery.",
                "pathways": [{
                    "pathway_id": "climate-fcv-on-project-1",
                    "pressure": "Erratic floods",
                    "mechanism": "Access roads close during insecure periods.",
                    "project_implication": (
                        "Landing-site rehabilitation may be delayed."
                    ),
                    "design_response": "Use seasonal work windows.",
                    "project_elements": ["Landing-site rehabilitation"],
                    "geographies": ["Upper Nile"],
                    "affected_groups": ["Fishing households"],
                    "systems_or_assets": ["Feeder roads"],
                    "time_horizons": [
                        "project-lifetime", "asset-system-lifetime"
                    ],
                    "confidence": "medium",
                    "evidence_gap": "Site thresholds are not defined.",
                }],
                "source_ids": ["peace-social-dividends"],
            }, {
                "direction_id": "project-on-climate-fcv",
                "summary": "Benefit rules could strengthen resilience or exclusion.",
                "pathways": [{
                    "pathway_id": "project-on-climate-fcv-1",
                    "pressure": "New access rules",
                    "mechanism": "Rules redistribute seasonal access.",
                    "project_implication": (
                        "Seasonal users may lose adaptive options."
                    ),
                    "design_response": "Represent seasonal users.",
                    "project_elements": ["BFMU governance"],
                    "geographies": ["Sudd"],
                    "affected_groups": ["Seasonal users"],
                    "time_horizons": ["current-near-term"],
                    "confidence": "medium",
                    "evidence_gap": "",
                }],
                "source_ids": ["peace-social-dividends"],
            }],
            "readout_sections": [{
                "section_id": "invest-in",
                "items": [{
                    "item_id": "institutional-capacity-legitimacy",
                    "status": "supported",
                    "mechanism": "Transparent allocation can strengthen legitimacy.",
                    "project_contribution": "Community institutions allocate resources.",
                    "strengthening_action": "Map seasonal users before approving rules.",
                    "evidence": ["Community institutions are a project mechanism."],
                    "evidence_gap": "Seasonal users are not mapped.",
                    "trade_off": "Formalization may exclude customary users.",
                    "source_ids": ["peace-social-dividends"],
                }, {
                    "item_id": "social-cohesion-inclusion",
                    "status": "not_material",
                    "project_contribution": "Do not render this pathway.",
                    "strengthening_action": "Do not render this pathway.",
                    "evidence": ["No credible entry point."],
                    "source_ids": ["peace-social-dividends"],
                }],
            }, {
                "section_id": "deliver-through",
                "items": [{
                    "item_id": "flexible-adaptive-delivery",
                    "status": "potential",
                    "mechanism": "Contingent delivery can respond to shocks.",
                    "project_contribution": "The project uses contingent delivery.",
                    "strengthening_action": "Define combined flood and access triggers.",
                    "evidence": ["Contingent delivery is described."],
                    "evidence_gap": "No trigger is defined.",
                    "trade_off": "Flexibility needs accountability.",
                    "source_ids": ["defueling-conflict"],
                }],
            }],
            "additional_pathways": [{
                "section_id": "invest-in",
                "title": "Shared ecosystem restoration",
                "status": "potential",
                "mechanism": "Joint restoration can create collective benefits.",
                "project_contribution": "The project restores shared watersheds.",
                "strengthening_action": "Add joint oversight and dispute resolution.",
                "evidence": ["Watershed restoration is included."],
                "source_ids": ["peace-social-dividends"],
            }],
            "other_pathways": [{
                "pathway": "mitigation-transition",
                "status": "not_material",
                "reason": "No clear transition pathway.",
            }],
            "reflections": [{
                "question_key": "cq2_maladaptation",
                "title": "Could the design lock in maladaptation?",
                "status_cue": "partial gap",
                "source": "FCV-Sensitive Climate Action Framework",
                "text": "Answer paragraph one about lock-in.\n\nAnswer paragraph two names BFMU governance.",
            }],
            "strengths_weaknesses": [
                {"side": "strength", "title": "Community delivery",
                 "text": "Fits weak centre and adapts to floods."},
                {"side": "gap", "title": "Flood-displacement link",
                 "text": "Named but no design response."},
            ],
        }], "findings": []},
        "lens_context_sources": [{
            "id": "context-ccdr",
            "lens_id": "climate",
            "source_type": "ccdr",
            "country": "Exampleland",
            "title": "Exampleland Country Climate and Development Report",
            "publication_date": "2025",
            "url": "https://www.worldbank.org/example-ccdr",
            "location": "p. 4",
            "summary": "Drought affects project areas.",
        }],
        "metadata": {"date_str": "21 July 2026"},
    })

    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert text.index("How relevant is climate to this project?") < text.index("Summary.")
    assert "High climate relevance" in text
    assert "Drought, access, and allocation shape project delivery. Why it matters: Drought and fragility affect delivery." in text
    # S/R sections are replaced by the integration line in the climate path
    assert "FCV Sensitivity" not in text
    assert "FCV Responsiveness" not in text
    assert text.index("How well does the project integrate climate and FCV?") < text.index(
        "How climate and FCV dynamics could affect this project"
    )
    assert "How this project could affect climate and FCV dynamics" in text
    assert "Key locations and components:" in text
    assert "Landing-site rehabilitation" in text
    assert "over the life of the assets" in text
    # Redesign: core-questions section (lay intro + theme answers with source) + S&W;
    # the standalone dividend-synthesis and wider-FCV sections are dropped in module mode.
    assert "Core climate and FCV questions" in text
    assert "Maximizing the Peace and Social Dividends of Climate Action" in text
    assert "Could the design lock in maladaptation?" in text
    assert "For further insights on why this matters, see: FCV-Sensitive Climate Action Framework" in text
    assert "[partial gap]" not in text
    assert "How the design holds up on climate and FCV" in text
    assert "Where the design is stronger" in text
    assert "Community delivery" in text
    assert "Named but no design response." in text
    assert "Wider FCV context" not in text
    # The priority still appears in the main Priority Actions table; the standalone
    # dividend-synthesis panel (with its "Priority N (title)" links) is dropped.
    assert "Inclusive seasonal access" in text
    assert "Differentiated approach note" not in text
    assert "Legacy differentiated note." not in text
    assert "Do not render this pathway" not in text
    assert "Other pathways considered" not in text
    assert "Core fallback risk to the project" not in text
    assert "Country Climate and Development Report" in text
    assert text.count("Country Climate and Development Report") <= 2


def test_downloaded_core_only_report_retains_differentiated_approach_note():
    from docx import Document

    response = app_module.app.test_client().post("/api/download-report", json={
        "summary": "# Core-only test\nSummary.",
        "priorities": [{
            "title": "Tailor delivery arrangements",
            "the_gap": "Delivery arrangements are not differentiated.",
            "country_category_relevance": (
                "Use a differentiated approach in high-risk areas."
            ),
        }],
        "metadata": {"date_str": "23 July 2026"},
    })

    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Differentiated approach note" in text
    assert "Use a differentiated approach in high-risk areas." in text
    assert "Climate, peace and social dividend contribution" not in text


def test_downloaded_report_scales_low_climate_materiality_without_empty_dividends():
    from docx import Document

    response = app_module.app.test_client().post("/api/download-report", json={
        "summary": "# Low test\nSummary.",
        "priorities": [],
        "sensitivity_summary": "Core sensitivity remains material.",
        "responsiveness_summary": "Core responsiveness remains limited.",
        "active_lenses": [{
            "id": "climate", "version": "1.1.0", "position": "primary"
        }],
        "lens_diagnostic": {
            "schema_version": CLIMATE_NATIVE_SCHEMA_VERSION,
            "fcv_baseline": _CANONICAL_BASELINE,
            "lenses": [{
            "lens_id": "climate",
            "applicability": "possible",
            "materiality_level": "low",
            "materiality_summary": "Climate entry points are limited.",
            "executive_summary": (
                "Seasonal rainfall creates a limited project interaction."
            ),
            "operating_context": {
                "fcv_setting": "Core FCV conditions remain material.",
                "climate_setting": "Seasonal rainfall may affect access.",
                "intersection": "The overlap is limited and localized.",
            },
            "interaction_readout": [{
                "direction_id": "climate-fcv-on-project",
                "summary": "Seasonal rainfall may modestly affect access.",
            }],
            "readout_sections": [],
            "additional_pathways": [],
        }], "findings": []},
        "metadata": {"date_str": "22 July 2026"},
    })

    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Low climate relevance" in text
    assert "Seasonal rainfall may modestly affect access" in text
    assert "Climate, peace and social dividends" not in text


def test_downloaded_report_uses_safe_climate_failure_and_core_risk_fallback():
    from docx import Document

    response = app_module.app.test_client().post("/api/download-report", json={
        "summary": "# Failure test\nSummary.",
        "priorities": [],
        "active_lenses": [{
            "id": "climate", "version": "1.1.0", "position": "primary"
        }],
        "lens_diagnostic": {},
        "risk_exposure": {
            "risks_to": "Insecurity could disrupt delivery.",
            "risks_from": "Exclusion could deepen grievance.",
        },
        "metadata": {"date_str": "22 July 2026"},
    })

    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "validated Climate-FCV diagnostic could not be produced" in text
    assert "FCV Risk Exposure" in text
    assert "Insecurity could disrupt delivery" in text
    assert "Climate, peace and social dividends" not in text


def test_frontend_contract_includes_selector_locking_v3_and_lens_rendering():
    html = (Path(app_module.__file__).parent / "index.html").read_text(encoding="utf-8")

    assert 'id="lens-selector"' in html
    assert "version: 3" in html
    assert "active_lenses:activeLenses" in html
    assert "lensSelectionLocked=true" in html
    assert "renderLensDiagnostic()" in html
    assert "pr.lens_ids" in html
    assert "function resumeExpressRun" not in html
    assert "restart Stage 1" in html


def test_final_stage3_lens_prompt_respects_combined_platform_budget():
    registry = load_registry(FIXTURE_ROOT)
    state = app_module.AnalysisState.from_payload({"active_lenses": ["test-agriculture"]})
    huge = {"findings": [{
        "lens_ids": ["test-agriculture"], "evidence": ["x" * 1000],
        "source_ids": ["agri-guidance"], "core_mappings": ["dnh:3"],
        "mechanism": f"mechanism-{index}", "geography": "north",
        "action_target": "targeting", "status": "gap",
    } for index in range(50)]}

    context = app_module.build_lens_stage_context(state, 3, registry, huge)

    assert context["estimated_tokens"] <= 1600
    assert context["truncated"] is True


def test_climate_diagnostic_retains_readouts_and_ccdr_context():
    module_root = Path(app_module.__file__).parent / "sector_lenses" / "modules"
    registry = load_registry(module_root)
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"]
    })
    payload = {"lenses": [{
        "lens_id": "climate",
        "applicability": "material",
        "materiality_level": "medium",
        "materiality_summary": "Drought affects delivery and livelihoods.",
        "analysis_emphasis": ["adaptation", "resource access"],
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "Drought and insecurity could disrupt delivery.",
        }, {
            "direction_id": "project-on-climate-fcv",
            "summary": "Resource allocation could affect resilience and trust.",
        }],
        "source_ids": [
            "peace-social-dividends", "context-ccdr", "invented"
        ],
        "readout_sections": [{
            "section_id": "invest-in",
            "items": [{
                "item_id": "institutional-capacity-legitimacy",
                "status": "supported",
                "mechanism": "Transparent allocation can strengthen legitimacy.",
                "evidence": ["Water committees allocate access."],
                "evidence_gap": "Seasonal users are not mapped.",
                "trade_off": "Formalization may exclude customary users.",
                "source_ids": ["peace-social-dividends", "context-ccdr"],
            }],
        }],
        "other_pathways": [{
            "pathway": "mitigation-transition",
            "status": "not_material",
            "reason": "No emissions or transition mechanism is documented.",
        }],
    }], "findings": []}
    sources = [{
        "id": "context-ccdr",
        "lens_id": "climate",
        "source_type": "ccdr",
        "country": "Exampleland",
        "title": "Example CCDR",
        "publication_date": "2025",
        "url": "https://www.worldbank.org/example",
        "location": "p. 4",
        "summary": "Drought affects project areas.",
    }, {
        "id": "context-ccdr",
        "lens_id": "climate",
        "source_type": "ccdr",
        "title": "Untrusted context",
        "url": "https://example.com/not-a-ccdr",
        "summary": "This source must not enter the diagnostic.",
    }]

    payload = _add_specific_climate_paths(payload)
    context = app_module.build_lens_stage_context(
        state, 3, registry, payload, sources
    )

    assert "Drought affects delivery" in context["prompt"]
    assert "context-ccdr" in context["prompt"]
    assert "invented" not in context["prompt"]
    assert context["lens_context_sources"] == sources[:1]


def test_active_climate_stage2_supersedes_lightweight_core_check():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"]
    })

    context = app_module.build_lens_stage_context(state, 2)

    assert "materiality_summary" in context["prompt"]
    assert "readout_sections" in context["prompt"]
    assert "supersedes the lightweight supplementary Climate-FCV Nexus" in (
        context["prompt"]
    )
    assert "do not produce a duplicate" in context["prompt"]


def test_native_climate_metadata_context_skips_unused_heavy_prompt():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"]
    })
    signals = " ".join(
        trigger
        for question in climate_question_bank.CLIMATE_QUESTION_BANK
        for trigger in question["triggers"]
    )
    climate_research = {
        "status": "complete",
        "attempts": 1,
        "sources": [
            {
                "id": "climate-source-1",
                "source_type": "ccdr",
                "title": "Country CCDR",
                "url": "https://www.worldbank.org/example",
            },
            {
                "id": "climate-source-2",
                "source_type": "scientific",
                "title": "Climate assessment",
                "url": "https://www.ipcc.ch/example",
            },
        ],
        "claims": [{
            "id": f"climate-claim-{index}",
            "claim": (
                "Climate pressure affects named project delivery systems and "
                "vulnerable groups through access and institutional constraints."
            ),
            "source_ids": ["climate-source-1", "climate-source-2"],
            "geographies": ["Project area"],
            "project_elements": ["Project component"],
            "affected_groups": ["Affected group"],
            "systems_or_assets": ["Delivery system"],
            "evidence_status": "projected",
            "confidence": "medium",
            "time_horizons": ["project-lifetime"],
            "evidence_gap": "Site evidence remains incomplete.",
        } for index in range(1, 4)],
        "failure_reason": "",
    }

    context = app_module.build_lens_stage_context(
        state,
        2,
        climate_research=climate_research,
        project_signals=signals,
        compose_prompt=False,
    )

    assert context["active_lenses"][0]["id"] == "climate"
    assert context["prompt"] == ""
    assert context["estimated_tokens"] == 0
    assert context["lens_context_sources"] == []


def test_active_climate_stage2_requests_materiality_interactions_and_pathways():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"]
    })

    prompt = app_module.build_lens_stage_context(state, 2)["prompt"]

    for field in (
        "materiality_level", "interaction_readout", "project_contribution",
        "strengthening_action", "additional_pathways",
    ):
        assert field in prompt
    assert "development project" in prompt
    assert "not its primary objective" in prompt


def test_climate_stage2_requires_project_specific_causal_contract():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"]
    })
    climate_research = {
        "status": "complete",
        "sources": [{
            "id": "climate-source-1",
            "source_type": "world-bank",
            "title": "Climate profile",
            "url": "https://www.worldbank.org/example",
        }],
        "claims": [{
            "id": "climate-claim-1",
            "source_ids": ["climate-source-1"],
            "claim": "Flood timing affects site access.",
            "project_elements": ["Landing sites"],
            "geographies": ["Upper Nile"],
            "affected_groups": [],
            "systems_or_assets": ["Access roads"],
            "evidence_status": "observed",
            "confidence": "medium",
            "time_horizons": ["project-lifetime"],
        }],
    }

    prompt = app_module.build_lens_stage_context(
        state,
        2,
        climate_research=climate_research,
    )["prompt"]

    for value in (
        "pressure -> mediated mechanism -> project implication -> design response",
        "current-near-term",
        "project-lifetime",
        "asset-system-lifetime",
        "research_claim_ids",
        "Suppress generic pathways",
        "climate-claim-1",
    ):
        assert value in prompt


def test_diagnostic_rejects_directions_without_specific_pathways():
    diagnostic = {
        "lenses": [{
            "lens_id": "climate",
            "materiality_level": "medium",
            "materiality_summary": "Material interactions.",
            "interaction_readout": [
                {"direction_id": "climate-fcv-on-project", "summary": "A"},
                {"direction_id": "project-on-climate-fcv", "summary": "B"},
            ],
        }],
        "findings": [],
    }

    failure = app_module.lens_diagnostic_failure_message(
        diagnostic, ["climate"]
    )

    # Both directions present but neither carries a specific causal pathway,
    # so the diagnostic is still rejected (graceful degradation requires at
    # least one fully-specified direction). Message consolidated in the
    # single-direction-acceptance fix.
    assert failure


def test_active_climate_stage3_preserves_option_a_layers_and_gradient():
    state = app_module.AnalysisState.from_payload({"active_lenses": ["climate"]})
    diagnostic = {"lenses": [{
        "lens_id": "climate",
        "applicability": "material",
        "materiality_level": "high",
        "materiality_summary": "Flood and FCV pressures are central.",
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "Flood and insecurity could disrupt delivery.",
        }, {
            "direction_id": "project-on-climate-fcv",
            "summary": "Benefit rules could affect resilience and trust.",
        }],
        "readout_sections": [],
        "additional_pathways": [],
    }], "findings": []}

    diagnostic = _add_specific_climate_paths(diagnostic)
    context = app_module.build_lens_stage_context(
        state, 3, lens_diagnostic=diagnostic
    )
    prompt = context["prompt"]

    assert "High, Medium, or Low" in prompt
    assert "executive summary" in prompt
    assert "two-way Climate-FCV interaction" in prompt
    assert "current contribution" in prompt
    assert "how it could be strengthened" in prompt
    assert "not a quota" in prompt
    assert "Flood and insecurity could disrupt delivery" in prompt
    assert context["lens_diagnostic"]["lenses"][0]["materiality_level"] == "high"


def test_high_climate_materiality_warns_when_priorities_drop_provenance(caplog):
    diagnostic = {"lenses": [{
        "lens_id": "climate",
        "materiality_level": "high",
    }]}

    with caplog.at_level("WARNING", logger=app_module.app.logger.name):
        warning = app_module.warn_on_missing_high_climate_priority(
            [{"title": "Core priority", "lens_ids": []}], diagnostic
        )

    assert warning is True
    assert "High Climate-FCV materiality" in caplog.text
    caplog.clear()
    assert app_module.warn_on_missing_high_climate_priority(
        [{"title": "Climate priority", "lens_ids": ["climate"]}], diagnostic
    ) is False


def test_lens_diagnostic_failure_names_parser_errors_and_missing_entries():
    assert app_module.lens_diagnostic_failure_message(
        {"error": True, "message": "Lens diagnostic block was not valid JSON."},
        ["climate"],
    ) == "Lens diagnostic block was not valid JSON."
    assert app_module.lens_diagnostic_failure_message(
        {"lenses": [], "findings": []}, ["climate"]
    ) == "The Climate-FCV diagnostic was omitted from the Stage 2 structured output."
    assert app_module.lens_diagnostic_failure_message(
        _add_specific_climate_paths({"lenses": [{
            "lens_id": "climate",
            "materiality_level": "medium",
            "materiality_summary": "Material interactions.",
            "interaction_readout": [
                {"direction_id": "climate-fcv-on-project", "summary": "A"},
                {"direction_id": "project-on-climate-fcv", "summary": "B"},
            ],
        }], "findings": []}), ["climate"]
    ) == ""
    assert "incomplete" in app_module.lens_diagnostic_failure_message(
        {"lenses": [{
            "lens_id": "climate",
            "materiality_level": "high",
            "materiality_summary": "Material interactions.",
            "interaction_readout": [],
        }], "findings": []}, ["climate"]
    ).lower()


def test_stage2_climate_prompt_requires_reflections_and_intersection():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    ctx = app_module.build_lens_stage_context(
        state, 2, climate_research={"status": "failed", "attempts": 0,
                                    "sources": [], "claims": [], "failure_reason": ""},
    )
    prompt = ctx["prompt"]
    assert "reflections" in prompt
    assert "integration_level" in prompt
    assert "cq2_maladaptation" in prompt
    assert "climate and an FCV" in prompt  # intersection rule wording
    # Reflection text must be prose, not a mechanical checklist entry, and the
    # status cue must be plain words rather than a snake_case token.
    assert "mechanical checklist entry" in prompt
    assert "never a snake_case token" in prompt.replace("\n", " ")


def test_dedicated_climate_stage2_prompt_is_not_under_hood_sibling():
    prompt = build_climate_stage2_prompt(
        instrument_type="IPF",
        document_type="PAD",
        temporal_guardrail="Treat as preparation-stage evidence.",
        regime_header="Preparation regime: current policy.",
        project_signals="flood displacement community infrastructure",
        climate_research={
            "status": "complete",
            "sources": [],
            "claims": [],
        },
        priority_questions=[],
    )

    assert prompt.count("%%%LENS_DIAGNOSTIC_START%%%") == 1
    assert prompt.count("%%%LENS_DIAGNOSTIC_END%%%") == 1
    assert CLIMATE_NATIVE_SCHEMA_VERSION in prompt
    assert "single source of truth" in prompt.lower()
    for generic_marker in (
        "%%%UNDER_HOOD_START%%%",
        "%%%UNDER_HOOD_END%%%",
        "%%%RECS_TABLE_START%%%",
        "%%%DNH_CHECKLIST_START%%%",
        "%%%QUESTIONS_MAP_START%%%",
    ):
        assert generic_marker not in prompt


def test_stage2_climate_prompt_injects_bank_and_requests_source_and_rating():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    ctx = app_module.build_lens_stage_context(
        state, 2,
        climate_research={"status": "failed", "attempts": 0, "sources": [], "claims": [], "failure_reason": ""},
        project_signals="IPF fisheries flooding displacement cold storage community co-management",
    )
    prompt = ctx["prompt"]
    # Bank questions surface as guidance
    assert "core climate-fcv questions" in prompt.lower()
    assert "FCV-Sensitive Climate Action Framework" in prompt  # a bank source
    # New field requests
    assert "integration_rating" in prompt
    assert "Extremely Low" in prompt and "Very Well Embedded" in prompt  # 6-tier scale
    assert "source" in prompt  # per-reflection source
    # Two-paragraph depth instruction
    assert "two" in prompt.lower() and "paragraph" in prompt.lower()


def test_climate_stage2_is_native_not_generic_engine():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    prompt = app_module.build_lens_stage_context(
        state, 2, climate_research={"status": "failed", "attempts": 0, "sources": [], "claims": [], "failure_reason": ""},
    )["prompt"]
    assert "core climate-fcv questions" in prompt.lower()
    # Sanity: a non-climate PAD Stage 2 still uses the generic engine unchanged.
    plain = app_module.AnalysisState.from_payload({"active_lenses": [], "lens_versions": {}, "doc_type": "PAD"})
    plain_prompt = app_module.build_lens_stage_context(plain, 2)["prompt"]
    assert "core climate-fcv questions" not in plain_prompt.lower()


def test_stage2_climate_prompt_has_opcs_calibration_guardrails():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    prompt = app_module.build_lens_stage_context(
        state, 2,
        climate_research={"status": "failed", "attempts": 0, "sources": [], "claims": [], "failure_reason": ""},
    )["prompt"]
    low = prompt.lower()
    assert "instrument-route" in low                     # 12.1
    assert "never determine" in low                      # advisory boundary / 12.2
    assert "asset-appropriate design horizon" in low     # 12.3 dropped universal 20-50yr
    assert "will cause conflict" in low                  # 12.6 names the banned deterministic phrasing
    assert "not an opcs policy" in low                   # 12.7 source labelling
    # Non-climate PAD Stage 2 does NOT carry the climate calibration block
    plain = app_module.AnalysisState.from_payload({"active_lenses": [], "lens_versions": {}, "doc_type": "PAD"})
    assert "asset-appropriate design horizon" not in app_module.build_lens_stage_context(plain, 2)["prompt"]


def _valid_climate_stage3_payload():
    # A "usable" climate diagnostic (both interaction directions) so the normal
    # Stage 3 climate prefix is exercised rather than the failure fallback branch.
    return _add_specific_climate_paths({"lenses": [{
        "lens_id": "climate", "applicability": "material", "materiality_level": "high",
        "materiality_summary": "Material climate-FCV interactions affect delivery and inclusion.",
        "interaction_readout": [
            {"direction_id": "climate-fcv-on-project", "summary": "Flood and insecurity could disrupt delivery."},
            {"direction_id": "project-on-climate-fcv", "summary": "Benefit rules could affect resilience and trust."},
        ],
        "readout_sections": [], "other_pathways": [],
    }], "findings": []})


def test_climate_stage3_does_not_request_wider_fcv_context():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    prompt = app_module.build_lens_stage_context(
        state, 3, lens_diagnostic=_valid_climate_stage3_payload())["prompt"]
    assert "wider_fcv_context" not in prompt


def test_stage3_climate_prompt_has_cerc_cdrs_and_authority_basis_guardrails():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    prompt = app_module.build_lens_stage_context(
        state, 3, lens_diagnostic=_valid_climate_stage3_payload())["prompt"]
    low = prompt.lower()
    assert "cerc" in low
    assert "named eligible emergency" in low or "activation pathway" in low   # 12.5
    assert "cdrs" in low                                                       # 12.9
    assert "af finances" in low or "what the af finances" in low              # 12.9 AF scoping
    assert "does not auto-restart" in low or "does not auto" in low           # 12.9 restructuring
    assert "phase level" in low or "phase-level" in low                       # 12.9 MPA
    assert "authority_basis" in prompt                                        # 5.5 tag


@pytest.mark.parametrize(
    ("response_text", "expected_status"),
    [
        ("no diagnostic delimiters", "missing_delimiters"),
        (
            app_module.LENS_DIAGNOSTIC_START
            + "{not valid json}"
            + app_module.LENS_DIAGNOSTIC_END,
            "invalid_json",
        ),
    ],
)
def test_lens_recovery_structure_classifies_unparseable_responses(
    response_text, expected_status
):
    summary = app_module.lens_recovery_structure(
        response_text,
        {"error": True, "message": "Recovery invalid."},
        ["climate"],
    )

    assert summary["json_status"] == expected_status
    assert summary["climate_entry_present"] is False
    assert summary["materiality_present"] is False
    assert summary["recognized_interactions"] == []


def test_lens_recovery_structure_reports_only_allowlisted_shape():
    sentinel = "SECRET PROJECT EVIDENCE MUST NOT LEAK"
    raw_payload = {
        "lenses": [{
            "lens_id": "climate",
            "materiality_summary": sentinel,
            "interaction_readout": [{
                "direction_id": "climate-fcv-on-project",
                "summary": sentinel,
                "untrusted_key": sentinel,
            }],
            "untrusted_key": sentinel,
        }],
        "findings": [{"evidence": [sentinel]}],
        "untrusted_key": sentinel,
    }
    response_text = (
        app_module.LENS_DIAGNOSTIC_START
        + json.dumps(raw_payload)
        + app_module.LENS_DIAGNOSTIC_END
    )
    normalized = {
        "lenses": [{
            "lens_id": "climate",
            "materiality_level": "",
            "materiality_summary": sentinel,
            "interaction_readout": [{
                "direction_id": "climate-fcv-on-project",
                "summary": sentinel,
            }],
        }],
        "findings": [],
    }

    summary = app_module.lens_recovery_structure(
        response_text, normalized, ["climate"]
    )
    serialized = json.dumps(summary, sort_keys=True)

    assert summary["json_status"] == "valid_object"
    assert summary["lenses_list"] is True
    assert summary["lens_count"] == 1
    assert summary["findings_list"] is True
    assert summary["finding_count"] == 1
    assert summary["climate_entry_present"] is True
    assert summary["materiality_present"] is False
    assert summary["materiality_valid"] is False
    assert summary["recognized_interactions"] == [
        "climate-fcv-on-project"
    ]
    assert summary["missing_required_interactions"] == []
    assert sentinel not in serialized
    assert "untrusted_key" not in serialized


def test_lens_recovery_structure_reports_missing_required_interaction():
    payload = {
        "lenses": [{
            "lens_id": "climate",
            "materiality_level": "medium",
            "materiality_summary": "Material interactions.",
            "interaction_readout": [{
                "direction_id": "climate-fcv-on-project",
                "summary": "Delivery risk.",
            }],
        }],
        "findings": [],
    }
    response_text = (
        app_module.LENS_DIAGNOSTIC_START
        + json.dumps(payload)
        + app_module.LENS_DIAGNOSTIC_END
    )

    summary = app_module.lens_recovery_structure(
        response_text, payload, ["climate"]
    )

    assert summary["materiality_valid"] is True
    assert summary["missing_required_interactions"] == [
        "project-on-climate-fcv"
    ]
def test_valid_inline_lens_diagnostic_bypasses_recovery(monkeypatch):
    state = app_module.AnalysisState.from_payload({"active_lenses": ["climate"]})
    active_lenses = app_module.build_lens_stage_context(state, stage=2)["active_lenses"]
    inline_payload = {
        "lenses": [{
            "lens_id": "climate",
            "applicability": "material",
            "materiality_level": "medium",
            "materiality_summary": "Flood and conflict pressures affect delivery.",
            "integration_level": "partly_integrated",
            "integration_rating": "Adequate",
            "integration_summary": "Climate-aware but allocation untreated.",
            "strengths_weaknesses": [{
                "side": "strength",
                "title": "Climate-aware design",
                "text": "The design recognizes material climate pressures.",
            }],
            "reflections": [{
                "question_key": "cq2_maladaptation",
                "title": "Maladaptation and lock-in",
                "status_cue": "partial gap",
                "text": "Siting is treated as engineering, not allocation.",
            }],
            "interaction_readout": [
                {
                    "direction_id": "climate-fcv-on-project",
                    "summary": "Flood and insecurity disrupt delivery.",
                },
                {
                    "direction_id": "project-on-climate-fcv",
                    "summary": "Benefit rules affect trust and access.",
                },
            ],
            "readout_sections": [],
            "additional_pathways": [],
        }],
        "findings": [],
    }
    stage2_output = (
        "Visible Stage 2 assessment\n"
        + app_module.LENS_DIAGNOSTIC_START
        + json.dumps(_add_specific_climate_paths(inline_payload))
        + app_module.LENS_DIAGNOSTIC_END
    )
    monkeypatch.setattr(
        app_module,
        "repair_lens_diagnostic",
        lambda *args, **kwargs: (_ for _ in ()).throw(
            AssertionError("recovery should not run for valid inline output")
        ),
    )

    diagnostic, recovered, failure = app_module.extract_or_repair_lens_diagnostic(
        stage2_output,
        active_lenses,
        [],
    )

    assert recovered is False
    assert failure == ""
    assert diagnostic["lenses"][0]["materiality_level"] == "medium"


def test_invalid_climate_diagnostic_gates_stage3_lens_instructions():
    state = app_module.AnalysisState.from_payload({"active_lenses": ["climate"]})
    invalid = {
        "error": True,
        "message": "Lens diagnostic block was not produced.",
        "lenses": [],
        "findings": [],
    }

    context = app_module.build_lens_stage_context(
        state,
        3,
        lens_diagnostic=invalid,
    )

    assert "validated sector-lens diagnostic is unavailable" in context["prompt"]
    assert "Preserve normal core-only Stage 3 behavior" in context["prompt"]
    assert "do not run the lightweight Climate-FCV check" in context["prompt"]
    assert "Treat the validated Climate-FCV materiality" not in context["prompt"]
    assert "two-way Climate-FCV interaction readout" not in context["prompt"]
def test_lens_recovery_client_has_bounded_timeout_and_no_sdk_retries(monkeypatch):
    captured = {}
    sentinel = object()

    def fake_anthropic(**kwargs):
        captured.update(kwargs)
        return sentinel

    monkeypatch.setattr(app_module.anthropic, "Anthropic", fake_anthropic)
    monkeypatch.setattr(app_module, "_lens_recovery_client", None, raising=False)

    client = app_module.get_lens_recovery_client()

    assert client is sentinel
    assert captured["max_retries"] == 0
    assert captured["timeout"].connect == 10.0
    assert captured["timeout"].read == 120.0


def test_frontend_persists_and_submits_lens_context_sources():
    html = (Path(app_module.__file__).parent / "index.html").read_text(
        encoding="utf-8"
    )

    assert "lensContextSources" in html
    assert "lens_context_sources:lensContextSources" in html
    assert "lens_context_sources: lensContextSources || []" in html


def test_large_climate_readout_respects_stage3_platform_budget():
    state = app_module.AnalysisState.from_payload({"active_lenses": ["climate"]})
    item_ids = [
        "social-cohesion-inclusion",
        "institutional-capacity-legitimacy",
        "livelihoods-opportunity",
    ]
    payload = {"lenses": [{
        "lens_id": "climate",
        "applicability": "material",
        "materiality_level": "medium",
        "materiality_summary": "m" * 600,
        "analysis_emphasis": ["e" * 100] * 5,
        "evidence": ["v" * 500] * 5,
        "source_ids": ["peace-social-dividends"],
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "Flood and insecurity could disrupt delivery.",
        }, {
            "direction_id": "project-on-climate-fcv",
            "summary": "Benefit rules could affect resilience and trust.",
        }],
        "readout_sections": [{
            "section_id": "invest-in",
            "items": [{
                "item_id": item_id,
                "status": "supported",
                "mechanism": "x" * 500,
                "evidence": ["y" * 500] * 5,
                "evidence_gap": "z" * 500,
                "trade_off": "t" * 500,
                "source_ids": ["peace-social-dividends"],
            } for item_id in item_ids],
        }],
        "other_pathways": [{
            "pathway": f"pathway-{index}",
            "status": "potential",
            "reason": "r" * 500,
        } for index in range(10)],
    }], "findings": []}

    payload = _add_specific_climate_paths(payload)
    context = app_module.build_lens_stage_context(state, 3, lens_diagnostic=payload)

    assert context["estimated_tokens"] <= 1600
    assert context["truncated"] is True


def test_active_climate_stage3_integrates_opening_and_uses_flexible_mix():
    module_root = Path(app_module.__file__).parent / "sector_lenses" / "modules"
    registry = load_registry(module_root)
    state = app_module.AnalysisState.from_payload({"active_lenses": ["climate"]})
    diagnostic = {"lenses": [{
        "lens_id": "climate",
        "applicability": "material",
        "materiality_level": "medium",
        "materiality_summary": "Drought and FCV pressures affect delivery.",
        "analysis_emphasis": ["adaptation"],
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "Drought and insecurity could disrupt delivery.",
        }, {
            "direction_id": "project-on-climate-fcv",
            "summary": "Benefit rules could affect resilience and trust.",
        }],
        "readout_sections": [],
        "other_pathways": [],
    }], "findings": []}

    diagnostic = _add_specific_climate_paths(diagnostic)
    context = app_module.build_lens_stage_context(
        state, 3, registry, diagnostic, []
    )
    prompt = context["prompt"].lower()

    assert "opening assessment" in prompt
    assert "operational context" in prompt
    assert "maximum of five substantive priorities" in prompt
    assert "not a quota" in prompt
    assert "may contain more climate-linked" in prompt
    assert "single existing priority list" in prompt


def test_climate_stage3_integrates_narrative_and_qualitative_dividends():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"]
    })
    diagnostic = _add_specific_climate_paths({"lenses": [{
        "lens_id": "climate",
        "applicability": "material",
        "materiality_level": "medium",
        "materiality_summary": "Flood and FCV pressures affect delivery.",
        "interaction_readout": [
            {"direction_id": "climate-fcv-on-project", "summary": "A"},
            {"direction_id": "project-on-climate-fcv", "summary": "B"},
        ],
        "readout_sections": [],
        "additional_pathways": [],
    }], "findings": []})

    context = app_module.build_lens_stage_context(
        state, 3, lens_diagnostic=diagnostic
    )
    prompt = context["prompt"]

    for value in (
        "bold opening assessment",
        "operational context",
        "strengths",
        "gaps",
        "FCV sensitivity",
        "FCV responsiveness",
        "two substantive interaction narratives",
        "qualitative Climate, peace and social dividends synthesis",
        "Do not produce dividend cards",
        "no more than five substantive priorities",
        "Adaptation and resilience are primary",
        "deep mitigation only when",
        "climate-fcv-on-project-1",
    ):
        assert value in prompt
    assert context["estimated_tokens"] <= 1600


def test_south_sudan_dual_use_fixture_crosses_stage3_and_docx_pipeline():
    from docx import Document

    fixture = json.loads(SOUTH_SUDAN_FIXTURE.read_text(encoding="utf-8"))
    research = app_module.normalize_climate_research_bundle(
        fixture["research_bundle"]
    )
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"],
        "lens_versions": {"climate": "1.1.0"},
    })
    context = app_module.build_lens_stage_context(
        state,
        3,
        lens_diagnostic=fixture["diagnostic"],
        lens_context_sources=research["sources"],
        climate_research=research,
    )
    diagnostic = context["lens_diagnostic"]
    climate = app_module.climate_lens_entry(diagnostic)

    assert app_module.lens_diagnostic_failure_message(
        diagnostic, ["climate"]
    ) == ""
    assert climate["materiality_level"] == fixture["expected"]["materiality"]
    assert len(climate["interaction_readout"]) == (
        fixture["expected"]["interaction_directions"]
    )
    pathways = [
        pathway
        for interaction in climate["interaction_readout"]
        for pathway in interaction["pathways"]
    ]
    assert len(pathways) >= fixture["expected"]["minimum_specific_pathways"]
    assert {
        horizon
        for pathway in pathways
        for horizon in pathway["time_horizons"]
    } == set(fixture["expected"]["time_horizons"])
    assert "climate-fcv-on-project-1" in context["prompt"]
    assert "project-on-climate-fcv-1" in context["prompt"]

    stage3_text = (
        "%%%JSON_START%%%"
        + json.dumps(fixture["stage3_block"])
        + "%%%JSON_END%%%"
    )
    parsed = app_module.extract_priorities(
        stage3_text,
        active_lens_ids=["climate"],
        lens_diagnostic=diagnostic,
    )

    assert parsed["error"] is False
    assert len(parsed["priorities"]) == (
        fixture["expected"]["substantive_priorities"]
    )
    assert all(
        priority.get("climate_links")
        for priority in parsed["priorities"]
    )
    assert {
        priority["climate_links"]["status"]
        for priority in parsed["priorities"]
    } == {"linked", "no-material-pathway"}

    response = app_module.app.test_client().post(
        "/api/download-report",
        json={
            "summary": (
                "# South Sudan synthetic dual-use test\n"
                "## Executive summary\n"
                "**Climate-FCV pressures shape access and governance.**"
            ),
            "priorities": parsed["priorities"],
            "fcv_rating": parsed["fcv_rating"],
            "fcv_responsiveness_rating": (
                parsed["fcv_responsiveness_rating"]
            ),
            "sensitivity_summary": parsed["sensitivity_summary"],
            "responsiveness_summary": parsed["responsiveness_summary"],
            "active_lenses": [{
                "id": "climate",
                "version": "1.1.0",
                "position": "primary",
            }],
            "lens_diagnostic": diagnostic,
            "lens_context_sources": research["sources"],
            "metadata": {"date_str": "23 July 2026"},
        },
    )

    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "How climate and FCV dynamics could affect this project" in text
    assert "How this project could affect climate and FCV dynamics" in text
    assert "Key locations and components:" in text
    assert "over the life of the assets" in text
    # Redesign: core-questions + strengths/weaknesses replace the standalone dividend section.
    assert "Core climate and FCV questions" in text
    assert "How the design holds up on climate and FCV" in text
    assert "Wider FCV context" not in text
    assert "Differentiated approach note" not in text


def test_docx_climate_notice_title_and_prose_interactions():
    """DOCX: notice heading, interaction direction labels, source signpost, prose strip."""
    from docx import Document

    fixture = json.loads(SOUTH_SUDAN_FIXTURE.read_text(encoding="utf-8"))
    research = app_module.normalize_climate_research_bundle(
        fixture["research_bundle"]
    )
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"],
        "lens_versions": {"climate": "1.1.0"},
    })
    context = app_module.build_lens_stage_context(
        state,
        3,
        lens_diagnostic=fixture["diagnostic"],
        lens_context_sources=research["sources"],
        climate_research=research,
    )
    diagnostic = context["lens_diagnostic"]
    climate = app_module.climate_lens_entry(diagnostic)
    stage3_text = (
        "%%%JSON_START%%%"
        + json.dumps(fixture["stage3_block"])
        + "%%%JSON_END%%%"
    )
    parsed = app_module.extract_priorities(
        stage3_text,
        active_lens_ids=["climate"],
        lens_diagnostic=diagnostic,
    )
    response = app_module.app.test_client().post(
        "/api/download-report",
        json={
            "summary": (
                "# South Sudan climate DOCX prose test\n"
                "## Executive summary\n"
                "**Climate-FCV pressures shape access and governance.**"
            ),
            "priorities": parsed["priorities"],
            "fcv_rating": parsed["fcv_rating"],
            "fcv_responsiveness_rating": parsed["fcv_responsiveness_rating"],
            "sensitivity_summary": parsed["sensitivity_summary"],
            "responsiveness_summary": parsed["responsiveness_summary"],
            "active_lenses": [{
                "id": "climate",
                "version": "1.1.0",
                "position": "primary",
            }],
            "lens_diagnostic": diagnostic,
            "lens_context_sources": research["sources"],
            "metadata": {"date_str": "23 July 2026"},
        },
    )
    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "How relevant is climate to this project?" in text
    assert "How climate and FCV dynamics could affect this project" in text
    assert "How this project could affect climate and FCV dynamics" in text
    assert "Defueling Conflict" in text


def test_climate_stage3_does_not_duplicate_lightweight_check():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"]
    })
    diagnostic = _add_specific_climate_paths({"lenses": [{
        "lens_id": "climate",
        "materiality_level": "low",
        "materiality_summary": "Limited materiality.",
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "A bounded interaction.",
        }],
    }], "findings": []})

    prompt = app_module.build_lens_stage_context(
        state, 3, lens_diagnostic=diagnostic
    )["prompt"]

    assert "lightweight conditional Climate-FCV check" not in prompt
    assert prompt.count("two substantive interaction narratives") == 1


def test_core_only_stage3_keeps_four_to_five_rule():
    context = app_module.build_lens_stage_context(
        app_module.AnalysisState.from_payload({}), 3
    )

    assert context["prompt"] == ""
    assert "4-5 priorities total" in app_module.DEFAULT_PROMPTS["3"]
    assert "Climate-FCV Nexus" in app_module.DEFAULT_PROMPTS["2"]


def test_priority_parser_derives_climate_provenance_from_valid_links():
    diagnostic = _add_specific_climate_paths({"lenses": [{
        "lens_id": "climate",
        "materiality_level": "medium",
        "materiality_summary": "Material interactions.",
        "interaction_readout": [
            {"direction_id": "climate-fcv-on-project", "summary": "A"},
            {"direction_id": "project-on-climate-fcv", "summary": "B"},
        ],
        "readout_sections": [],
        "additional_pathways": [],
    }], "findings": []})
    pathway_id = diagnostic["lenses"][0]["interaction_readout"][0][
        "pathways"
    ][0]["pathway_id"]
    block = {
        "fcv_rating": "Moderate",
        "fcv_responsiveness_rating": "Emerging",
        "sensitivity_summary": "Summary",
        "responsiveness_summary": "Summary",
        "risk_exposure": {"risks_to": "A", "risks_from": "B"},
        "priorities": [{
            "title": "Protect seasonal access",
            "the_gap": "Access safeguards are missing.",
            "why_it_matters": "Seasonal users could lose access.",
            "actions": [],
            "climate_links": {
                "status": "linked",
                "interaction_pathway_ids": [pathway_id],
                "dividend_pathway_ids": [],
                "finding_ids": [],
                "contribution": "The priority protects inclusive access.",
                "strengthening_effect": "It preserves adaptive options.",
                "reason": "",
            },
        }],
    }

    parsed = app_module.extract_priorities(
        "%%%JSON_START%%%" + json.dumps(block) + "%%%JSON_END%%%",
        active_lens_ids=["climate"],
        lens_diagnostic=diagnostic,
    )

    assert parsed["error"] is False
    assert parsed["priorities"][0]["lens_ids"] == ["climate"]
    assert parsed["priorities"][0]["climate_links"]["status"] == "linked"


def test_every_climate_priority_missing_link_degrades_gracefully():
    # Previously this test asserted error=True when climate_links was absent;
    # graceful-degradation contract (Task 1): priorities are kept, climate tag
    # is omitted, and the unlinked counter increments instead of a hard fail.
    diagnostic = _add_specific_climate_paths({"lenses": [{
        "lens_id": "climate",
        "materiality_level": "low",
        "materiality_summary": "Limited but specific interaction.",
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "A",
        }],
    }], "findings": []})
    block = {
        "fcv_rating": "Moderate",
        "fcv_responsiveness_rating": "Emerging",
        "sensitivity_summary": "Summary",
        "responsiveness_summary": "Summary",
        "risk_exposure": {"risks_to": "A", "risks_from": "B"},
        "priorities": [{
            "title": "Core SEA/SH safeguard",
            "the_gap": "A core gap.",
            "why_it_matters": "Material on FCV grounds.",
            "actions": [],
        }],
    }

    parsed = app_module.extract_priorities(
        "%%%JSON_START%%%" + json.dumps(block) + "%%%JSON_END%%%",
        active_lens_ids=["climate"],
        lens_diagnostic=diagnostic,
    )

    assert parsed["error"] is False
    assert len(parsed["priorities"]) == 1
    assert parsed["climate_unlinked"] >= 1
    assert "climate" not in parsed["priorities"][0]["lens_ids"]
    assert parsed["priorities"][0]["climate_links"] is None


def test_climate_active_research_plan_balances_core_and_climate():
    plan = app_module.build_stage1_research_plan(
        active_lens_ids=["climate"],
        country="South Sudan",
        sector="Natural resources",
        doc_parts=[{
            "label": "PROJECT DOCUMENT",
            "name": "Concept Note",
            "raw_text": (
                "Sites: Upper Nile and Jonglei. "
                "The project rehabilitates landing sites and conservancies."
            ),
        }],
    )

    assert plan["core"] == {"max_tokens": 4000, "max_uses": 3}
    assert plan["climate"]["enabled"] is True
    assert "Upper Nile" in plan["project_profile"]["document_excerpt"]
    assert plan["project_profile"]["documents"] == ["Concept Note"]


def test_core_only_research_plan_preserves_current_budget():
    plan = app_module.build_stage1_research_plan(
        active_lens_ids=[],
        country="Exampleland",
        sector="Water",
        doc_parts=[],
    )

    assert plan["core"] == {"max_tokens": 5500, "max_uses": 4}
    assert plan["climate"]["enabled"] is False
    assert plan["project_profile"]["document_excerpt"] == ""


def test_express_and_step_routes_emit_climate_research_context():
    source = Path(app_module.__file__).read_text(encoding="utf-8")

    assert source.count("'climate_research': climate_research") >= 2
    assert source.count("format_climate_research_context(climate_research)") >= 2
    assert source.count("_iter_stage1_research(") >= 3
    assert source.count("research_plan, assessment_id") >= 2


def test_stage3_climate_prompt_uses_prose_and_wider_context():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    diagnostic = _add_specific_climate_paths({"lenses": [{
        "lens_id": "climate",
        "materiality_level": "high",
        "materiality_summary": "Flood and FCV pressures are central.",
        "interaction_readout": [
            {"direction_id": "climate-fcv-on-project", "summary": "A"},
            {"direction_id": "project-on-climate-fcv", "summary": "B"},
        ],
    }], "findings": []})
    ctx = app_module.build_lens_stage_context(state, 3, lens_diagnostic=diagnostic)
    prompt = ctx["prompt"]
    # Phase 4 (Task 4.1): the dedicated module no longer surfaces wider_fcv_context.
    assert "wider_fcv_context" not in prompt
    assert "causal strip" not in prompt.lower()
    assert "prose" in prompt.lower()


def test_climate_integration_payload_helper():
    diagnostic = {"lenses": [{"lens_id": "climate",
                              "integration_level": "partly_integrated",
                              "integration_rating": "Adequate",
                              "integration_summary": "Aware but allocation untreated."}]}
    out = app_module.climate_integration_payload(diagnostic)
    assert out == {"level": "partly_integrated", "rating": "Adequate",
                   "summary": "Aware but allocation untreated."}
    assert app_module.climate_integration_payload({"lenses": []}) is None
    assert app_module.climate_integration_payload({"lenses": [{"lens_id": "climate"}]}) is None


def test_climate_prompts_carry_opcs_guardrails():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    s2 = app_module.build_lens_stage_context(
        state, 2, climate_research={"status": "failed", "attempts": 0,
                                    "sources": [], "claims": [], "failure_reason": ""},
    )["prompt"]
    assert "POLICY BOUNDARY" in s2
    assert "do not apply IPF/ESF terms" in s2
    assert "policy_status" in s2
    assert "specialist_referral" in s2
    # Build a valid climate diagnostic (routes to the normal Stage 3 branch, not failure)
    diagnostic = {"lenses": [{
        "lens_id": "climate",
        "applicability": "material",
        "materiality_level": "high",
        "materiality_summary": "Flood and fragility pressures are central to this project.",
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "Flood and insecurity could disrupt delivery.",
        }, {
            "direction_id": "project-on-climate-fcv",
            "summary": "Benefit rules could affect resilience and trust.",
        }],
        "readout_sections": [],
        "additional_pathways": [],
    }], "findings": []}
    diagnostic = _add_specific_climate_paths(diagnostic)
    s3 = app_module.build_lens_stage_context(state, 3, lens_diagnostic=diagnostic)["prompt"]
    # Guard: must be exercising the valid branch (branch 2), not the failure branch
    assert "Preserve normal" not in s3
    assert "policy_status" in s3
    assert "does not determine ESF" in s3


def test_docx_climate_reflections_integration_wider_fcv_boundary_compliance():
    """DOCX: reflections heading, integration line, wider FCV section, policy boundary, per-priority compliance, order."""
    from docx import Document

    fixture = json.loads(SOUTH_SUDAN_FIXTURE.read_text(encoding="utf-8"))
    research = app_module.normalize_climate_research_bundle(
        fixture["research_bundle"]
    )
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"],
        "lens_versions": {"climate": "1.1.0"},
    })
    context = app_module.build_lens_stage_context(
        state,
        3,
        lens_diagnostic=fixture["diagnostic"],
        lens_context_sources=research["sources"],
        climate_research=research,
    )
    diagnostic = context["lens_diagnostic"]
    stage3_text = (
        "%%%JSON_START%%%"
        + json.dumps(fixture["stage3_block"])
        + "%%%JSON_END%%%"
    )
    parsed = app_module.extract_priorities(
        stage3_text,
        active_lens_ids=["climate"],
        lens_diagnostic=diagnostic,
    )
    response = app_module.app.test_client().post(
        "/api/download-report",
        json={
            "summary": (
                "# South Sudan climate reflections DOCX test\n"
                "## Executive summary\n"
                "**Climate-FCV pressures shape access and governance.**"
            ),
            "priorities": parsed["priorities"],
            "fcv_rating": parsed["fcv_rating"],
            "fcv_responsiveness_rating": parsed["fcv_responsiveness_rating"],
            "sensitivity_summary": parsed["sensitivity_summary"],
            "responsiveness_summary": parsed["responsiveness_summary"],
            "wider_fcv_context": fixture["stage3_block"].get("wider_fcv_context", ""),
            "active_lenses": [{
                "id": "climate",
                "version": "1.1.0",
                "position": "primary",
            }],
            "lens_diagnostic": diagnostic,
            "lens_context_sources": research["sources"],
            "metadata": {"date_str": "23 July 2026"},
        },
    )
    assert response.status_code == 200
    document = Document(io.BytesIO(response.data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    # Redesign contract: core-questions section (with source lines) + strengths/weaknesses;
    # standalone reflections/dividends/wider-FCV DOCX sections are gone in module mode.
    assert "Core climate and FCV questions" in text
    assert "How well does the project integrate climate and FCV?" in text
    assert "How the design holds up on climate and FCV" in text        # strengths/weaknesses
    assert "Wider FCV context" not in text
    assert "does not determine ESF or ESS compliance" in text          # policy boundary
    assert "Task Team E&S specialist" in text                          # specialist_referral route

    # order: integration line -> strengths&weaknesses -> core questions (interactions + answers)
    i_sw = text.index("How the design holds up on climate and FCV")
    i_q = text.index("Core climate and FCV questions")
    i_int = text.index("How climate and FCV dynamics could affect this project")
    assert i_sw < i_q < i_int


def test_stage3_sse_payloads_include_wider_fcv_context():
    """C1: both Stage-3 SSE payload construction paths reference wider_fcv_context."""
    import re
    app_src = (
        Path(__file__).resolve().parents[1] / "app.py"
    ).read_text(encoding="utf-8")
    # Step-by-step: done_data['wider_fcv_context'] inside the elif stage == 3 block
    assert "done_data['wider_fcv_context']" in app_src, (
        "Step-by-step Stage-3 done_data missing wider_fcv_context"
    )
    # Express: 'wider_fcv_context': parsed.get('wider_fcv_context') in stage_done:3 event
    # The express stage_done line contains both stage_done and wider_fcv_context
    express_match = re.search(
        r"stage_done.*wider_fcv_context|wider_fcv_context.*stage_done",
        app_src,
    )
    assert express_match, (
        "Express Stage-3 stage_done event missing wider_fcv_context"
    )


def test_extract_priorities_returns_climate_link_counts_on_success():
    """Task 3 contract: extract_priorities returns climate_unlinked + climate_total keys
    for climate runs; both default to 0 for a non-climate run."""
    # --- climate run: one priority with a valid linked climate_links ---
    diag = _add_specific_climate_paths({"lenses": [{
        "lens_id": "climate",
        "materiality_level": "high",
        "materiality_summary": "s",
        "interaction_readout": [{
            "direction_id": "climate-fcv-on-project",
            "summary": "A",
        }],
    }], "findings": []})
    block = {
        "fcv_rating": "Moderate", "fcv_responsiveness_rating": "Moderate",
        "sensitivity_summary": "s", "responsiveness_summary": "r",
        "risk_exposure": {"risks_to": "x", "risks_from": "y"},
        "priorities": [{
            "title": "Climate priority in Juba",
            "fcv_dimension": "Contextual awareness",
            "tag": "[S]",
            "refresh_shift": "Shift A: Anticipate",
            "risk_level": "High",
            "the_gap": "Gap in Juba area.",
            "why_it_matters": "Why it matters.",
            "actions": [{"document_element": "PAD", "guidance": "Do X in Juba.", "suggested_language": ""}],
            "who_acts": "TTL",
            "when": "before appraisal",
            "action_timing": "required-before-appraisal",
            "resources": "r",
            "pad_sections": "SORT",
            "implementation_note": "n",
            "cpf_alignment": None,
            "climate_links": {
                "status": "linked",
                "interaction_pathway_ids": ["climate-fcv-on-project-1"],
                "contribution": "c",
                "strengthening_effect": "s",
            },
        }],
    }
    result = app_module.extract_priorities(
        "%%%JSON_START%%%" + json.dumps(block) + "%%%JSON_END%%%",
        active_lens_ids=["climate"],
        lens_diagnostic=diag,
    )
    assert result["error"] is False
    assert "climate_unlinked" in result
    assert "climate_total" in result
    assert result["climate_total"] == 1

    # --- non-climate run: counts are zero ---
    plain_block = {
        "fcv_rating": "Moderate", "fcv_responsiveness_rating": "Moderate",
        "sensitivity_summary": "s", "responsiveness_summary": "r",
        "risk_exposure": {"risks_to": "x", "risks_from": "y"},
        "priorities": [{
            "title": "FCV priority in Juba",
            "fcv_dimension": "Contextual awareness",
            "tag": "[S]",
            "refresh_shift": "Shift A: Anticipate",
            "risk_level": "High",
            "the_gap": "Gap in Juba area.",
            "why_it_matters": "Why.",
            "actions": [{"document_element": "PAD", "guidance": "Do X in Juba.", "suggested_language": ""}],
            "who_acts": "TTL",
            "when": "before appraisal",
            "action_timing": "required-before-appraisal",
            "resources": "r",
            "pad_sections": "SORT",
            "implementation_note": "n",
            "cpf_alignment": None,
        }],
    }
    plain_result = app_module.extract_priorities(
        "%%%JSON_START%%%" + json.dumps(plain_block) + "%%%JSON_END%%%",
    )
    assert plain_result["error"] is False
    assert plain_result["climate_unlinked"] == 0
    assert plain_result["climate_total"] == 0


# --- graceful degradation: single-direction climate diagnostic ---
def _single_direction_climate_diag(level="high"):
    return {
        "lenses": [{
            "lens_id": "climate",
            "materiality_level": level,
            "materiality_summary": "Water scarcity and armed conflict interact across the fisheries sites.",
            "interaction_readout": [{
                "direction_id": "climate-fcv-on-project",
                "summary": "Drought and flooding cut access to landing sites.",
                "pathways": [{
                    "pathway_id": "climate-fcv-on-project-1",
                    "pressure": "drought",
                    "mechanism": "road closure",
                    "project_implication": "delayed works",
                    "design_response": "seasonal windows",
                }],
            }],
        }],
        "findings": [],
    }


def test_climate_diagnostic_usable_with_single_specific_direction():
    # High/Medium materiality with ONE fully-specified direction is now usable
    # (graceful degradation), not discarded for lacking the second direction.
    assert app_module.lens_diagnostic_failure_message(
        _single_direction_climate_diag("high"), ["climate"]) == ""
    assert app_module.lens_diagnostic_failure_message(
        _single_direction_climate_diag("medium"), ["climate"]) == ""


def test_climate_diagnostic_still_fails_without_any_specific_direction():
    d = _single_direction_climate_diag("high")
    d["lenses"][0]["interaction_readout"][0]["pathways"] = []
    assert app_module.lens_diagnostic_failure_message(d, ["climate"]) != ""


def test_climate_diagnostic_still_fails_without_summary():
    d = _single_direction_climate_diag("high")
    d["lenses"][0]["materiality_summary"] = ""
    assert app_module.lens_diagnostic_failure_message(d, ["climate"]) != ""


def test_climate_diagnostic_still_fails_on_invalid_materiality():
    d = _single_direction_climate_diag("amazing")
    assert app_module.lens_diagnostic_failure_message(d, ["climate"]) != ""


def test_stage2_climate_prompt_caps_pathways_and_requires_completion():
    state = app_module.AnalysisState.from_payload({
        "active_lenses": ["climate"], "lens_versions": {}, "doc_type": "PAD",
    })
    prompt = app_module.build_lens_stage_context(
        state, 2, climate_research={"status": "failed", "attempts": 0,
                                    "sources": [], "claims": [], "failure_reason": ""},
    )["prompt"]
    assert "one or two pathways" in prompt
    assert "Always complete and close the hidden diagnostic block" in prompt
    assert "one to four pathways" not in prompt


def test_workflow_bridge_streams_events_then_completes():
    # Happy path: bridge yields each workflow event then stops on the sentinel.
    def workflow_events():
        yield "data: one\n\n"
        yield "data: two\n\n"

    with app_module.app.test_request_context("/api/run-express"):
        out = list(
            app_module._stream_workflow_events(
                workflow_events, "assess-x", poll_interval=0.05
            )
        )
    assert out == ["data: one\n\n", "data: two\n\n"]


def test_workflow_bridge_surfaces_crash_without_hanging():
    # A workflow that raises must yield an error event and terminate, never hang.
    def workflow_events():
        yield "data: partial\n\n"
        raise RuntimeError("boom")

    with app_module.app.test_request_context("/api/run-express"):
        out = list(
            app_module._stream_workflow_events(
                workflow_events, "assess-y", poll_interval=0.05
            )
        )
    assert out[0] == "data: partial\n\n"
    assert any('"error"' in chunk and "failed_stage" in chunk for chunk in out)


def test_workflow_bridge_keepalive_and_idle_deadline_on_silent_workflow():
    # A workflow that produces nothing for a while must not hang silently: the
    # bridge emits keepalives and then a clean idle-deadline error (never blocks).
    import time as _time

    def workflow_events():
        _time.sleep(1.0)  # silent window longer than the idle deadline
        yield "data: never-reached\n\n"

    with app_module.app.test_request_context("/api/run-express"):
        out = list(
            app_module._stream_workflow_events(
                workflow_events, "assess-z",
                poll_interval=0.05, idle_deadline=0.2,
            )
        )
    assert any('"keepalive"' in chunk for chunk in out)
    assert any('"error"' in chunk and "failed_stage" in chunk for chunk in out)


def test_workflow_bridge_does_not_kill_slow_but_streaming_run():
    # A run that keeps emitting events (even slowly) must never trip the idle
    # backstop, even when total elapsed exceeds the idle_deadline.
    import time as _time

    def workflow_events():
        for i in range(4):
            _time.sleep(0.15)
            yield f"data: chunk-{i}\n\n"

    with app_module.app.test_request_context("/api/run-express"):
        out = list(
            app_module._stream_workflow_events(
                workflow_events, "assess-w",
                poll_interval=0.05, idle_deadline=0.2,
            )
        )
    # All four workflow chunks delivered; no error/deadline event.
    assert [c for c in out if c.startswith("data: chunk-")] == [
        "data: chunk-0\n\n", "data: chunk-1\n\n",
        "data: chunk-2\n\n", "data: chunk-3\n\n",
    ]
    assert not any('"error"' in chunk for chunk in out)



def _canonical_recovery_fixture():
    return json.loads(SOUTH_SUDAN_FIXTURE.read_text(encoding="utf-8"))["diagnostic"]


def _active_climate_for_recovery():
    state = app_module.AnalysisState.from_payload({"active_lenses": ["climate"]})
    return app_module.build_lens_stage_context(state, stage=2)["active_lenses"]


def test_field_level_climate_recovery_is_bounded_and_requests_missing_path():
    primary = _canonical_recovery_fixture()
    primary["lenses"][0]["integration_summary"] = ""
    repair = _canonical_recovery_fixture()
    repair["lenses"][0]["integration_summary"] = "Repaired integration summary."
    captured = {}

    class FakeMessages:
        def create(self, **kwargs):
            captured.update(kwargs)
            text = (
                app_module.LENS_DIAGNOSTIC_START
                + json.dumps(repair)
                + app_module.LENS_DIAGNOSTIC_END
            )
            return type("Response", (), {
                "content": [type("Text", (), {"text": text})()]
            })()

    client = type("Client", (), {"messages": FakeMessages()})()
    output = (
        app_module.LENS_DIAGNOSTIC_START
        + json.dumps(primary)
        + app_module.LENS_DIAGNOSTIC_END
    )
    events = list(app_module._iter_native_climate_stage2_diagnostic(
        stage2_output=output,
        active_lenses=_active_climate_for_recovery(),
        context_sources=[],
        assessment_id="assessment-field-repair",
        client=client,
        max_seconds=90,
        keepalive_interval=0.01,
    ))

    terminal = events[-1]
    prompt = captured["messages"][0]["content"]
    assert captured["max_tokens"] == 4500
    assert captured["timeout"] == 90
    assert "- lenses.climate.integration_summary" in prompt
    assert "Do not regenerate or rewrite valid fields" in prompt
    assert terminal["error_code"] == ""
    assert terminal["recovered"] is True
    assert terminal["result"]["lenses"][0]["integration_summary"] == (
        "Repaired integration summary."
    )



def test_field_level_climate_recovery_rejects_incomplete_repair():
    primary = _canonical_recovery_fixture()
    primary["lenses"][0]["integration_summary"] = ""

    class FakeMessages:
        def create(self, **kwargs):
            text = (
                app_module.LENS_DIAGNOSTIC_START
                + json.dumps({"lenses": [], "findings": []})
                + app_module.LENS_DIAGNOSTIC_END
            )
            return type("Response", (), {
                "content": [type("Text", (), {"text": text})()]
            })()

    output = (
        app_module.LENS_DIAGNOSTIC_START
        + json.dumps(primary)
        + app_module.LENS_DIAGNOSTIC_END
    )
    events = list(app_module._iter_native_climate_stage2_diagnostic(
        stage2_output=output,
        active_lenses=_active_climate_for_recovery(),
        context_sources=[],
        assessment_id="assessment-invalid-repair",
        client=type("Client", (), {"messages": FakeMessages()})(),
        max_seconds=1,
        keepalive_interval=0.01,
    ))

    assert events[-1]["recovered"] is False
    assert events[-1]["error_code"] == "climate_diagnostic_invalid"



def test_non_climate_recovery_retains_legacy_generic_contract():
    captured = {}

    class FakeMessages:
        def create(self, **kwargs):
            captured.update(kwargs)
            text = (
                app_module.LENS_DIAGNOSTIC_START
                + json.dumps({
                    "lenses": [{
                        "lens_id": "agriculture",
                        "applicability": "material",
                        "materiality_summary": "Material delivery interaction.",
                        "analysis_emphasis": [],
                        "evidence": [],
                        "source_ids": [],
                        "readout_sections": [],
                        "other_pathways": [],
                    }],
                    "findings": [],
                })
                + app_module.LENS_DIAGNOSTIC_END
            )
            return type("Response", (), {
                "content": [type("Text", (), {"text": text})()]
            })()

    diagnostic, recovered = app_module.repair_lens_diagnostic(
        "Visible Stage 2 assessment",
        ["agriculture"],
        {"agriculture": set()},
        {"agriculture": {}},
        client=type("Client", (), {"messages": FakeMessages()})(),
    )

    assert recovered is True
    assert diagnostic["lenses"][0]["lens_id"] == "agriculture"
    assert captured["model"] == "claude-sonnet-4-6"
    assert captured["max_tokens"] == 8000
    assert "Recover only the missing structured sector-lens diagnostic" in (
        captured["messages"][0]["content"]
    )
    assert "timeout" not in captured


def test_verified_express_is_limited_to_climate_only_design_runs():
    climate = app_module.AnalysisState(active_lenses=["climate"])
    mixed = app_module.AnalysisState(active_lenses=["climate", "agriculture"])

    assert app_module._is_verified_climate_express(climate, False) is True
    assert app_module._is_verified_climate_express(climate, True) is False
    assert app_module._is_verified_climate_express(mixed, False) is False


def test_verified_client_builder_uses_server_smoke_profile(monkeypatch):
    monkeypatch.setenv("CLIMATE_VERIFIED_RUN_MODE", "smoke")
    monkeypatch.setattr(app_module, "get_client", lambda: object())
    monkeypatch.setattr(app_module, "get_lens_recovery_client", lambda: object())

    clients = app_module._build_verified_pipeline_clients()

    assert clients.assessment._model == "claude-haiku-4-5-20251001"
    assert clients.reviewer._model == "claude-haiku-4-5-20251001"


def test_verified_runtime_bridge_emits_keepalives_then_result(monkeypatch):
    import time as _time

    expected = {
        "assessment": {"schema_version": "climate-verified-v2.1"},
        "reader": {"executive_readout": "Verified."},
        "source_warnings": [],
    }
    captured = {}

    def fake_run(**kwargs):
        captured.update(kwargs)
        _time.sleep(0.04)
        return expected

    monkeypatch.setattr(app_module, "run_verified_from_doc_parts", fake_run)
    events = list(app_module._iter_verified_climate_assessment(
        doc_parts=[],
        climate_grounding={},
        clients=object(),
        run_id="verified-runtime-test",
        doc_type="PCN",
        instrument_type="IPF",
        keepalive_interval=0.01,
    ))

    assert any(item.get("keepalive") is True for item in events[:-1])
    assert events[-1] == {"result": expected}
    assert captured["doc_type"] == "PCN"
    assert captured["instrument_type"] == "IPF"


def test_verified_runtime_bridge_cancels_after_wall_clock(monkeypatch):
    import time as _time

    captured = {}

    def fake_run(**kwargs):
        captured["cancel_event"] = kwargs["cancel_event"]
        while not kwargs["cancel_event"].is_set():
            _time.sleep(0.005)
        raise RuntimeError("cancelled")

    monkeypatch.setattr(app_module, "run_verified_from_doc_parts", fake_run)
    with pytest.raises(TimeoutError, match="14 minutes"):
        list(app_module._iter_verified_climate_assessment(
            doc_parts=[],
            climate_grounding={},
            clients=object(),
            run_id="verified-timeout-test",
            keepalive_interval=0.005,
            maximum_wait_seconds=0.02,
        ))

    assert captured["cancel_event"].is_set()


def test_verified_climate_ui_contract_is_ranked_and_multidimensional():
    html = (Path(app_module.__file__).parent / "index.html").read_text(
        encoding="utf-8"
    )

    assert "renderClimateVerifiedAssessment" in html
    for dimension in (
        "relevance",
        "sensitivity",
        "responsiveness",
        "operationalization",
    ):
        assert f'data-climate-dimension="${{esc(j.dimension)}}"' in html
    assert "Points to check before the decision meeting" in html
    start = html.index("function renderClimateVerifiedAssessment")
    end = html.index("\n  function ", start + 20)
    body = html[start:end]
    assert "priority.rank" in body
    assert "priority.priority_label" not in body
    assert "High priority" not in body
    assert "Smoke test: validates workflow completion only" in body
    assert (
        "passed the checks but ${admittedCount===1?'was':'were'} held back on review"
        in body
    )
    assert "recommendation_admitted_count" in body
    assert "semantic_reviewer_verdict" in body
    assert "current_document_drafting" in body
    assert "operational_instrument_drafting" in body
    assert "Suggested drafting for the current document" in body
    assert "Suggested drafting for an operational instrument" in body
    # The card leads with the model narrative; the useful structured fields fold
    # into a "Recommendation details" collapsible, and app-internal routing/coded
    # references are dropped from the user view entirely.
    assert "Recommendation details" in body
    assert "pc-narr" in body
    assert "priority_summary" in body
    assert "live_research_count" in body
    assert "drafting_language" not in body
    assert "recommendation_reason_codes" in body


def test_express_route_dispatches_verified_assessment_contract():
    source = Path(app_module.__file__).read_text(encoding="utf-8")

    assert "run_verified_from_doc_parts" in source
    assert "'climate_assessment': verified_assessment" in source
    assert "'climate_reader': verified_reader" in source
    assert "Climate recommendation diagnostics" in source
    assert "'recommendation_diagnostics'" in source


def test_verified_climate_docx_route_uses_canonical_reader():
    sentence = (
        "Verified project evidence supports a material pathway and a bounded response. "
    )
    assessment = {
        "schema_version": "climate-verified-v2.1",
        "run_id": "route-run",
        "bank_release_id": "2026.08",
        "evidence_status": "preview; not approved",
        "executive_readout": (sentence * 50).strip(),
        "judgments": {
            "relevance": {"value": "high", "rationale": "Material."},
            "sensitivity": {"value": "moderate", "rationale": "Partial."},
            "responsiveness": {"value": "emerging", "rationale": "Emerging."},
            "operationalization": {"value": "partial", "rationale": "Partial."},
        },
        "priorities": [],
        "review_readiness_flags": [],
        "validation": {"status": "passed"},
    }

    response = app_module.app.test_client().post(
        "/api/download-report", json={
            "climate_assessment": assessment,
            "climate_reader": {"runtime_mode": "smoke"},
        }
    )

    assert response.status_code == 200
    from docx import Document
    document = Document(io.BytesIO(response.data))
    text = "\n".join(item.text for item in document.paragraphs)
    assert "Climate-FCV judgments" in text
    assert "preview; not approved" in text
    assert "Smoke test: validates workflow completion only" in text


def test_browser_exports_verified_climate_assessment_object():
    html = (Path(app_module.__file__).parent / "index.html").read_text(encoding="utf-8")
    assert "climate_assessment: climateVerifiedAssessment" in html
    assert "renderClimateVerifiedAssessment(climateVerifiedReader)" in html


def test_verified_climate_reader_drives_ui_followon_and_persistence():
    html = (Path(app_module.__file__).parent / "index.html").read_text(
        encoding="utf-8"
    )

    assert "climateVerifiedReader=p.climate_reader" in html
    assert "renderClimateVerifiedAssessment(climateVerifiedReader)" in html
    assert "JSON.stringify(climateVerifiedReader" in html
    assert "climateVerifiedAssessment: climateVerifiedAssessment" in html
    assert "climateVerifiedReader: climateVerifiedReader" in html


def test_verified_climate_express_timeout_covers_full_automatic_review():
    html = (Path(app_module.__file__).parent / "index.html").read_text(encoding="utf-8")
    assert "2:15*60*1000" in html
    assert "2:'15 minutes'" in html
    assert html.count("climateVerifiedAssessment=null") >= 5


def test_climate_only_express_route_returns_verified_v2_without_legacy_stage(monkeypatch):
    monkeypatch.setenv("CLIMATE_VERIFIED_RUN_MODE", "smoke")
    assessment = {
        "schema_version": "climate-verified-v2.1",
        "run_id": "route-v2",
        "bank_release_id": "2026.08",
        "evidence_status": "preview; not approved",
        "executive_readout": "Verified Climate-FCV readout.",
        "judgment_summary": "High relevance; moderate sensitivity.",
        "judgments": {
            "relevance": {"value": "high", "rationale": "Material."},
            "sensitivity": {"value": "moderate", "rationale": "Partial."},
            "responsiveness": {"value": "emerging", "rationale": "Emerging."},
            "operationalization": {"value": "partial", "rationale": "Partial."},
        },
        "priorities": [],
        "review_readiness_flags": [],
        "validation": {"status": "passed"},
    }
    reader = {"executive_readout": assessment["executive_readout"]}

    monkeypatch.setattr(app_module, "get_fast_client", lambda: object())
    monkeypatch.setattr(app_module, "extract_country_name", lambda *_: "South Sudan")
    monkeypatch.setattr(app_module, "extract_sector_name", lambda *_: "Fisheries")
    monkeypatch.setattr(app_module, "build_stage1_research_plan", lambda *_args, **_kwargs: {})
    monkeypatch.setattr(app_module, "_iter_stage1_research", lambda *_args, **_kwargs: iter([{
        "result": {
            "core_brief": "",
            "climate_research": {},
            "lens_context_sources": [],
            "climate_grounding": {},
        }
    }]))
    grounding = {
        "state": "bank-only",
        "content_version": "2026.08",
        "candidate_preview": True,
        "bank_sources": [],
        "bank_evidence_records": [],
        "bank_pathways": [],
        "live_claims": [],
    }
    monkeypatch.setattr(
        app_module, "resolve_climate_grounding", lambda *_args, **_kwargs: (grounding, {})
    )
    monkeypatch.setattr(app_module, "_build_verified_pipeline_clients", lambda: object())
    monkeypatch.setattr(app_module, "_iter_verified_climate_assessment", lambda **_kwargs: iter([{
        "result": {"assessment": assessment, "reader": reader, "source_warnings": []}
    }]))
    monkeypatch.setattr(
        app_module, "_stream_stage",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(AssertionError("legacy stage called")),
    )

    response = app_module.app.test_client().post("/api/run-express", json={
        "assessment_id": "route-v2",
        "documents": [{
            "name": "pcn.txt",
            "type": "text",
            "content": "South Sudan fisheries project with flood and conflict risks.",
            "docRole": "primary",
        }],
        "active_lenses": ["climate"],
        "review_mode": "design",
    })

    body = response.get_data(as_text=True)
    assert response.status_code == 200
    assert '"schema_version": "climate-verified-v2.1"' in body
    assert '"stage_done": 3' in body
    assert '"express_done": true' in body
    assert '"runtime_mode": "smoke"' in body


def test_verified_climate_failure_log_includes_bounded_schema_reason(caplog):
    diagnostic = {
        "stage": "recommendation_compiler",
        "attempt": 1,
        "elapsed_ms": 274,
        "exception_type": "BadRequestError",
        "status_code": 400,
        "prompt_chars": 37495,
        "timeout_seconds": 240,
        "remaining_seconds": 239,
        "provider_error_type": "invalid_request_error",
        "provider_failure_code": "schema_rejected",
        "schema_path": "properties.recommendation_candidates.items.type",
    }

    with caplog.at_level("WARNING"):
        app_module._log_verified_climate_call_failure(diagnostic)

    assert "provider_failure_code=schema_rejected" in caplog.text
    assert "schema_path=properties.recommendation_candidates.items.type" in caplog.text
