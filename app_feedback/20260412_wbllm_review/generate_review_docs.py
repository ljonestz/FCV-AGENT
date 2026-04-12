"""
Generate WB LLM Review Package — v8 Output Quality Review (2026-04-12)

Produces three .docx reference documents from the live source files:
  01_Knowledge_Base_v8.docx         — WB_INSTRUMENT_GUIDE + WB_PROCESS_GUIDE + FCV_GLOSSARY
  02_Stage_Prompts_v8.docx          — DEFAULT_PROMPTS (stages 1, 2, 3, deeper_playbook, followon)
  03_FCV_Reference_Constants_v8.docx — FCV_GUIDE, FCV_OPERATIONAL_MANUAL, FCV_REFRESH_FRAMEWORK,
                                       PLAYBOOKs, STAGE_GUIDANCE_MAP

Run from repo root:
    python app_feedback/20260412_wbllm_review/generate_review_docs.py
"""

import re
import sys
import os
from pathlib import Path

# ── Path setup ───────────────────────────────────────────────────────────────
REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(REPO_ROOT))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")

import background_docs  # pure Python, safe to import directly

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0, 41, 68)  # WB navy
    return p


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_body(doc, text):
    if not text:
        return
    doc.add_paragraph(str(text).strip())


def add_separator(doc):
    doc.add_paragraph("─" * 60)


def add_meta(doc, label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(str(value))


# ── Extract DEFAULT_PROMPTS from app.py (without importing Flask) ─────────────

def extract_default_prompts(app_py_path: Path) -> dict:
    """
    Reads app.py and extracts DEFAULT_PROMPTS using regex on triple-quoted strings.
    Returns a dict of {key: prompt_text}.
    """
    content = app_py_path.read_text(encoding="utf-8")

    # Find start of DEFAULT_PROMPTS block
    start_idx = content.find("DEFAULT_PROMPTS = {")
    if start_idx == -1:
        raise ValueError("DEFAULT_PROMPTS not found in app.py")

    block = content[start_idx:]

    # Match each key and its triple-quoted string value
    pattern = re.compile(r'"([^"]+)":\s*\'\'\'(.*?)\'\'\'', re.DOTALL)
    matches = pattern.findall(block)

    if not matches:
        raise ValueError("No prompt keys found in DEFAULT_PROMPTS block")

    return {key: text.strip() for key, text in matches}


# ── Document 1: Knowledge Base ────────────────────────────────────────────────

def write_knowledge_base(output_path: Path):
    doc = Document()
    add_title(doc, "FCV Project Screener v8 — Knowledge Base")
    doc.add_paragraph(
        "Generated: 2026-04-12 | Source: background_docs.py\n"
        "Contains: WB Instrument Guide, WB Process Guide, FCV Glossary"
    )

    # ── Section 1: WB Instrument Guide ──────────────────────────────────────
    add_h1(doc, "Section 1: WB Instrument Guide")
    doc.add_paragraph(
        "Per-instrument reference covering FCV levers, scope limitations, typical structure, "
        "common FCV considerations, preparation/supervision processes, and policy transitions. "
        "Only the relevant instrument slice is injected into each analysis stage."
    )

    field_labels = {
        "name": "Full Name",
        "description": "Description",
        "fcv_levers": "FCV Levers Available",
        "not_applicable": "Not Applicable / Guardrails",
        "typical_structure": "Typical Structure",
        "common_fcv_considerations": "Common FCV Considerations",
        "preparation_process": "Preparation Process",
        "supervision_process": "Supervision Process",
        "policy_transitions": "Policy Transitions",
        "recent_changes": "Recent Procedural Changes",
        "ost_applicability": "OST Recommendation Applicability",
        "dnh_applicability": "Do No Harm Principle Applicability",
    }

    for instrument_key, instrument_data in background_docs.WB_INSTRUMENT_GUIDE.items():
        add_h2(doc, f"Instrument: {instrument_key}")
        for field, label in field_labels.items():
            value = instrument_data.get(field)
            if not value:
                continue
            add_h3(doc, label)
            if isinstance(value, dict):
                for sub_key, sub_val in value.items():
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(f"{sub_key}: ")
                    run.bold = True
                    p.add_run(str(sub_val))
            else:
                add_body(doc, value)
        add_separator(doc)

    # ── Section 2: WB Process Guide ─────────────────────────────────────────
    add_h1(doc, "Section 2: WB Process Guide")
    doc.add_paragraph(
        "Per-process reference for implementation-stage reviews (MTR, ISR, AF, Restructuring, ICR). "
        "Covers purpose, scope, key policies, typical documents, FCV considerations, common pitfalls, "
        "and backward/forward look guidance."
    )

    process_field_labels = {
        "purpose": "Purpose",
        "scope": "Scope",
        "key_policies": "Key Policies",
        "typical_documents": "Typical Documents",
        "fcv_considerations": "FCV Considerations",
        "common_pitfalls": "Common Pitfalls",
        "backward_forward_look": "Backward / Forward Look",
    }

    for process_key, process_data in background_docs.WB_PROCESS_GUIDE.items():
        add_h2(doc, f"Process: {process_key}")
        for field, label in process_field_labels.items():
            value = process_data.get(field)
            if not value:
                continue
            add_h3(doc, label)
            add_body(doc, value)
        add_separator(doc)

    # ── Section 3: FCV Glossary ──────────────────────────────────────────────
    add_h1(doc, "Section 3: FCV Glossary")
    doc.add_paragraph(
        "29 FCV terms with definitions, measurement approaches, and authoritative sources. "
        "Injected into Stage 2 for consistent terminology. Also served to the frontend "
        "for glossary tooltip rendering."
    )

    for term_key, term_data in background_docs.FCV_GLOSSARY.items():
        add_h2(doc, term_data.get("term", term_key))
        add_h3(doc, "Definition")
        add_body(doc, term_data.get("definition"))
        add_h3(doc, "How Measured / Applied")
        add_body(doc, term_data.get("measurement"))
        source = term_data.get("source")
        if source:
            add_meta(doc, "Source", source)
        doc.add_paragraph()

    doc.save(output_path)
    print(f"  Written: {output_path.name}")


# ── Document 2: Stage Prompts ─────────────────────────────────────────────────

def write_stage_prompts(output_path: Path, prompts: dict):
    doc = Document()
    add_title(doc, "FCV Project Screener v8 — Stage Prompts")
    doc.add_paragraph(
        "Generated: 2026-04-12 | Source: app.py — DEFAULT_PROMPTS\n"
        "Contains: Stage 1 (Context & Extraction), Stage 2 (FCV Assessment), "
        "Stage 3 (Recommendations Note), Go Deeper Playbook, Follow-on"
    )

    prompt_labels = {
        "1": "Stage 1 — Context & Extraction",
        "2": "Stage 2 — FCV Assessment",
        "3": "Stage 3 — Recommendations Note",
        "deeper_playbook": "Go Deeper — FCV Playbook References",
        "followon": "Follow-on Analysis",
        "deeper": "Go Deeper — Alternatives (legacy)",
    }

    display_order = ["1", "2", "3", "deeper_playbook", "followon", "deeper"]

    for key in display_order:
        if key not in prompts:
            continue
        label = prompt_labels.get(key, f"Prompt: {key}")
        add_h1(doc, label)
        doc.add_paragraph(f"Key: \"{key}\"").runs[0].italic = True

        # Add prompt text as a monospace-style paragraph
        text = prompts[key]
        # Split by double newlines for readability; add as paragraphs
        paragraphs = text.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph(para_text)
            p.style = doc.styles["Normal"]

        add_separator(doc)

    doc.save(output_path)
    print(f"  Written: {output_path.name}")


# ── Document 3: FCV Reference Constants ──────────────────────────────────────

def write_fcv_constants(output_path: Path):
    doc = Document()
    add_title(doc, "FCV Project Screener v8 — FCV Reference Constants")
    doc.add_paragraph(
        "Generated: 2026-04-12 | Source: background_docs.py\n"
        "Contains: FCV_GUIDE, FCV_OPERATIONAL_MANUAL, FCV_REFRESH_FRAMEWORK, "
        "PLAYBOOK_DIAGNOSTICS, PLAYBOOK_PREPARATION, PLAYBOOK_IMPLEMENTATION, "
        "PLAYBOOK_CLOSING, STAGE_GUIDANCE_MAP"
    )

    constants = [
        ("FCV_GUIDE", "FCV Framework — Sensitivity, Responsiveness, and Strategic Shifts",
         background_docs.FCV_GUIDE),
        ("FCV_OPERATIONAL_MANUAL", "FCV Operational Manual — Core Framework for FCV-Sensitive Project Design",
         background_docs.FCV_OPERATIONAL_MANUAL),
        ("FCV_REFRESH_FRAMEWORK", "FCV Strategy Refresh Framework (January 2026)",
         background_docs.FCV_REFRESH_FRAMEWORK),
        ("PLAYBOOK_DIAGNOSTICS", "FCV Operational Playbook — Diagnostics Phase",
         background_docs.PLAYBOOK_DIAGNOSTICS),
        ("PLAYBOOK_PREPARATION", "FCV Operational Playbook — Project Preparation Phase",
         background_docs.PLAYBOOK_PREPARATION),
        ("PLAYBOOK_IMPLEMENTATION", "FCV Operational Playbook — Implementation Phase",
         background_docs.PLAYBOOK_IMPLEMENTATION),
        ("PLAYBOOK_CLOSING", "FCV Operational Playbook — Project Closing Phase",
         background_docs.PLAYBOOK_CLOSING),
    ]

    for const_key, const_label, const_value in constants:
        add_h1(doc, const_label)
        doc.add_paragraph(f"Constant: {const_key}").runs[0].italic = True

        if isinstance(const_value, str):
            for para_text in const_value.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)
        elif isinstance(const_value, dict):
            for sub_key, sub_val in const_value.items():
                add_h2(doc, str(sub_key))
                if isinstance(sub_val, dict):
                    for field_key, field_val in sub_val.items():
                        add_h3(doc, field_key)
                        add_body(doc, field_val)
                else:
                    add_body(doc, sub_val)

        add_separator(doc)

    # STAGE_GUIDANCE_MAP
    add_h1(doc, "Stage Guidance Map")
    doc.add_paragraph("Constant: STAGE_GUIDANCE_MAP").runs[0].italic = True
    for stage_key, stage_val in background_docs.STAGE_GUIDANCE_MAP.items():
        add_h2(doc, str(stage_key))
        if isinstance(stage_val, dict):
            for k, v in stage_val.items():
                add_meta(doc, k, v)
        else:
            add_body(doc, stage_val)

    doc.save(output_path)
    print(f"  Written: {output_path.name}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print("FCV Project Screener v8 — Generating WB LLM Review Package")
    print(f"Output directory: {OUTPUT_DIR}\n")

    # Extract prompts
    app_py_path = REPO_ROOT / "app.py"
    print("Extracting DEFAULT_PROMPTS from app.py...")
    prompts = extract_default_prompts(app_py_path)
    print(f"  Found {len(prompts)} prompt keys: {list(prompts.keys())}\n")

    # Document 1
    print("Writing 01_Knowledge_Base_v8.docx...")
    write_knowledge_base(OUTPUT_DIR / "01_Knowledge_Base_v8.docx")

    # Document 2
    print("Writing 02_Stage_Prompts_v8.docx...")
    write_stage_prompts(OUTPUT_DIR / "02_Stage_Prompts_v8.docx", prompts)

    # Document 3
    print("Writing 03_FCV_Reference_Constants_v8.docx...")
    write_fcv_constants(OUTPUT_DIR / "03_FCV_Reference_Constants_v8.docx")

    print("\nDone. Share the following files with the WB LLM:")
    print("  01_Knowledge_Base_v8.docx")
    print("  02_Stage_Prompts_v8.docx")
    print("  03_FCV_Reference_Constants_v8.docx")
    print("  stage 1 output.docx  (existing)")
    print("  stage 2 output.docx  (existing)")
    print("  stage 3 output.docx  (existing)")
    print("  PROMPT_Output_Quality_Review.md  (copy as prompt text)")


if __name__ == "__main__":
    main()
